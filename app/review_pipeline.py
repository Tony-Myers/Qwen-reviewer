#!/usr/bin/env python3
"""
Local peer-review pipeline using Qwen.

The LLM is reached through llm_backend, which serves GGUF files via llama.cpp
and MLX repo ids via mlx_lm. The default model is the Qwen3.8-27B GGUF; the
previously used MLX models remain selectable with --model.

Architecture:
  1. Document parsing  (PDF/docx/csv/xlsx/txt/md, optional Marker)
  2. Evidence structuring  (typed blocks: TABLE, MODEL, DESIGN, RESULT, ...)
  3. Method classification  (rule-based framework detection)
  4. Method-aware critique  (framework-specific prompt expectations)
  5. Synthesis + validation  (LLM synthesis, programmatic + LLM post-checks)
  6. Output  (review report + evidence appendix with typed blocks)
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import os
import re
import sys
import math
import unicodedata
from contextlib import contextmanager
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum, auto
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

import pdfplumber
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

# The LLM backend is a drop-in replacement for the mlx_lm API: it exposes
# load/generate/make_sampler with identical signatures and dispatches to
# llama.cpp (for .gguf models) or mlx_lm (for MLX repo ids).
sys.path.insert(0, str(Path(__file__).resolve().parent))

import design_expectations as de  # noqa: E402
import manuscript_checks as mchk  # noqa: E402

from llm_backend import (  # noqa: E402
    available_context,
    BackendError,
    current_backend,
    generate,
    load,
    make_sampler,
    set_backend,
    strip_reasoning,
    thinking_enabled,
    thinking_token_allowance,
)


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Default model: Qwen3.8-27B, Unsloth dynamic 4-bit GGUF, served by llama.cpp.
QWEN38_27B_GGUF = str(
    Path.home()
    / ".cache" / "huggingface" / "hub"
    / "models--unsloth--Qwen3.8-27B-GGUF"
    / "snapshots" / "4ca720788d1e01f1bff70c033e0d0028fd02e502"
    / "Qwen3.8-27B-UD-Q4_K_XL.gguf"
)

# Previous MLX models, still selectable with --model.
QWEN36_35B_MLX = "mlx-community/Qwen3.6-35B-A3B-4bit"
QWEN36_27B_MLX = "mlx-community/Qwen3.6-27B-6bit"
QWEN36_27B_4BIT_MLX = "mlx-community/Qwen3.6-27B-4bit"
QWEN36_27B_8BIT_MLX = "mlx-community/Qwen3.6-27B-8bit"
GEMMA4_26B_MLX = "mlx-community/gemma-4-26b-a4b-it-4bit"
GEMMA4_31B_MLX = "mlx-community/gemma-4-31b-it-4bit"

# Short names accepted by --model, matching the launcher and the web UI.
# This is the single source of truth: server.py folds these in as well, so the
# same alias means the same model whichever entry point is used.
MODEL_ALIASES = {
    "qwen38": QWEN38_27B_GGUF,
    "qwen38-27b": QWEN38_27B_GGUF,
    "27b-gguf": QWEN38_27B_GGUF,
    "gguf": QWEN38_27B_GGUF,
    "35b": QWEN36_35B_MLX,
    "35b-4bit": QWEN36_35B_MLX,
    "qwen35": QWEN36_35B_MLX,
    "27b": QWEN36_27B_MLX,
    "27b-6bit": QWEN36_27B_MLX,
    "qwen27": QWEN36_27B_MLX,
    "27b-4bit": QWEN36_27B_4BIT_MLX,
    "27b-8bit": QWEN36_27B_8BIT_MLX,
    "gemma4": GEMMA4_26B_MLX,
    "gemma4-26b": GEMMA4_26B_MLX,
    "gemma4-26b-it": GEMMA4_26B_MLX,
    "gemma4-31b": GEMMA4_31B_MLX,
    "gemma4-31b-it": GEMMA4_31B_MLX,
}


def resolve_model_alias(value: str) -> str:
    """
    Expand a short model alias to a .gguf path or an MLX repo id.

    Anything that is not a known alias is returned unchanged, so full repo ids
    and absolute paths still work. Without this, an alias such as "35b" was
    passed to the loader verbatim and treated as a Hugging Face repo id.
    """
    if not value:
        return value
    return MODEL_ALIASES.get(value.strip().lower(), value)


def list_model_aliases() -> str:
    """Human-readable alias listing for --list-models."""
    grouped: Dict[str, List[str]] = {}
    for alias, target in MODEL_ALIASES.items():
        grouped.setdefault(target, []).append(alias)
    lines = ["Model aliases accepted by --model:", ""]
    for target, aliases in grouped.items():
        backend = "llama.cpp" if target.lower().endswith(".gguf") else "MLX"
        lines.append(f"  {', '.join(sorted(aliases))}")
        lines.append(f"      -> {model_display_name(target)}  [{backend}]")
    lines.append("")
    lines.append("A full Hugging Face repo id or an absolute path to a .gguf "
                 "file is also accepted.")
    return "\n".join(lines)


MODEL_NAME = os.environ.get("MODEL_NAME") or QWEN38_27B_GGUF
OUTPUT_DOMAIN = "general"


def model_display_name(name: str = None) -> str:
    """Short label for reports: the file name for a GGUF, the id otherwise."""
    name = MODEL_NAME if name is None else name
    return Path(name).name if name.lower().endswith(".gguf") else name

MAX_SECTION_CHARS = 7000
SECTION_MAX_TOKENS = 900
FILE_SYNTHESIS_MAX_TOKENS = 1200
# Raised from 1600 when strengths gained an Evidence line. Confidence and the
# item table are computed after generation and cost nothing here.
SYNTHESIS_MAX_TOKENS = int(os.environ.get("QWEN_SYNTHESIS_MAX_TOKENS", "2000"))
VALIDATION_MAX_TOKENS = 1800

# Sampling. Overridable from the environment so a configuration can be compared
# against another without editing the file -- see tests/sampler_sweep.py.
#
# Unsloth publish 0.7 / 0.80 / 20 for Qwen3.8 in non-thinking mode. top_p and
# top_k already agree; the temperature is deliberately lower for a factual task,
# and REPEAT_PASS_TEMPERATURE below restores their value for the extra passes
# that exist to find what a cautious pass misses.
#
# repetition_penalty is sent explicitly at 1.0. This is NOT a quality change:
# llama-server's /props reports 1.0 as its own default, so this sends the value
# already in force. It is here so the setting is stated rather than inherited,
# and cannot change under us when llama.cpp changes a default.
#
# It is worth recording why, because the reasoning was wrong before it was
# right. A penalty above 1.0 would discourage reproducing tokens, which is the
# opposite of what verbatim quotation needs, and the pipeline's habit of
# paraphrasing instead of quoting made that an appealing explanation. A sweep
# over four papers appeared to confirm it -- until /props showed the default was
# already 1.0 and the comparison had been an A/A test all along. The paraphrasing
# is not a sampler artefact.
#
# presence_penalty stays unset (server default 0.0). Unsloth suggest 1.5 for
# chat, where repeating yourself is the failure; here the failure is the
# opposite, and a presence penalty pushes against quoting the source.
TEMPERATURE = float(os.environ.get("QWEN_TEMPERATURE", "0.2"))
TOP_P = float(os.environ.get("QWEN_TOP_P", "0.8"))
TOP_K = int(os.environ.get("QWEN_TOP_K", "20"))
REPETITION_PENALTY = float(os.environ.get("QWEN_REPETITION_PENALTY", "1.0"))
PRESENCE_PENALTY = (float(os.environ["QWEN_PRESENCE_PENALTY"])
                    if os.environ.get("QWEN_PRESENCE_PENALTY") else None)

MAX_TABLE_ROWS = 80
MAX_TABLE_COLS = 14
MAX_TABLE_CHARS_PER_CELL = 120
MAX_APPENDIX_CHUNK_PREVIEW = 1800


# ---------------------------------------------------------------------------
# Phase 2: Evidence block types
# ---------------------------------------------------------------------------

class BlockType(Enum):
    TABLE = auto()
    MODEL = auto()
    DESIGN = auto()
    RESULT = auto()
    FIGURE_CAPTION = auto()
    NARRATIVE = auto()


@dataclass
class EvidenceBlock:
    """A typed unit of extracted evidence with source tracking."""
    block_type: BlockType
    text: str
    source_name: str
    page: Optional[int] = None
    label: Optional[str] = None
    confidence: str = "high"  # high | medium | low
    model_method_class: Optional["MethodClass"] = None


# Appended wherever this pipeline shows the model a shortened piece of the
# manuscript. A preview is for orientation, not for quotation: the cut can land
# mid-word, and a quotation taken from one cannot be verified against the paper
# because the fragment is not in the paper.
PREVIEW_TRUNCATION_MARKER = " [PREVIEW CUT SHORT - DO NOT QUOTE FROM HERE]"


@dataclass
class EvidenceManifest:
    """Summary of all structured evidence available for a single file."""
    source_name: str
    blocks: List[EvidenceBlock] = field(default_factory=list)
    method_class: Optional["MethodClass"] = None
    additional_method_classes: List["MethodClass"] = field(default_factory=list)
    # Study design is orthogonal to analysis method and neither replaces the
    # other: a meta-analysis may pool with REML or with a Bayesian hierarchical
    # model, and the method expectations remain correct for each.
    design_class: Optional["de.DesignClass"] = None
    design_signals: List[str] = field(default_factory=list)
    synthesis_method_class: Optional["MethodClass"] = None
    table_labels: List[str] = field(default_factory=list)
    has_equations: bool = False
    has_model_spec: bool = False
    has_effect_sizes: bool = False
    has_confidence_intervals: bool = False
    has_p_values: bool = False
    has_sample_description: bool = False
    has_randomisation: bool = False
    has_standard_errors: bool = False
    has_variance_components: bool = False
    has_model_fit_stats: bool = False
    n_tables: int = 0
    n_figures_mentioned: int = 0

    def summary_text(self) -> str:
        """Human-readable manifest for injection into review prompts."""
        lines = [f"Evidence manifest for: {self.source_name}"]
        lines.append(f"  Method classification: {self.method_class.value if self.method_class else 'unclassified'}")
        if self.design_class is not None:
            lines.append(f"  Study design: {self.design_class.value}"
                         + (f" (signals: {', '.join(self.design_signals)})"
                            if self.design_signals else ""))
        if self.synthesis_method_class is not None:
            lines.append(
                "  Analysis performed by the review itself: "
                f"{self.synthesis_method_class.value} (scoped to the methods "
                "section; method words elsewhere may belong to the included "
                "studies)")
        if self.additional_method_classes:
            lines.append(
                "  Additional method classifications: "
                + ", ".join(m.value for m in self.additional_method_classes)
            )
        lines.append(f"  Tables extracted: {self.n_tables} ({', '.join(self.table_labels) if self.table_labels else 'none labelled'})")
        lines.append(f"  Figures mentioned in text: {self.n_figures_mentioned}")
        lines.append(f"  Model specification present: {self.has_model_spec}")
        lines.append(f"  Equations present: {self.has_equations}")
        lines.append(f"  Standard errors reported: {self.has_standard_errors}")
        lines.append(f"  Variance components reported: {self.has_variance_components}")
        lines.append(f"  Model fit statistics reported: {self.has_model_fit_stats}")
        lines.append(f"  Effect sizes reported: {self.has_effect_sizes}")
        lines.append(f"  Confidence intervals reported: {self.has_confidence_intervals}")
        lines.append(f"  P-values reported: {self.has_p_values}")
        lines.append(f"  Sample/participants described: {self.has_sample_description}")
        lines.append(f"  Randomisation described: {self.has_randomisation}")

        type_counts = {}
        for b in self.blocks:
            type_counts[b.block_type.name] = type_counts.get(b.block_type.name, 0) + 1
        lines.append(f"  Block counts: {type_counts}")

        return "\n".join(lines)
        
        
    def model_block_summary(self) -> str:
        """
        Summarise model/result blocks and any model-level classifications found.

        These are previews: a block is cut at 280 characters, and a cut can land
        mid-word. One report quoted such a preview and produced
        "finding the qualitative inte[rsion] (text truncated in extraction)" --
        an honest quotation of a fragment that exists nowhere in the manuscript,
        which the citation check then reported as an unverifiable quotation.
        Truncated previews are now marked, so the model has somewhere to look
        rather than something to quote.
        """
        lines = []
        for b in self.blocks:
            if b.block_type not in {BlockType.MODEL, BlockType.RESULT}:
                continue

            method_label = b.model_method_class.value if b.model_method_class else "unclassified_model_block"
            page_text = f"Page {b.page}" if b.page is not None else "No page"
            preview = re.sub(r"\s+", " ", b.text.strip())
            if len(preview) > 280:
                preview = preview[:280] + PREVIEW_TRUNCATION_MARKER

            lines.append(f"- [{page_text}] {method_label}: {preview}")

        if not lines:
            return "No model/result blocks classified."

        return "\n".join(lines)
# ---------------------------------------------------------------------------
# Phase 3: Method classification
# ---------------------------------------------------------------------------

class MethodClass(Enum):
    BAYESIAN_MIXED_EFFECTS = "bayesian_mixed_effects"
    BAYESIAN_REGRESSION = "bayesian_regression"
    BAYESIAN_QUANTILE_REGRESSION = "bayesian_quantile_regression"
    FREQUENTIST_PROFILE_MODEL = "frequentist_profile_model"
    FREQUENTIST_ALLOMETRIC = "frequentist_allometric"
    FREQUENTIST_MIXED_EFFECTS = "frequentist_mixed_effects"
    FREQUENTIST_ROBUST = "frequentist_robust"
    FREQUENTIST_REGRESSION = "frequentist_regression"
    FREQUENTIST_ANOVA = "frequentist_anova"
    FREQUENTIST_SIMPLE_COMPARISON = "frequentist_simple_comparison"
    FREQUENTIST_CORRELATION = "frequentist_correlation"
    CROSS_CLASSIFIED_MCMC = "cross_classified_mcmc"
    MULTILEVEL_MCMC = "multilevel_mcmc"
    BAYESIAN_MODEL = "bayesian_model"
    DISTRIBUTIONAL_MODEL = "distributional_model"
    PREDICTIVE_ML = "predictive_ml"
    DESCRIPTIVE = "descriptive"
    MIXED_METHODS = "mixed_methods"
    UNCLASSIFIED = "unclassified"
    

# Keyword sets for classification — order matters (more specific first)
_METHOD_PATTERNS: List[Tuple[MethodClass, List[str], List[str]]] = [
    # (class, required_any, exclude_any)

    (MethodClass.CROSS_CLASSIFIED_MCMC, [
        r"cross.?classif",
        r"\bmlwin\b",
        r"\bmcmc\b",
        r"\bdic\b",
        ], [
        r"\bprior[s]?\b",
        r"\bposterior[s]?\b",
        r"\bcredible\s+interval\b",
        r"\bposterior\s+predict",
        r"\br-hat\b",
        r"\beffective\s+sample\s+size\b",
    ]),
    
    (MethodClass.MULTILEVEL_MCMC, [
        r"\bmlwin\b",
        r"multilevel.*mcmc",
        r"hierarchical.*mcmc",
        r"mcmc.*multilevel",
        r"mcmc.*hierarchical",
        r"\bdic\b",
    ], [
        r"\bprior[s]?\b",
        r"\bposterior[s]?\b",
        r"\bcredible\s+interval\b",
        r"\bposterior\s+predict",
        r"\br-hat\b",
        r"\beffective\s+sample\s+size\b",
    ]),
    
        (MethodClass.BAYESIAN_MIXED_EFFECTS, [
        # These patterns genuinely require hierarchical/multilevel structure
        r"\bposterior\b.*\brandom\s+intercept",
        r"\bposterior\b.*\brandom\s+slope",
        r"\bcredible\s+interval\b.*\brandom\s+intercept",
        r"\bcredible\s+interval\b.*\brandom\s+slope",
        r"\bprior[s]?\b.*\bmultilevel\b",
        r"\bprior[s]?\b.*\bhierarch(?:ical|y)\b",
        r"\bgroup[\-\s]?level\s+(?:effect|variance|intercept|slope)",
        r"\brandom\s+(?:intercept|slope|effect).*\bprior",
        r"\bbrms\b.*\brandom\s+(?:intercept|slope|effect)",
        r"\brstanarm\b.*\brandom\s+(?:intercept|slope|effect)",
        r"\b(?:brms|rstanarm)\b.*\b(?:multilevel|hierarchic)",
    ], []),

    # Bayesian quantile regression (more specific, check before generic Bayesian regression)
    (MethodClass.BAYESIAN_QUANTILE_REGRESSION, [
        r"\bbayesian\s+quantile\s+regress",
        r"\bquantile\s+regress.*\bprior",
        r"\bquantile\s+regress.*\bposterior",
        r"\bquantile\s+regress.*\bcredible\s+interval",
        r"\basymmetric\s+laplace",
        r"\bbrqr\b",
    ], []),

    # Bayesian regression (non-hierarchical) — brms/rstanarm without random effects
    (MethodClass.BAYESIAN_REGRESSION, [
        r"\bbayesian\s+(?:linear\s+)?regress",
        r"\bbayesian\s+(?:log[\-\s]?linear|logistic|allometric)\s+(?:model|regress)",
        r"\bbrms\b",
        r"\brstanarm\b",
    ], [
        # Exclude if hierarchical/multilevel signals are present
        r"\brandom\s+intercept",
        r"\brandom\s+slope",
        r"\bmultilevel\b.*\bprior",
        r"\bhierarch(?:ical|y)\b.*\bprior",
        r"\bgroup[\-\s]?level\s+effect",
    ]),
    
 (MethodClass.BAYESIAN_MODEL, [
    r"\bprior[s]?\b.*\bposterior[s]?\b",
    r"\bposterior\s+predict",
    r"\bcredible\s+interval\b",
    r"\bhdi\b",
    r"\brope\b",
    r"bayes\s*factor",
    r"\bposterior\s+distribution\b",
    r"\bposterior\s+mean\b",
    r"\bposterior\s+median\b",
], [
    r"\bmlwin\b.*\bmcmc\b",
    r"\bcross.?classif",
    r"\bdic\b",
]),
    (MethodClass.DISTRIBUTIONAL_MODEL, [
        r"\bgamlss\b",
        r"distributional\s+model",
        r"distributional\s+regression",
        r"\bvgam\b",
        r"location.?scale.?shape",
        r"randomi[sz]ed\s+quantile\s+residual",
        r"worm\s+plot",
    ], []),

    (MethodClass.FREQUENTIST_PROFILE_MODEL, [
        r"torque[–\- ]cadence",
        r"power[–\- ]cadence",
        r"second[ -]?order polynomial",
        r"optimal cadence",
        r"maximum cadence",
        r"maximum torque",
        r"modelled maximum power",
        r"\brpmopt\b",
        r"\brpmmax\b",
        r"\btmax\b",
        r"\bpmax\b",
    ], []),

    (MethodClass.FREQUENTIST_ALLOMETRIC, [
        r"\ballometric\b",
        r"multiplicative\s+model",
        r"gamma\s+function",
        r"log[-\s]?transformed",
        r"\bancova\b",
        r"maturity\s+offset",
        r"arm\s+span",
    ], []),

    (MethodClass.FREQUENTIST_MIXED_EFFECTS, [
        r"mixed.?effects?\s+model",
        r"linear\s+mixed",
        r"generalized\s+linear\s+mixed",
        r"\bglmm\b",
        r"\blmm\b",
        r"\blme4\b",
        r"\bnlme\b",
        r"\brandom\s+intercept",
        r"\brandom\s+slope",
        r"restricted\s+maximum\s+likelihood",
        r"\bREML\b",
        r"subjects?\s+were\s+included\s+as\s+a\s+random\s+effect",
        r"subjects?\s+included\s+as\s+a\s+random\s+effect",
    ], [
        r"\bprior\b",
        r"\bposterior\b",
        r"\bcredible\s+interval",
        r"\br-hat\b",
    ]),

    (MethodClass.FREQUENTIST_ROBUST, [
        r"yuen\s+paired\s*t",
        r"yuen\s+test",
        r"trimmed\s+mean",
        r"trimmed\s+means",
        r"winsori[sz]ed",
        r"\bwrs2\b",
        r"benjamini",
        r"hochberg",
        r"\bs\s+value\b",
        r"surprisal",
        r"robust\s+standardi[sz]ed\s+effect\s+size",
    ], []),

    (MethodClass.FREQUENTIST_REGRESSION, [
        r"linear\s+regress",
        r"logistic\s+regress",
        r"multiple\s+regress",
        r"\bols\b",
        r"\bglm\b(?!.*gamlss)",
        r"generalised\s+linear",
        r"generaliz(?:ed|ed)\s+linear",
    ], [
        # Do not classify as frequentist regression if Bayesian signals present
        r"\bprior[s]?\b.*\bposterior[s]?\b",
        r"\bcredible\s+interval",
        r"\bbrms\b",
        r"\brstanarm\b",
        r"\bposterior\s+predict",
    ]),

    (MethodClass.FREQUENTIST_ANOVA, [
        r"\banova'?s?\b",
        r"\bancova'?s?\b",
        r"\bmanova'?s?\b",
        r"\bmancova'?s?\b",
        r"repeated\s*measures",
        r"factorial\s+design",
        r"between.?subject.*within.?subject",
    ], []),

    (MethodClass.FREQUENTIST_SIMPLE_COMPARISON, [
        r"\bt-?test\b",
        r"paired\s+t-?test",
        r"independent\s+t-?test",
        r"\bwelch\b",
        r"mann[ -]?whitney",
        r"\bwilcoxon\b",
        r"\bkruskal[ -]?wallis\b",
        r"\bcohen'?s\s+d\b",
    ], []),

    (MethodClass.FREQUENTIST_CORRELATION, [
        r"\bpearson'?s?\b",
        r"\bspearman'?s?\b",
        r"\bcorrelation\b",
        r"\bcorrelated\b",
        r"\br\s*=",
        r"\brho\b",
    ], [
        r"linear\s+regress",
        r"logistic\s+regress",
        r"mixed.?effects?\s+model",
    ]),

    (MethodClass.PREDICTIVE_ML, [
        r"\bcross.?valid",
        r"\broc\b",
        r"\bauc\b",
        r"\bbrier\b",
        r"random\s*forest",
        r"\bxgboost\b",
        r"\blasso\b",
        r"\belastic\s*net\b",
        r"neural\s*net",
        r"deep\s*learn",
        r"train.*test\s*split",
        r"hyperparamet",
    ], []),

    (MethodClass.DESCRIPTIVE, [
        r"descriptive\s+statistic",
        r"descriptive\s+analys",
        r"prevalence\s+stud",
        r"cross.?sectional\s+surv",
    ], []),
]

# Priority ordering for method classification: lower number = higher priority.
# When multiple method patterns match, the most complex/specific one wins.
# This prevents a paper using correlations + ANCOVA + MANOVA + GAMLSS
# from collapsing to 'frequentist_correlation'.
_METHOD_PRIORITY: Dict[MethodClass, int] = {
    MethodClass.CROSS_CLASSIFIED_MCMC: 1,
    MethodClass.MULTILEVEL_MCMC: 2,
    MethodClass.BAYESIAN_MIXED_EFFECTS: 3,
    MethodClass.BAYESIAN_QUANTILE_REGRESSION: 4,
    MethodClass.BAYESIAN_REGRESSION: 5,
    MethodClass.BAYESIAN_MODEL: 6,
    MethodClass.DISTRIBUTIONAL_MODEL: 7,
    MethodClass.FREQUENTIST_MIXED_EFFECTS: 8,
    MethodClass.FREQUENTIST_ROBUST: 9,
    MethodClass.FREQUENTIST_ALLOMETRIC: 10,
    MethodClass.FREQUENTIST_PROFILE_MODEL: 11,
    MethodClass.PREDICTIVE_ML: 12,
    MethodClass.FREQUENTIST_ANOVA: 13,
    MethodClass.FREQUENTIST_REGRESSION: 14,
    MethodClass.FREQUENTIST_CORRELATION: 15,
    MethodClass.FREQUENTIST_SIMPLE_COMPARISON: 16,
    MethodClass.DESCRIPTIVE: 17,
    MethodClass.MIXED_METHODS: 18,
    MethodClass.UNCLASSIFIED: 99,
}


def classify_method(full_text: str) -> MethodClass:
    """
    Rule-based classification of the primary statistical method.

    Strategy:
    1. Strip references/bibliography to avoid false matches on cited methods.
    2. Apply specific early-routing rules for complex designs that need
       multi-keyword conjunction tests (these are hard to express as single
       regex patterns in _METHOD_PATTERNS).
    3. Iterate _METHOD_PATTERNS to collect ALL matching method classes.
    4. Return the highest-priority match, where priority favours more complex
       / specific methods over simpler ones.  This prevents a paper that uses
       correlations AND ANCOVA AND MANOVA from collapsing to
       'frequentist_correlation'.
    """
    text_lower = full_text.lower()

    # Ignore references/bibliography when classifying the statistical method.
    #
    # Splitting at the FIRST occurrence of the bare word was wrong on any paper
    # that uses it in the body. RJSP-2025-0796 says "all identified references
    # were imported into the online Rayyan application" at 15% of the document,
    # so the classifier saw 10,258 of 67,582 characters and returned
    # unclassified. One paper in the development corpus kept only 7% of its
    # text. de.strip_back_matter() looks for the word as a section heading in
    # the back of the document instead, and keeps everything when it cannot
    # find one, because retaining a reference list is a smaller error than
    # discarding a results section. All sixteen corpus classifications are
    # unchanged by this.
    text_for_classification = de.strip_back_matter(full_text)

    # --- Early routing for designs that require multi-keyword conjunctions ---

    # Explicit frequentist mixed-effects
    if (
        re.search(r"mixed.?effects?\s+model", text_for_classification)
        and re.search(r"restricted\s+maximum\s+likelihood|\bREML\b", text_for_classification)
    ):
        return MethodClass.FREQUENTIST_MIXED_EFFECTS

    if (
        re.search(r"subjects?\s+were\s+included\s+as\s+a\s+random\s+effect", text_for_classification)
        and re.search(r"\bintervention\b", text_for_classification)
        and re.search(r"\btime\b", text_for_classification)
    ):
        return MethodClass.FREQUENTIST_MIXED_EFFECTS

    # Cross-classified MCMC
    if (
        re.search(r"cross.?classif", text_for_classification)
        and re.search(r"\bmlwin\b", text_for_classification)
        and re.search(r"\bmcmc\b", text_for_classification)
    ):
        return MethodClass.CROSS_CLASSIFIED_MCMC

    # Multilevel MCMC
    if (
        re.search(r"\bmultilevel\b", text_for_classification)
        and re.search(r"\bmultinomial\s+logit\b", text_for_classification)
        and re.search(r"\bmlwin\b", text_for_classification)
        and re.search(r"\bmcmc\b", text_for_classification)
    ):
        return MethodClass.MULTILEVEL_MCMC

    # --- Iterate _METHOD_PATTERNS to find all matching classes ---
    matched: List[MethodClass] = []
    for method_class, required_any, exclude_any in _METHOD_PATTERNS:
        # At least one required pattern must match
        if not any(re.search(pat, text_for_classification) for pat in required_any):
            continue
        # None of the exclusion patterns may match
        if exclude_any and any(re.search(pat, text_for_classification) for pat in exclude_any):
            continue
        matched.append(method_class)

    if not matched:
        return MethodClass.UNCLASSIFIED

    if len(matched) == 1:
        return matched[0]

    # Multiple matches: return the highest-priority (most complex/specific)
    return min(matched, key=lambda m: _METHOD_PRIORITY.get(m, 99))
    
def classify_additional_methods(full_text: str, primary: MethodClass) -> List[MethodClass]:
    """
    Identify additional analysis types present in the paper, beyond the primary
    classification. This is useful for papers that combine multiple model types.
    """
    text_lower = full_text.lower()
    text_for_classification = re.split(
        r"\b(?:references|bibliography|acknowledgments?)\b",
        text_lower,
        maxsplit=1,
    )[0]

    found: List[MethodClass] = []

    if (
        re.search(r"cross.?classif(?:ied|ication)?", text_for_classification)
        and re.search(r"\bmlwin\b", text_for_classification)
        and re.search(r"\bmcmc\b", text_for_classification)
    ):
        if primary != MethodClass.CROSS_CLASSIFIED_MCMC:
            found.append(MethodClass.CROSS_CLASSIFIED_MCMC)

    if (
        re.search(r"\bmultilevel\b", text_for_classification)
        and re.search(r"\bmultinomial\s+logit\b", text_for_classification)
        and re.search(r"\bmlwin\b", text_for_classification)
        and re.search(r"\bmcmc\b", text_for_classification)
    ):
        if primary != MethodClass.MULTILEVEL_MCMC:
            found.append(MethodClass.MULTILEVEL_MCMC)

    if (
        re.search(r"\bmixed.?effects?\s+model", text_for_classification)
        or re.search(r"\blme4\b|\bnlme\b|\bglmm\b|\blmm\b", text_for_classification)
    ):
        if primary != MethodClass.FREQUENTIST_MIXED_EFFECTS:
            found.append(MethodClass.FREQUENTIST_MIXED_EFFECTS)

    if (
        re.search(r"\bbrms\b|\brstanarm\b", text_for_classification)
        or (
            re.search(r"\bprior[s]?\b", text_for_classification)
            and re.search(r"\bposterior[s]?\b", text_for_classification)
        )
    ):
        # Determine which Bayesian sub-type
        has_hierarchical = bool(re.search(
            r"\brandom\s+(?:intercept|slope|effect)"
            r"|\bmultilevel\b.*\bprior"
            r"|\bhierarch(?:ical|y)\b.*\bprior"
            r"|\bgroup[\-\s]?level\s+(?:effect|variance)",
            text_for_classification,
        ))
        has_quantile = bool(re.search(
            r"\bquantile\s+regress|\basymmetric\s+laplace|\bbrqr\b",
            text_for_classification,
        ))

        if has_hierarchical and primary != MethodClass.BAYESIAN_MIXED_EFFECTS:
            found.append(MethodClass.BAYESIAN_MIXED_EFFECTS)
        elif has_quantile and primary != MethodClass.BAYESIAN_QUANTILE_REGRESSION:
            found.append(MethodClass.BAYESIAN_QUANTILE_REGRESSION)
        elif re.search(r"\bregress|\ballometric|\blogistic|\blog[\-\s]?linear", text_for_classification):
            if primary != MethodClass.BAYESIAN_REGRESSION:
                found.append(MethodClass.BAYESIAN_REGRESSION)
        elif primary not in (MethodClass.BAYESIAN_MIXED_EFFECTS, MethodClass.BAYESIAN_MODEL,
                             MethodClass.BAYESIAN_REGRESSION, MethodClass.BAYESIAN_QUANTILE_REGRESSION):
            found.append(MethodClass.BAYESIAN_MODEL)

    # ANOVA / ANCOVA / MANOVA / MANCOVA
    if re.search(
        r"\banova'?s?\b|\bancova'?s?\b|\bmanova'?s?\b|\bmancova'?s?\b|"
        r"willk'?s\s+lambda|wilks'?\s+lambda|"
        r"repeated\s*measures",
        text_for_classification,
    ):
        if primary != MethodClass.FREQUENTIST_ANOVA:
            found.append(MethodClass.FREQUENTIST_ANOVA)

    # GAMLSS / distributional modelling
    if re.search(
        r"\bgamlss\b|distributional\s+model|distributional\s+regression|"
        r"\bgaic\b|location.?scale.?shape",
        text_for_classification,
    ):
        if primary != MethodClass.DISTRIBUTIONAL_MODEL:
            found.append(MethodClass.DISTRIBUTIONAL_MODEL)

    # Correlation — only flag as additional if it appears to be a substantive
    # analysis method, not just mentioned in background or as a descriptive
    # statistic.  Exclude when primary is Bayesian (correlation terms often
    # appear in literature reviews without being the paper's own analysis).
    if primary not in (
        MethodClass.BAYESIAN_MIXED_EFFECTS, MethodClass.BAYESIAN_REGRESSION,
        MethodClass.BAYESIAN_QUANTILE_REGRESSION, MethodClass.BAYESIAN_MODEL,
    ):
        # Require stronger signals: "correlation analysis", "correlations were
        # calculated/computed", or a methods-section context — not just
        # "Pearson" appearing in a literature review paragraph.
        if re.search(
            r"correlation\s+(?:analysis|coefficient\s+was|coefficients?\s+were|was\s+(?:used|calculated|computed))"
            r"|\bcorrelations?\s+(?:were|was)\s+(?:calculated|computed|used|performed|conducted)"
            r"|(?:we\s+)?(?:used|conducted|performed|calculated|computed)\s+(?:simple\s+)?correlations?",
            text_for_classification,
        ):
            if primary != MethodClass.FREQUENTIST_CORRELATION:
                found.append(MethodClass.FREQUENTIST_CORRELATION)

    # Regression (linear, logistic, multiple) — only flag as additional if
    # the primary is not already a Bayesian class (since Bayesian papers
    # naturally describe their model as "regression" without being frequentist).
    if primary not in (
        MethodClass.BAYESIAN_MIXED_EFFECTS, MethodClass.BAYESIAN_REGRESSION,
        MethodClass.BAYESIAN_QUANTILE_REGRESSION, MethodClass.BAYESIAN_MODEL,
    ):
        if re.search(
            r"linear\s+regress|logistic\s+regress|multiple\s+regress|\bols\b",
            text_for_classification,
        ):
            # Exclude if Bayesian terms are nearby (the word "regression" is
            # describing a Bayesian model, not a separate frequentist analysis)
            if not re.search(
                r"\bprior[s]?\b.*\bregress|\bregress.*\bprior[s]?\b"
                r"|\bposterior\b.*\bregress|\bregress.*\bposterior\b"
                r"|\bbayesian\s+(?:linear|logistic|quantile)?\s*regress"
                r"|\bbrms\b|\brstanarm\b",
                text_for_classification,
            ):
                if primary != MethodClass.FREQUENTIST_REGRESSION:
                    found.append(MethodClass.FREQUENTIST_REGRESSION)

    # Simple comparison (t-test, Wilcoxon, etc.)
    if re.search(
        r"\bt-?test\b|\bwilcoxon\b|\bmann[ -]?whitney\b|\bkruskal[ -]?wallis\b",
        text_for_classification,
    ):
        if primary != MethodClass.FREQUENTIST_SIMPLE_COMPARISON:
            found.append(MethodClass.FREQUENTIST_SIMPLE_COMPARISON)

    # Deduplicate while preserving order
    deduped: List[MethodClass] = []
    seen = set()
    for item in found:
        if item not in seen:
            deduped.append(item)
            seen.add(item)

    return deduped
    
def classify_model_block(block_text: str) -> Optional[MethodClass]:
    """
    Classify a single model-related block. This is deliberately narrower than
    whole-paper classification and is used to detect multiple analysis types
    within one manuscript.
    """
    text_lower = block_text.lower()

    has_mlwin = bool(re.search(r"\bmlwin\b", text_lower))
    has_mcmc = bool(re.search(r"\bmcmc\b", text_lower))
    has_dic = bool(re.search(r"\bdic\b|deviance\s+information\s+criter", text_lower))
    has_multinomial_logit = bool(re.search(r"\bmultinomial\s+logit\b", text_lower))
    has_cross_classified = bool(re.search(r"cross.?classif(?:ied|ication)?", text_lower))
    has_classified_by_respondent_problem = bool(
        re.search(r"classified\s+by\s+respondent\s+and\s+problem\s+type", text_lower)
    )
    has_nested_within_respondent = bool(
        re.search(r"nested\s+within\s+respondent", text_lower)
    )
    has_multilevel = bool(re.search(r"\bmultilevel\b|\bhierarchical\b", text_lower))

    # Most specific first
    if has_cross_classified and has_multinomial_logit:
        return MethodClass.CROSS_CLASSIFIED_MCMC

    if has_classified_by_respondent_problem and has_multinomial_logit:
        return MethodClass.CROSS_CLASSIFIED_MCMC

    if has_multilevel and has_multinomial_logit:
        return MethodClass.MULTILEVEL_MCMC

    if has_nested_within_respondent and has_multinomial_logit:
        return MethodClass.MULTILEVEL_MCMC

    # Secondary support for MLwiN / MCMC / DIC phrasing
    if (has_mlwin or has_mcmc or has_dic) and has_cross_classified:
        return MethodClass.CROSS_CLASSIFIED_MCMC

    if (has_mlwin or has_mcmc or has_dic) and (has_multilevel or has_nested_within_respondent):
        return MethodClass.MULTILEVEL_MCMC

    if (
        re.search(r"\bmixed.?effects?\s+model", text_lower)
        or re.search(r"\blme4\b|\bnlme\b|\bglmm\b|\blmm\b", text_lower)
    ):
        return MethodClass.FREQUENTIST_MIXED_EFFECTS

    if (
        re.search(r"\bbrms\b|\brstanarm\b", text_lower)
        or (
            re.search(r"\bprior[s]?\b", text_lower)
            and re.search(r"\bposterior[s]?\b", text_lower)
        )
    ):
        # Distinguish Bayesian sub-types at block level
        has_hierarchical = bool(re.search(
            r"\brandom\s+(?:intercept|slope|effect)"
            r"|\bmultilevel\b|\bhierarch(?:ical|y)\b"
            r"|\bgroup[\-\s]?level\s+(?:effect|variance)",
            text_lower,
        ))
        has_quantile = bool(re.search(
            r"\bquantile\s+regress|\basymmetric\s+laplace|\bbrqr\b",
            text_lower,
        ))

        if has_hierarchical:
            return MethodClass.BAYESIAN_MIXED_EFFECTS
        if has_quantile:
            return MethodClass.BAYESIAN_QUANTILE_REGRESSION
        if re.search(
            r"\bregress|\ballometric|\blinear\s+model|\blogistic|\blog[\-\s]?linear",
            text_lower,
        ):
            return MethodClass.BAYESIAN_REGRESSION
        return MethodClass.BAYESIAN_MODEL

    # GAMLSS / distributional models
    if re.search(
        r"\bgamlss\b|\bgaic\b|distributional\s+(?:model|regression)|"
        r"location.?scale.?shape|"
        r"box[\-\s]?cox[\-\s]?(?:t|cole|power)|"
        r"randomi[sz]ed\s+quantile\s+residual|worm\s+plot",
        text_lower,
    ):
        return MethodClass.DISTRIBUTIONAL_MODEL

    # MANOVA / MANCOVA (check before ANOVA since MANOVA is more specific)
    if re.search(
        r"\bmanova'?s?\b|\bmancova'?s?\b|willk'?s\s+lambda|wilks'?\s+lambda|"
        r"multivariate\s+(?:dependent\s+variable|analysis\s+of\s+variance)",
        text_lower,
    ):
        return MethodClass.FREQUENTIST_ANOVA

    # ANOVA / ANCOVA
    if re.search(
        r"\banova'?s?\b|\bancova'?s?\b|\bf[\-\s]?ratio|\bf\s*\(\s*\d+\s*,\s*\d+\s*\)|"
        r"partial\s+eta[\-\s]?squared|"
        r"between[\-\s]?subject.*within[\-\s]?subject|"
        r"repeated\s*measures?\s+(?:design|analysis|anova)",
        text_lower,
    ):
        return MethodClass.FREQUENTIST_ANOVA

    # Regression (linear, logistic, multiple)
    if re.search(
        r"linear\s+regress|logistic\s+regress|multiple\s+regress|"
        r"\bols\b|\bglm\b(?!.*gamlss)|"
        r"generali[sz]ed\s+linear\s+model",
        text_lower,
    ):
        return MethodClass.FREQUENTIST_REGRESSION

    # Correlation / association
    if re.search(
        r"\bpearson'?s?\s+correlat|\bspearman'?s?\s+correlat|"
        r"correlation\s+coefficien|"
        r"\br\s*=\s*[\-\d]",
        text_lower,
    ):
        return MethodClass.FREQUENTIST_CORRELATION

    # Simple comparison
    if re.search(
        r"\bt[\-\s]?test\b|\bwilcoxon\b|\bmann[\-\s]?whitney\b|"
        r"\bkruskal[\-\s]?wallis\b|\bcohen'?s\s+d\b",
        text_lower,
    ):
        return MethodClass.FREQUENTIST_SIMPLE_COMPARISON

    if has_multinomial_logit:
        return MethodClass.FREQUENTIST_REGRESSION

    return None
    
# Framework-specific review expectations loaded by method class
_METHOD_EXPECTATIONS: Dict[MethodClass, str] = {
MethodClass.BAYESIAN_MIXED_EFFECTS: """Method-specific expectations (Bayesian mixed-effects):
- Appropriate focus: prior specification and justification for both fixed and group-level effects, convergence diagnostics (R-hat, ESS, divergences where relevant), posterior predictive checks, and hierarchical model structure.
- Reporting: fixed effects, group-level variance/correlation components, credible intervals, model-comparison criteria where used, and clear description of random intercepts/slopes.
- Appropriate concerns: prior sensitivity, weakly identified variance components, and whether the multilevel structure matches the design.
- Do not request frequentist p-values or classical mixed-model diagnostics as if they were mandatory Bayesian outputs.
""",

MethodClass.FREQUENTIST_PROFILE_MODEL: """Method-specific expectations (frequentist profile / mechanistic model):
- Appropriate focus: transparency of fitted profiling equations, derivation of modelled parameters (e.g. optimal cadence, maximum torque, modelled maximum power), and robustness of curve fitting.
- Reporting: explicit equations, fitted coefficients where relevant, uncertainty for derived parameters where possible, and clear distinction between measured and modelled quantities.
- Appropriate concerns: filtering choices, outlier handling, sensitivity to sparse or low-frequency data, and whether model fit metrics are accompanied by enough detail to assess robustness.
- Do not treat the manuscript as a plain ANOVA paper if the core analysis is based on fitted torque-cadence or power-cadence relationships.
""",

MethodClass.FREQUENTIST_ALLOMETRIC: """Method-specific expectations (frequentist allometric / ANCOVA-regression):
- Appropriate focus: justification of the multiplicative/allometric form, interpretability of the log-linearised model, and clarity of gamma-function or growth/decay mass terms.
- Reporting: transformed model equations, coefficients with standard errors or confidence intervals, fit statistics (e.g. R-squared), and clear explanation of derived optima where reported.
- Appropriate concerns: whether the allometric and ANCOVA layers are clearly distinguished, whether derived optimal values have uncertainty, and whether the model form is biologically/statistically justified.
- Do not say the functional form is absent when equations and transformed model terms are explicitly reported.
""",

MethodClass.FREQUENTIST_MIXED_EFFECTS: """Method-specific expectations (frequentist mixed-effects):
- Appropriate focus: random-effects structure, repeated-measures dependence, variance components, and justification for random intercepts/slopes.
- Reporting: fixed-effect estimates, standard errors or confidence intervals, variance components, model-comparison criteria where relevant, and enough detail to understand the nesting/crossing structure.
- Appropriate concerns: singular fits, over-complex random-effects structures, and whether model assumptions are checked at the right level.
- Do not default to plain ANOVA or simple regression expectations when mixed-effects structure is explicit.
""",

MethodClass.FREQUENTIST_SIMPLE_COMPARISON: """Method-specific expectations (frequentist simple comparison):
- Appropriate focus: group-comparison design, test choice, assumption checks where relevant, effect sizes, and confidence intervals.
- Reporting: group means/medians, spread, exact or near-exact p-values where appropriate, and effect sizes.
- Appropriate concerns: multiplicity where many tests are run, whether paired vs unpaired structure was respected, and whether assumptions match the chosen test.
""",

MethodClass.FREQUENTIST_CORRELATION: """Method-specific expectations (frequentist correlation / association):
- Appropriate focus: strength and direction of association, uncertainty around the association, and whether the analysis is descriptive association rather than causal inference.
- Reporting: coefficient values, confidence intervals where available, sample size, and clear distinction between correlation and prediction/causation.
- Appropriate concerns: outliers, non-linearity, and over-interpretation of association as explanation.
""",
    
    MethodClass.FREQUENTIST_ROBUST: """Method-specific expectations (robust frequentist):
- Appropriate focus: trimming rationale, robustness to outliers/non-normality, clarity of paired or grouped comparison structure.
- Reporting: trimmed means or robust estimates, interval estimates, effect sizes, multiplicity correction if relevant.
- Prefer concerns about sensitivity to trimming choices, agreement versus association, and whether robust and conventional conclusions diverge.
- Do not default to mixed-model or variance-component expectations unless such models were actually used.
- Do not overemphasise regression diagnostics unless regression is central to the primary inferential claims.
- Do not request Bayesian diagnostics (R-hat, ESS, divergences, priors).
- Do not request prediction metrics unless the paper is explicitly framed as predictive.""",

    MethodClass.FREQUENTIST_REGRESSION: """Method-specific expectations (frequentist regression):
- Appropriate diagnostics: residual plots, multicollinearity checks (VIF where >1 predictor), influence diagnostics, normality of residuals.
- Reporting: coefficients with standard errors or confidence intervals, R-squared or adjusted R-squared, F-test or likelihood ratio test.
- Do not request Bayesian diagnostics (R-hat, ESS, divergences, priors).
- Do not request prediction metrics (AUC, calibration) unless the paper frames it as a prediction study.""",

    MethodClass.FREQUENTIST_ANOVA: """Method-specific expectations (frequentist ANOVA/ANCOVA):
- Appropriate diagnostics: homogeneity of variance (Levene), sphericity (Mauchly for RM), normality, Box's M for MANOVA.
- Reporting: F-values, degrees of freedom, p-values, effect sizes (partial eta-squared, Cohen's d, or similar).
- Post-hoc tests with multiplicity correction if applicable.
- Do not request Bayesian diagnostics or prediction metrics.""",

    MethodClass.CROSS_CLASSIFIED_MCMC: """Method-specific expectations (cross-classified model via MCMC):
- MCMC is an estimation method here, NOT a Bayesian analysis. Treat this as a frequentist multilevel model estimated via MCMC because standard ML/REML cannot handle cross-classified structures efficiently.
- Appropriate diagnostics: convergence (chain mixing, burn-in), DIC or similar information criterion for model comparison, variance component estimates with standard errors.
- Reporting: fixed effect estimates with standard errors, variance components, model comparison statistics (DIC), significance tests (Wald-type chi-square).
- Do NOT request: prior specifications, posterior distributions, credible intervals, Bayes factors, prior sensitivity analyses, R-hat, or any Bayesian reporting elements.
- Do NOT describe the framework as unclear, ambiguous, or hybrid.
- Do NOT request VIF for single-predictor models.
- Do NOT request residual plots as a standard expectation for MCMC-estimated models (convergence diagnostics serve a different role).""",

    MethodClass.MULTILEVEL_MCMC: """Method-specific expectations (multilevel model via MCMC):
- MCMC is an estimation method here, NOT necessarily a Bayesian analysis. If no priors, posteriors, or credible intervals are mentioned, treat as frequentist multilevel modelling using MCMC for estimation.
- Appropriate diagnostics: convergence assessment (burn-in, chain mixing), DIC for model comparison, variance partition coefficients.
- Reporting: fixed effects with standard errors, random effect variance components, model fit statistics.
- Do NOT request Bayesian diagnostics unless explicitly Bayesian language is present in the manuscript.
- Do NOT describe the framework as unclear if MCMC is used without Bayesian terminology.""",

    MethodClass.BAYESIAN_MODEL: """Method-specific expectations (Bayesian model):
- Appropriate diagnostics: prior specification and justification, prior sensitivity analysis, R-hat and effective sample size, divergent transitions (if HMC/NUTS), posterior predictive checks.
- Reporting: posterior summaries (means/medians with credible intervals), prior choices, model comparison (LOO, WAIC, Bayes factors as appropriate).
- Do not request frequentist p-values or confidence intervals.""",

    MethodClass.BAYESIAN_REGRESSION: """Method-specific expectations (Bayesian regression):
- Appropriate focus: prior specification and justification for regression coefficients, posterior summaries with credible intervals, convergence diagnostics (R-hat, ESS), posterior predictive checks.
- Reporting: regression coefficients with credible intervals, prior choices, model comparison (LOO, WAIC) where relevant, R-squared or Bayesian R-squared if reported.
- Do NOT request random-effects structure, variance components, or hierarchical model diagnostics unless the text explicitly describes multilevel/hierarchical structure.
- Do NOT assume that brms or rstanarm implies a mixed-effects model — both packages support simple regression.
- Do not request frequentist p-values or confidence intervals.""",

    MethodClass.BAYESIAN_QUANTILE_REGRESSION: """Method-specific expectations (Bayesian quantile regression):
- Appropriate focus: prior specification, choice of quantile(s), asymmetric Laplace likelihood or equivalent, posterior summaries with credible intervals at each quantile.
- Reporting: quantile-specific coefficients with credible intervals, convergence diagnostics, comparison across quantiles where relevant.
- Do NOT request random-effects structure or hierarchical model diagnostics unless the text explicitly describes multilevel/hierarchical structure.
- Do not request frequentist p-values or confidence intervals.
- Quantile regression does NOT split the sample. Every quantile model is fitted to the whole sample using an asymmetric loss; the tau = 0.1 and tau = 0.9 fits use all n observations, not the bottom or top ten percent. Never state or imply that a quantile was estimated from a subset, and never compute a per-quantile sample size such as "n = 59 x 0.1, so about 6 swimmers". That is a misreading of the method and authors will reject it.
- The legitimate small-sample concern is different and may be raised where it applies: estimates at extreme quantiles carry more posterior uncertainty than those at the median, because fewer observations are informative about the tail. Judge this from the width of the reported credible intervals at each quantile, not from an invented subsample size.""",

    MethodClass.DISTRIBUTIONAL_MODEL: """Method-specific expectations (distributional/GAMLSS model):
- These models simultaneously model location, scale, and shape parameters. Do not treat scale/shape submodels as unusual.
- Appropriate diagnostics: worm plots or Q-Q plots for distributional fit, randomised quantile residuals, GAIC for model/distribution selection.
- Reporting: coefficients for each distribution parameter, distribution family chosen with justification, model selection criteria.
- Do not request standard OLS residual plots. Do not request VIF unless multiple collinear predictors are modelled.
- Do not treat small coefficients on scale/shape parameters as data errors.""",

    MethodClass.PREDICTIVE_ML: """Method-specific expectations (predictive/ML model):
- Appropriate diagnostics: internal validation (cross-validation, bootstrap), calibration (calibration plots, Hosmer-Lemeshow), discrimination (AUC/C-statistic), decision curve analysis if clinical.
- Reporting: performance metrics with confidence intervals or standard errors, comparison with simpler benchmarks, description of feature engineering and hyperparameter tuning.
- Do not request causal interpretation of feature importance unless the paper frames it causally.""",

    MethodClass.DESCRIPTIVE: """Method-specific expectations (descriptive/epidemiological):
- Appropriate reporting: summary statistics with measures of spread, response rates, sampling strategy, representativeness.
- Do not request inferential diagnostics beyond what the study design supports.""",
}


def get_method_expectations(
    method_class: MethodClass,
    additional_classes: Optional[List["MethodClass"]] = None,
    design_class: Optional["de.DesignClass"] = None,
) -> str:
    """
    Return framework-specific review expectations for the primary method,
    plus any additional methods detected.  This ensures that a paper using
    e.g. correlation + ANCOVA + MANOVA + GAMLSS gets expectations for all
    of those frameworks rather than only the primary.
    """
    primary = _METHOD_EXPECTATIONS.get(method_class, """Method-specific expectations:
- Method not confidently classified. Apply general standards: appropriate diagnostics for the analysis type, clear reporting of estimates and uncertainty, and model assumptions addressed.
- Match your diagnostic expectations to the actual analysis used. Do not default to OLS assumptions.""")

    design_block = de.get_design_expectations(design_class) if design_class else ""

    if not additional_classes:
        return (primary + "\n" + design_block) if design_block else primary

    parts = [primary]
    if design_block:
        parts.append(design_block)
    seen = {method_class}
    for mc in additional_classes:
        if mc in seen:
            continue
        seen.add(mc)
        extra = _METHOD_EXPECTATIONS.get(mc)
        if extra:
            parts.append(f"\nAdditional method expectations ({mc.value}):\n{extra}")

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Evidence block classification (heuristic, not LLM)
# ---------------------------------------------------------------------------

# Patterns that signal specific block types
_TABLE_SIGNALS = re.compile(
    r"\[TABLE_START\]|^\s*\|.+\|.+\|", re.MULTILINE
)
_MODEL_SIGNALS = re.compile(
    r"(?i)(?:"
    r"model\s+equation|regression\s+model|fitted\s+model|model\s+specification"
    r"|(?:fixed|random)\s+effect|variance\s+component"
    r"|coefficient.*estimate|~\s*N\(|\\?beta_?\d"
    r"|equation\s+\d|eq\.\s*\d|formula:"
    r"|AIC\b|BIC\b|DIC\b|WAIC\b|LOO\b|GAIC\b"
    r"|R[\-\s]?squared|adjusted\s+R|deviance"
    r"|logit\(|log\(|link\s+function"
    r"|cross[\-\s]?classif.*model|multilevel\s+model|hierarchical\s+model"
    r"|MCMC.*model|model.*MCMC"
    r"|between[\-\s]?\w+\s+variance|within[\-\s]?\w+\s+variance"
    r"|ANCOVA\b|MANCOVA\b|ANOVA\b|MANOVA\b"
    r"|Willk'?s\s+Lambda|Wilks'?\s+Lambda"
    r"|GAMLSS\b|distributional\s+model|distributional\s+regression"
    r"|centile\s+curve|Box[\-\s]?Cox"
    r"|F[\-\s]?ratio|partial\s+eta[\-\s]?squared"
    r")"
)
_DESIGN_SIGNALS = re.compile(
    r"(?i)(?:"
    # Narrowed: removed generic terms like "condition", "procedure", "protocol"
    # that appear in results and discussion sections too
    r"participant[s]?\s+(?:were|gave|volunteer|complet|recruit)"
    r"|sample\s+size|recruit(?:ed|ment)|volunteer(?:ed|s)"
    r"|informed\s+consent|written\s+consent"
    r"|inclusion\s+criter|exclusion\s+criter|eligib"
    r"|randomi[sz](?:ed|ation)\s+(?:to|into|between)"
    r"|control\s+group|experimental\s+group"
    r"|between[\-\s]?subject.*design|within[\-\s]?subject.*design"
    r"|counterbalance"
    r"|ethical\s+approv|IRB|ethics\s+committee"
    r"|materials?\s+and\s+methods\b"
    r")"
)
_RESULT_SIGNALS = re.compile(
    r"(?i)(?:"
    r"p\s*[<=<]\s*0?\.\d|p\s*=\s*0?\.\d"
    r"|chi[\-\s]?square|chi2|χ2|χ²"
    r"|F\s*\(\s*\d|t\s*\(\s*\d|z\s*=\s*[\-\d]"
    r"|cohen['\u2019]?s\s*d|eta[\-\s]?squared|partial\s+eta"
    r"|odds\s+ratio|hazard\s+ratio|risk\s+ratio|relative\s+risk"
    r"|confidence\s+interval|CI\s*[=:\[]|95\s*%\s*CI"
    r"|standard\s+error|SE\s*=|SD\s*="
    r"|mean\s+difference|effect\s+size"
    r"|statistically\s+significant|non[\-\s]?significant"
    r"|β\s*=|b\s*=\s*[\-\d]"
    r"|the\s+results?\s+(?:show|indicat|suggest|support|reveal)"
    r"|difference\s+of\s+\d"
    r"|practical(?:ly)?\s+significant"
    r")"
)
_FIGURE_CAPTION_SIGNALS = re.compile(
    r"(?i)(?:^|\n)\s*(?:Fig(?:ure)?\.?\s+\d|FIGURE\s+\d)", re.MULTILINE
)

# Section headings that override paragraph-level classification
_SECTION_HEADING_HINTS = {
    BlockType.RESULT: re.compile(
        r"(?i)(?:^|\n)\s*(?:RESULTS?\b|FINDINGS?\b|CROSS[\-\s]?CLASSIFIED\s+MODEL"
        r"|PRACTICAL\s+SIGNIFICANCE|STATISTICAL\s+ANALYSIS)",
    ),
    BlockType.MODEL: re.compile(
        r"(?i)(?:^|\n)\s*(?:ANALYSIS\b|STATISTICAL\s+MODEL|MODEL\s+SPECIFICATION"
        r"|MODELL?ING\s+APPROACH|CROSS[\-\s]?CLASSIFIED\s+(?:MODEL|ANALYSIS)"
        r"|STATISTICAL\s+METHOD|CENTILE\s+CURVE|GAMLSS\b)",
    ),
    BlockType.DESIGN: re.compile(
        r"(?i)(?:^|\n)\s*(?:MATERIALS?\s+AND\s+METHODS\b|METHOD(?:S|OLOGY)\b"
        r"|PARTICIPANTS?\b|PROCEDURE\b|STUDY\s+DESIGN\b|DATA\s+COLLECTION\b)",
    ),
}


def classify_paragraph_block(text: str, page: Optional[int] = None) -> BlockType:
    """Classify a paragraph-level text segment into a block type."""
    # TABLE_START marker is definitive
    if _TABLE_SIGNALS.search(text):
        return BlockType.TABLE

    # Check for section heading overrides in the first 120 characters
    head = text[:120]
    for block_type, pattern in _SECTION_HEADING_HINTS.items():
        if pattern.search(head):
            return block_type

    if _FIGURE_CAPTION_SIGNALS.search(text) and len(text) < 600:
        return BlockType.FIGURE_CAPTION

    # Count signal density for competing types, with weighting
    model_hits = len(_MODEL_SIGNALS.findall(text))
    design_hits = len(_DESIGN_SIGNALS.findall(text))
    result_hits = len(_RESULT_SIGNALS.findall(text))

    # MODEL and RESULT signals are weighted more heavily because DESIGN
    # terms (participant, condition, procedure) appear across all sections,
    # while model specifications and test statistics are more localised
    scores = {
        BlockType.MODEL: model_hits * 2.0,
        BlockType.DESIGN: design_hits * 1.0,
        BlockType.RESULT: result_hits * 1.5,
    }
    best = max(scores, key=scores.get)
    if scores[best] >= 2:
        return best

    return BlockType.NARRATIVE


def structure_evidence(
    source_name: str,
    text: str,
    table_blocks: List[Tuple[int, str]],
) -> EvidenceManifest:
    """
    Convert raw extracted text and table blocks into typed EvidenceBlocks
    and build a manifest summarising what evidence is available.
    """
    manifest = EvidenceManifest(source_name=source_name)
    # Folded for detection only: the manuscript writes "p < 0.05" as "\U0001d45d < 0.05"
    # with a mathematical italic p, and the manifest reported "P-values
    # reported: False" on a paper whose results table is nothing but p-values.
    # The manifest is not cosmetic -- the prompt forbids the model from calling
    # a statistic absent when the manifest says it is present, so a false
    # negative there invites exactly the wrong criticism. The text itself is
    # left alone; only the copy these patterns read is folded.
    text_lower = fold_mathematical_letters(text).lower()

    # --- Structured table blocks (from pdfplumber/camelot) ---
    for page_num, block_text in table_blocks:
        label_match = re.search(r"Label:\s*(.+)", block_text)
        label = label_match.group(1).strip() if label_match else None
        manifest.blocks.append(EvidenceBlock(
            block_type=BlockType.TABLE,
            text=block_text,
            source_name=source_name,
            page=page_num,
            label=label,
        ))
        if label:
            manifest.table_labels.append(label)

    # --- Inline table recovery from prose text ---
    inline_tables = detect_inline_tables(text)
    for itbl in inline_tables:
        # Avoid duplicating tables already captured structurally
        if not _inline_table_duplicates_existing(itbl["text"], manifest.blocks):
            manifest.blocks.append(EvidenceBlock(
                block_type=BlockType.TABLE,
                text=itbl["text"],
                source_name=source_name,
                label=itbl.get("label"),
                confidence="medium",
            ))
            if itbl.get("label"):
                manifest.table_labels.append(itbl["label"])

    manifest.n_tables = len([b for b in manifest.blocks if b.block_type == BlockType.TABLE])

    # --- Paragraph-level classification ---
    paragraphs = text.split("\n\n")
    current_page: Optional[int] = None

    # Detect where the references section begins so we can suppress
    # model/result classification for bibliography paragraphs.
    refs_boundary_idx: Optional[int] = None
    for idx, para in enumerate(paragraphs):
        stripped = para.strip()
        # Standalone heading: "References" or "Bibliography" as the whole paragraph
        if re.match(
            r"\s*(?:references|bibliography)\s*$",
            stripped,
            re.IGNORECASE,
        ):
            refs_boundary_idx = idx
            break
        # Heading followed immediately by citations in the same paragraph
        if (
            re.match(r"\s*(?:references|bibliography)\b", stripped, re.IGNORECASE)
            and re.search(r"\[\d+\]", stripped)
        ):
            refs_boundary_idx = idx
            break
        # "References" on its own line within a multi-line paragraph
        # (e.g. "Appendix A...\nReferences\n[1] ...")
        if re.search(r"(?:^|\n)\s*(?:references|bibliography)\s*\n", para, re.IGNORECASE):
            refs_boundary_idx = idx
            break
        # Catch paragraphs where the heading ends the previous paragraph
        # and this paragraph opens with a run of numbered citation entries
        if re.search(r"\breferences\b", para.strip(), re.IGNORECASE) and \
           len(re.findall(r"\[\d+\]", para)) >= 3:
            refs_boundary_idx = idx
            break

    # Fallback: if no heading was found, detect paragraphs that are
    # predominantly citation entries (e.g. [15] Author, ... [16] Author, ...)
    # This catches the common case where the heading was at the end of
    # the previous page and the citation entries start on a new page.
    if refs_boundary_idx is None:
        for idx, para in enumerate(paragraphs):
            stripped = para.strip()
            citation_markers = re.findall(r"^\s*\[\d+\]", stripped, re.MULTILINE)
            if len(citation_markers) >= 3:
                refs_boundary_idx = idx
                break

    for para_idx, para in enumerate(paragraphs):
        para = para.strip()
        if not para:
            continue

        page_match = re.search(r"\[Page\s+(\d+)\]", para, re.IGNORECASE)
        if page_match:
            current_page = int(page_match.group(1))

        if len(para) < 30:
            continue

        # Force everything in the references/bibliography section to NARRATIVE
        # to prevent cited method names (e.g. "GAMLSS" in a citation) from
        # being classified as model blocks.
        in_references = (refs_boundary_idx is not None and para_idx >= refs_boundary_idx)

        # Detect abstract/introduction paragraphs: these often mention method
        # names (ANCOVA, MANOVA, GAMLSS) as part of a brief summary but are
        # not the methods section.  Classify them as NARRATIVE to prevent
        # spurious MODEL block assignments on page 1.
        #
        # Scan the full paragraph (not just [:300]) because title+author blocks
        # can push "Abstract" far into the text.  Also detect first-page
        # structural cues: received/accepted dates, handling editor, keywords.
        is_abstract_intro = bool(re.search(
            r"(?i)(?:"
            r"\babstract\b|"
            r"\bbackground\s+and\s+aims?\b|"
            r"\bmethods?\s+and\s+results?\b|"
            r"\bconclusions?\s*:|"
            r"\bkeywords?\s*:|"
            r"\breceived\s+\d+\s+\w+\s+\d{4}\b|"             # "Received 10 December 2021"
            r"\bhandling\s+editor\b|"                          # journal front-matter
            r"\bavailable\s+online\s+\d|"                      # "Available online 10 April 2022"
            r"\bintroduction\b.*\n.*\bbody\s+mass\s+index\b"
            r")",
            para,
        ))

        if in_references or is_abstract_intro:
            block_type = BlockType.NARRATIVE
        else:
            block_type = classify_paragraph_block(para)

        # Do not duplicate TABLE blocks already captured
        if block_type == BlockType.TABLE and "[TABLE_START]" in para:
            continue

        model_method_class = None
        if not in_references and not is_abstract_intro and block_type in {BlockType.MODEL, BlockType.RESULT}:
            model_method_class = classify_model_block(para)

        manifest.blocks.append(EvidenceBlock(
            block_type=block_type,
            text=para,
            source_name=source_name,
            page=current_page,
            model_method_class=model_method_class,
        ))

    # --- Manifest flags from full text ---
    # A manuscript typeset in Word renders its formulae as Mathematical
    # Alphanumeric Symbols and numbers them "(1)" on the same line, with no
    # literal word "equation" anywhere. One submission with numbered display
    # equations and a full appendix of model specification was reported as
    # having none.
    manifest.has_equations = bool(re.search(
        r"(?:equation\s+\d|eq\.\s*\d|~\s*N\(|Score.*=.*β|model\s+equation)", text_lower
    ) or re.search(r"[\U0001D400-\U0001D7FF]", text or "")
       or re.search(r"(?m)^.{0,120}[=∼~].{0,60}\(\d{1,2}\)\s*$", text or ""))
    manifest.has_model_spec = bool(re.search(
        r"(?:model\s+equation|regression\s+model|cross.?classif|fitted\s+model|model\s+specification"
        r"|\bancova\b|\bmancova\b|\banova\b|\bmanova\b|\bgamlss\b"
        r"|distributional\s+(?:model|regression)|centile\s+curve)", text_lower
    ))
    # R-squared, adjusted R-squared and Wilks' Lambda are standardised effect
    # sizes and are what ANCOVA/MANCOVA papers report. Recognising only Cohen's
    # d and eta-squared made the manifest say "Effect sizes reported: False"
    # for a paper whose tables give F, R2 and Adj R2 for every covariate, and
    # the review then criticised it for omitting effect sizes entirely.
    manifest.has_effect_sizes = bool(re.search(
        r"(?:cohen|hedges|glass|cliff|eta.?squared|partial\s+eta|omega.?squared"
        r"|effect\s+size|odds\s+ratio|hazard\s+ratio|risk\s+ratio|relative\s+risk"
        r"|rate\s+ratio|cram[e\u00e9]r|wilks|\br\s*\^?\s*2\b|r\u00b2"
        r"|r-?squared|adj\w*\.?\s*r\s*\^?\s*2?\b"
        r"|standardi[sz]ed\s+(?:beta|coefficient|mean\s+difference))", text_lower
    ))
    manifest.has_confidence_intervals = bool(re.search(
        r"(?:confidence\s+interval|95\s*%\s*ci|\bci\s*[=:\[])", text_lower
    ))
    manifest.has_p_values = bool(re.search(
        r"(?:p\s*[<=<]\s*0?\.\d|p\s*=\s*0?\.\d)", text_lower
    ))
    manifest.has_sample_description = bool(re.search(
        r"(?:participant|sample\s+size|n\s*=\s*\d|volunteer|recruit)", text_lower
    ))
    # --- Randomisation flag (keep this conservative) ---
    randomisation_positive = bool(re.search(
        r"(?:\brandomi[sz](?:ed|ation)?\b|\brandom\s+allocat(?:ed|ion)?\b|\brandom\s+assign(?:ed|ment)?\b|\bcounterbalanc(?:ed|ing)\b|\brandomi[sz]ed\s+order\b)",
        text_lower
    ))

    sampling_nonrandom = bool(re.search(
        r"(?:\bpurposive\s+sampling\b|\bconvenience\s+sampling\b|\bconsecutive\s+sampling\b|\bquota\s+sampling\b|\bsnowball\s+sampling\b|\bparticipants\s+were\s+invited\b|\bparticipants\s+were\s+selected\b)",
        text_lower
    ))

    manifest.has_randomisation = randomisation_positive and not sampling_nonrandom

    manifest.has_standard_errors = bool(re.search(
        r"(?:standard\s+error|SE\s*=|s\.e\.\s*=|\(\s*\d+\.\d+\s*\)\s*$)", text_lower
    ))
    manifest.has_variance_components = bool(re.search(
        r"(?:variance\s+component|between.?\w+\s+variance|random.*variance|σ²|sigma\s*squared)", text_lower
    ))
    manifest.has_model_fit_stats = bool(re.search(
        r"(?:\bDIC\b|\bAIC\b|\bBIC\b|\bWAIC\b|\bLOO\b|deviance\s+information|"
        r"R[\-\s]?squared|adjusted\s+R|log[\-\s]?likelihood|-2\s*log)", text_lower
    ))
    manifest.n_figures_mentioned = len(re.findall(
        r"(?i)(?:fig(?:ure)?\.?\s+\d)", text
    ))

    # --- Method classification ---
    manifest.method_class = classify_method(text)
    manifest.additional_method_classes = classify_additional_methods(text, manifest.method_class)

    verdict = de.classify_design(text)
    if verdict.design_class is not de.DesignClass.UNCLASSIFIED:
        manifest.design_class = verdict.design_class
        manifest.design_signals = verdict.matched
    # In an evidence synthesis the body describes the included studies'
    # methods, so the paper-level classification may not describe this paper.
    # Both are reported and a disagreement is stated rather than resolved.
    if manifest.design_class is de.DesignClass.EVIDENCE_SYNTHESIS:
        scoped = de.scoped_synthesis_text(text)
        if scoped:
            manifest.synthesis_method_class = classify_method(scoped)
    
    combined_table_text = "\n\n".join(block_text for _, block_text in table_blocks)
    combined_block_text = "\n\n".join(b.text for b in manifest.blocks if b.block_type == BlockType.TABLE)
    full_evidence_text = text + "\n\n" + combined_table_text + "\n\n" + combined_block_text
    manifest = refine_manifest_flags(manifest, full_evidence_text)

    return manifest


def detect_inline_tables(text: str) -> List[dict]:
    """
    Detect tabular data embedded in prose text, including:
    1. Markdown pipe tables from Marker output
    2. Numeric column patterns (rows of aligned numbers)
    3. Table headers followed by numeric rows

    Excludes numeric patterns found in reference/bibliography sections.
    """
    tables: List[dict] = []

    # --- Identify where the references section begins (if any) ---
    ref_match = re.search(
        r"\b(?:references|bibliography)\b",
        text,
        re.IGNORECASE,
    )
    ref_start_pos = ref_match.start() if ref_match else len(text)

    # --- Pattern 1: Markdown pipe tables (from Marker) ---
    pipe_table_re = re.compile(
        r"((?:^[ \t]*\|.+\|[ \t]*\n){3,})",
        re.MULTILINE,
    )
    for m in pipe_table_re.finditer(text):
        # Skip tables found inside the references section
        if m.start() >= ref_start_pos:
            continue
        block_text = m.group(1).strip()
        # Skip separator-only tables
        rows = [r for r in block_text.split("\n") if r.strip() and not re.match(r"^\s*\|[\s\-:]+\|\s*$", r)]
        if len(rows) >= 2:
            label = _infer_inline_table_label(text, m.start())
            tables.append({
                "text": f"[TABLE_START]\nSource: inline_markdown\n{block_text}\n[TABLE_END]",
                "label": label,
            })

    # --- Pattern 2: Rows of aligned numbers preceded by a header ---
    lines = text.split("\n")
    # Pre-compute a character offset for each line so we can check
    # whether a run falls inside the references section
    line_offsets: List[int] = []
    offset = 0
    for line in lines:
        line_offsets.append(offset)
        offset += len(line) + 1  # +1 for the newline

    i = 0
    while i < len(lines):
        run_start = None
        run_lines = []
        j = i
        while j < len(lines):
            line = lines[j].strip()
            # Count numeric tokens (integers, decimals, negatives)
            nums = re.findall(r"-?\d+\.?\d*", line)
            if len(nums) >= 2 and len(line.split()) <= 12:
                if run_start is None:
                    run_start = j
                run_lines.append(line)
            else:
                if run_lines and len(run_lines) >= 3:
                    break
                run_start = None
                run_lines = []
            j += 1

        if run_lines and len(run_lines) >= 3:
            # --- Filter: skip runs inside the references section ---
            run_char_offset = line_offsets[run_start] if run_start < len(line_offsets) else len(text)
            if run_char_offset >= ref_start_pos:
                i = j + 1 if j > i else i + 1
                continue

            # --- Filter: skip runs that look like reference entries ---
            ref_line_count = sum(
                1 for rl in run_lines
                if re.search(
                    r"^\s*\[\d+\]|"          # [1], [2], ...
                    r"\d{4}\s*[;:]\s*\d|"    # 2018;18(3):248
                    r"\bet\s+al\b|"          # et al.
                    r"\bvol\.\s*\d|"         # vol. 2
                    r"\bpp?\.\s*\d",          # p. 63 or pp. 63
                    rl, re.IGNORECASE,
                )
            )
            if ref_line_count >= len(run_lines) * 0.5:
                i = j + 1 if j > i else i + 1
                continue

            # --- Filter: skip runs that are mostly prose with incidental numbers ---
            # Tabular data has a high number-to-word ratio; prose does not.
            prose_line_count = sum(
                1 for rl in run_lines
                if (
                    len(rl.split()) > 0
                    and len(re.findall(r"-?\d+\.?\d*", rl)) / len(rl.split()) < 0.3
                )
            )
            if prose_line_count >= len(run_lines) * 0.5:
                i = j + 1 if j > i else i + 1
                continue

            # --- Filter: check column-count consistency ---
            # Genuine tables have roughly the same number of numeric tokens per row.
            num_counts = [len(re.findall(r"-?\d+\.?\d*", rl)) for rl in run_lines]
            median_count = sorted(num_counts)[len(num_counts) // 2]
            consistent_rows = sum(1 for c in num_counts if abs(c - median_count) <= 1)
            if consistent_rows < len(run_lines) * 0.6:
                i = j + 1 if j > i else i + 1
                continue

            # Include a header line if the line before the run has words
            header = ""
            if run_start and run_start > 0:
                candidate = lines[run_start - 1].strip()
                if candidate and re.search(r"[A-Za-z]", candidate):
                    header = candidate + "\n"

            block_text = header + "\n".join(run_lines)
            label = _infer_inline_table_label(text, text.find(run_lines[0]))
            tables.append({
                "text": f"[TABLE_START]\nSource: inline_numeric\n{block_text}\n[TABLE_END]",
                "label": label,
            })
            i = j
        else:
            i = j + 1 if j > i else i + 1

    return tables


def _infer_inline_table_label(text: str, position: int) -> Optional[str]:
    """Look backwards from a table position for a Table N label."""
    preceding = text[max(0, position - 300):position]
    m = re.search(r"(Table\s+\d+[a-z]?)\b", preceding, re.IGNORECASE)
    return m.group(1) if m else None


def _inline_table_duplicates_existing(inline_text: str, existing_blocks: List[EvidenceBlock]) -> bool:
    """Check whether an inline table duplicates data already in a structured block."""
    # Extract the numeric values from the inline table
    inline_nums = set(re.findall(r"-?\d+\.?\d*", inline_text))
    if len(inline_nums) < 4:
        return False
    for block in existing_blocks:
        if block.block_type == BlockType.TABLE:
            block_nums = set(re.findall(r"-?\d+\.?\d*", block.text))
            # If >70% of inline numbers appear in an existing block, it is a duplicate
            if inline_nums and len(inline_nums & block_nums) / len(inline_nums) > 0.7:
                return True
    return False


# ---------------------------------------------------------------------------
# Phase 6: Programmatic post-validation
# ---------------------------------------------------------------------------

def programmatic_post_checks(report_text: str, manifest: EvidenceManifest) -> List[str]:
    """
    Catch common LLM reviewer errors by comparing claims in the report
    against what the evidence manifest confirms is present.

    Returns a list of correction instructions to feed into LLM validation.
    """
    corrections: List[str] = []
    report_lower = report_text.lower()

    # 1. "Table missing" when tables are present
    if manifest.n_tables > 0:
        if re.search(r"(?:table[s]?\s+(?:are|is|were|was)\s+(?:missing|absent|not\s+provided|not\s+reported))", report_lower):
            corrections.append(
                f"CORRECTION: The report claims tables are missing, but {manifest.n_tables} table(s) were extracted "
                f"({', '.join(manifest.table_labels) if manifest.table_labels else 'unlabelled'}). "
                f"Remove or revise this claim."
            )

    # 2. "Equation/model spec missing" when present
    if manifest.has_model_spec or manifest.has_equations:
        if re.search(r"(?:equation|model\s+(?:specification|form|equation))\s+(?:is|are|was|were)\s+(?:missing|absent|not\s+provided|not\s+reported|unclear)", report_lower):
            corrections.append(
                "CORRECTION: The report claims the model specification or equation is missing/unclear, "
                "but the extracted evidence contains model specification content. Revise to acknowledge this."
            )

    # 3. Requesting VIF for single-predictor models
    if re.search(r"\bvif\b", report_lower):
        model_blocks = [b for b in manifest.blocks if b.block_type == BlockType.MODEL]
        all_model_text = " ".join(b.text for b in model_blocks).lower()
        # Also check the full text for single-predictor indicators
        full_check = all_model_text + " " + " ".join(
            b.text.lower() for b in manifest.blocks
            if b.block_type in (BlockType.RESULT, BlockType.DESIGN)
        )
        if re.search(r"single\s+fixed\s+(?:factor|effect|predictor)|single\s+predictor|one\s+predictor|single\s+.*\bfactor\b", full_check):
            corrections.append(
                "CORRECTION: The report requests VIF checks, but the model has a single predictor. "
                "VIF is not applicable. Remove this request."
            )

    # 4. Requesting Bayesian diagnostics for non-Bayesian MCMC models
    mc = manifest.method_class
    if mc in (MethodClass.CROSS_CLASSIFIED_MCMC, MethodClass.MULTILEVEL_MCMC):
        bayesian_requests = re.findall(
            r"(?:prior\s+(?:specification|sensitivity|justification|choice)"
            r"|posterior\s+(?:distribution|summary|check)"
            r"|credible\s+interval"
            r"|bayes\s+factor"
            r"|framework\s+(?:is|remains|appears)\s+(?:unclear|ambiguous|hybrid))",
            report_lower
        )
        if bayesian_requests:
            corrections.append(
                f"CORRECTION: The model uses MCMC as an estimation method (classified as {mc.value}), "
                f"not as a Bayesian analysis. The report contains {len(bayesian_requests)} reference(s) to "
                f"Bayesian diagnostics or calls the framework unclear. Remove all such references."
            )

    # 5. Requesting OLS-style residual plots for MCMC/Bayesian/GAMLSS models
    if mc in (MethodClass.CROSS_CLASSIFIED_MCMC, MethodClass.MULTILEVEL_MCMC,
              MethodClass.BAYESIAN_MODEL, MethodClass.BAYESIAN_MIXED_EFFECTS,
              MethodClass.BAYESIAN_REGRESSION, MethodClass.BAYESIAN_QUANTILE_REGRESSION,
              MethodClass.DISTRIBUTIONAL_MODEL):
        if re.search(r"residual\s+plot", report_lower):
            corrections.append(
                f"CORRECTION: The report requests residual plots, but the analysis uses {mc.value}. "
                f"Standard OLS residual plots are not the primary diagnostic for this framework. "
                f"Revise to request framework-appropriate diagnostics (e.g., convergence diagnostics for MCMC, "
                f"randomised quantile residuals for GAMLSS) or remove if not applicable."
            )

    # 5b. Requesting hierarchical/random-effects diagnostics for non-hierarchical Bayesian models
    if mc in (MethodClass.BAYESIAN_REGRESSION, MethodClass.BAYESIAN_QUANTILE_REGRESSION,
              MethodClass.BAYESIAN_MODEL):
        hierarchical_requests = re.findall(
            r"(?:random\s+(?:intercept|slope|effect)\s+(?:structure|specification|justification)"
            r"|variance\s+component\s+(?:report|estimat)"
            r"|hierarchical\s+(?:structure|model)\s+(?:is|was|should|need)"
            r"|multilevel\s+(?:structure|model)\s+(?:is|was|should|need)"
            r"|group[\-\s]?level\s+(?:effect|variance)\s+(?:is|was|should|need)"
            r"|nesting\s+(?:structure|level))",
            report_lower,
        )
        if hierarchical_requests:
            corrections.append(
                f"CORRECTION: The model is classified as {mc.value} (non-hierarchical). "
                f"The report contains {len(hierarchical_requests)} reference(s) to hierarchical/random-effects "
                f"structure or diagnostics. Remove these — they are not applicable to this model type."
            )

    # 5c. Survey-design concerns in non-survey papers
    # Check if the paper uses survey-design language; if not, flag any
    # survey-specific concerns in the report for removal.
    all_block_text = " ".join(b.text.lower() for b in manifest.blocks)
    is_survey_design = bool(re.search(
        r"\bstratified\s+random\s+probability\s+sample"
        r"|\bhealth\s+survey\b"
        r"|\bnational\s+health\b"
        r"|\bnhanes\b"
        r"|\bpopulation[\-\s]?based\s+survey"
        r"|\bpopulation[\-\s]?level\s+survey"
        r"|\bcomplex\s+survey\s+design"
        r"|\bsurvey\s+weight",
        all_block_text,
    ))
    if not is_survey_design:
        survey_concerns = re.findall(
            r"survey\s+weight|clustering\s+adjust|medication\s+sensitivity|"
            r"hierarchical\s+structure\s+of\s+(?:swimmers|participants|patients|subjects)\s+within\s+(?:teams|clubs|regions|centres)",
            report_lower,
        )
        if survey_concerns:
            corrections.append(
                "CORRECTION: The report raises survey-design concerns (survey weights, clustering "
                "adjustments) but the paper does not describe a population-level survey design. "
                "Remove all survey-specific scrutiny points."
            )

    # 6. "Sample description missing" when present
    if manifest.has_sample_description:
        if re.search(r"sample\s+(?:size|description)\s+(?:is|are|was|were)\s+(?:missing|not\s+reported|absent)", report_lower):
            corrections.append(
                "CORRECTION: The report claims sample description is missing, but participant/sample "
                "information was found in the extracted evidence. Revise this claim."
            )

    # 7. "Randomisation unclear" when described
    if manifest.has_randomisation:
        if re.search(r"randomi[sz]ation\s+(?:is|was|remains)\s+(?:unclear|ambiguous|not\s+described|not\s+specified)", report_lower):
            corrections.append(
                "CORRECTION: The report claims randomisation is unclear or not described, but "
                "randomisation details were found in the extracted evidence. Revise: the criticism "
                "should target specific aspects that remain unclear, not claim absence."
            )

    # 8. "P-values/CIs not reported" when present
    if manifest.has_p_values:
        if re.search(
            r"p[\-\s]?value[s]?"
            r"\s+(?:are|is|were|was)"
            r"\s+(?:not\s+(?:\w+\s+)?(?:reported|provided|included|accessible|available)"
            r"|missing|absent)",
            report_lower,
        ):
            corrections.append(
                "CORRECTION: The report claims p-values are not reported, but p-values were found "
                "in the extracted evidence. Revise to specify which particular p-values are missing "
                "rather than claiming general absence."
            )

    if manifest.has_confidence_intervals:
        if re.search(
            r"confidence\s+interval[s]?"
            r"\s+(?:are|is|were|was)"
            r"\s+(?:not\s+(?:\w+\s+)?(?:reported|provided|included|accessible|available)"
            r"|missing|absent)",
            report_lower,
        ):
            corrections.append(
                "CORRECTION: The report claims confidence intervals are not reported, but CIs were "
                "found in the extracted evidence. Revise to specify which particular CIs are missing "
                "rather than claiming general absence."
            )
        # Also catch comma-separated lists: "standard errors, confidence intervals, ... are not reported"
        elif re.search(
            r"(?:standard\s+error|SE|p[\-\s]?value|uncertainty)[s,\s]+(?:and\s+)?confidence\s+interval[s]?"
            r".{0,40}?"
            r"(?:not\s+(?:\w+\s+)?(?:reported|provided|included|accessible|available)"
            r"|missing|absent)",
            report_lower,
        ):
            corrections.append(
                "CORRECTION: The report claims confidence intervals (alongside other statistics) are "
                "not reported, but the evidence manifest confirms CIs are present. Revise this claim."
            )

    # 9. "Standard errors not reported" when present
    if manifest.has_standard_errors:
        if re.search(
            r"standard\s+error[s]?"
            r"\s+(?:are|is|were|was)"
            r"\s+(?:not\s+(?:\w+\s+)?(?:reported|provided|included|accessible|available)"
            r"|missing|absent)",
            report_lower,
        ):
            corrections.append(
                "CORRECTION: The report claims standard errors are not reported, but standard errors "
                "were found in the extracted evidence. Revise to specify which particular SEs are "
                "missing rather than claiming general absence."
            )

    # 10. "Variance components not reported/tested" when present
    if manifest.has_variance_components:
        if re.search(r"variance\s+component[s]?\s+(?:are|is|were|was)\s+(?:not\s+reported|missing|absent)", report_lower):
            corrections.append(
                "CORRECTION: The report claims variance components are not reported, but variance "
                "component estimates were found in the extracted evidence. Revise this claim."
            )

    # 11. "Model fit not reported" when DIC/AIC/BIC is present
    if manifest.has_model_fit_stats:
        if re.search(r"(?:model\s+fit|goodness[\-\s]of[\-\s]fit|fit\s+statistic)[s]?\s+(?:are|is|were|was)\s+(?:not\s+reported|missing|absent)", report_lower):
            corrections.append(
                "CORRECTION: The report claims model fit statistics are not reported, but fit "
                "statistics (e.g., DIC, AIC, deviance) were found in the extracted evidence."
            )

    # 12. Cross-contamination: numbers from cited studies confused with current study
    # Detect when the report claims a sample discrepancy using a number that
    # belongs to a different study cited in the text
    contamination_issues = _detect_cross_contamination(report_lower, manifest)
    corrections.extend(contamination_issues)

    return corrections
    
def refine_manifest_flags(manifest: EvidenceManifest, full_text: str) -> EvidenceManifest:

    """
    Upgrade manifest flags when evidence is clearly present in text or table blocks
    but was missed by the initial coarse regex checks.
    """
    full_low = full_text.lower()

    def seen(*patterns: str) -> bool:
        return any(p.lower() in full_low for p in patterns)

    # Model specification / equations
    if seen("the following equation", "equation was applied", "model equation",
            "ln(vo2max", "p = w", "work–time model", "work-time model",
            "ancova", "mancova", "manova", "gamlss", "centile curve",
            "willk", "wilks"):
        manifest.has_model_spec = True
        manifest.has_equations = True

    # Confidence intervals
    if seen("95% ci", "confidence interval", "lower bound", "upper bound", " lower ", " upper "):
        manifest.has_confidence_intervals = True

    # Standard errors / SD-like reporting
    if seen(" se ", "\tse\t", "standard error", "trimmed sd", "sd (0.2)", "winsorized variances"):
        manifest.has_standard_errors = True

    # Fit statistics / performance metrics / effect sizes
    if seen("r squared", "adjusted r squared", "aic", "bic", "waic", "loo", "looic", "dic", "rmse", "mae", "effect size", "s value", "surprisal"):
        manifest.has_model_fit_stats = True

    if seen("effect size", "cohen", "robust version of cohen", "eta squared", "partial eta", "odds ratio", "hazard ratio", "s value"):
        manifest.has_effect_sizes = True

    # P-values
    if seen("p =", "p <", "p ≤", "sig.", "benjamini", "hochberg"):
        manifest.has_p_values = True

    # Randomisation
    if seen("randomized", "randomised", "randomized into", "randomised into", "randomization", "randomisation", "order of test completion"):
        manifest.has_randomisation = True

    return manifest

def _detect_cross_contamination(report_lower: str, manifest: EvidenceManifest) -> List[str]:
    """
    Detect when the LLM has confused numbers from cited studies with the
    current study's data, creating spurious discrepancy claims.
    """
    corrections: List[str] = []

    # Look for "discrepancy" or "unexplained" claims involving specific numbers
    discrepancy_patterns = re.findall(
        r"(?:unexplained|discrepan|inconsisten|contradict)\w*"
        r".{0,80}?"  # allow up to 80 chars between keyword and first number
        r"(\d{2,})"
        r".{0,40}?"  # allow up to 40 chars between first number and comparison word
        r"(?:vs\.?|versus|and|compared\s+to|but|yet|however)"
        r".{0,40}?"  # allow up to 40 chars between comparison word and second number
        r"(\d{2,})",
        report_lower,
    )

    if not discrepancy_patterns:
        return corrections

    # Gather numbers that appear in citation contexts (near "et al.", year references,
    # or phrases indicating a different study)
    all_text = " ".join(b.text for b in manifest.blocks)
    cited_numbers: set = set()

    # Match numbers near citation markers (et al., year references)
    for m in re.finditer(r"(?:et\s+al\.?\s*[\(,]?\s*\d{4}\)?|[\(]\d{4}[\)])", all_text):
        context_start = max(0, m.start() - 200)
        context_end = min(len(all_text), m.end() + 200)
        context = all_text[context_start:context_end]
        for num in re.findall(r"\b(\d{2,})\b", context):
            cited_numbers.add(num)

    # Also flag numbers near "previous study", "larger dataset", "earlier study" etc.
    for m in re.finditer(r"(?:previous\s+\w*\s*(?:study|dataset|research|investigation)"
                         r"|larger\s+dataset|earlier\s+\w*\s*study|separate\s+\w*\s*study"
                         r"|prior\s+\w*\s*(?:study|research))", all_text, re.IGNORECASE):
        context_start = max(0, m.start() - 200)
        context_end = min(len(all_text), m.end() + 200)
        context = all_text[context_start:context_end]
        for num in re.findall(r"\b(\d{2,})\b", context):
            cited_numbers.add(num)

    for num1, num2 in discrepancy_patterns:
        if num1 in cited_numbers or num2 in cited_numbers:
            corrections.append(
                f"CORRECTION: The report identifies a discrepancy between {num1} and {num2}, but "
                f"at least one of these numbers appears in the context of a cited study (e.g., "
                f"a different dataset referenced by the authors), not the current study's data. "
                f"Verify that this is a genuine within-study discrepancy before claiming inconsistency. "
                f"If the numbers come from different studies, remove the discrepancy claim."
            )
            break  # One correction is sufficient to flag the issue

    return corrections


# ---------------------------------------------------------------------------
# OCR and text cleaning (preserved from original)
# ---------------------------------------------------------------------------

SUPERSCRIPT_MAP = str.maketrans({
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
    "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾",
    "n": "ⁿ", "i": "ⁱ",
})

SUBSCRIPT_MAP = str.maketrans({
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
    "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎",
    "a": "ₐ", "e": "ₑ", "i": "ᵢ", "j": "ⱼ", "o": "ₒ",
    "r": "ᵣ", "u": "ᵤ", "v": "ᵥ", "x": "ₓ",
})


def _to_superscript(s: str) -> str:
    return s.translate(SUPERSCRIPT_MAP)


def _to_subscript(s: str) -> str:
    return s.translate(SUBSCRIPT_MAP)


def clean_ocr_artifacts(text: str) -> str:
    """Centralised function to fix PDF OCR corruption across both text and tables."""
    text = text.replace("\x00", " ").replace("", "")
    text = text.replace("(cid:3)", "-").replace("(cid:2)", "±")

    # Fix the specific index name
    text = text.replace("WHT$5R", "WHT.5R")
    text = text.replace("WHT$5", "WHT.5")

    # Fix the $ replacing decimal points in numbers (e.g., 0$249 -> 0.249, 177$5 -> 177.5)
    text = re.sub(r"(\d)\$(\d)", r"\1.\2", text)

    # Fix OCR 'Z' being used instead of '='
    text = re.sub(r"\bWHT\.5R Z\b", "WHT.5R =", text)
    text = re.sub(r"\br Z\b", "r =", text)
    text = re.sub(r"\bF Z\b", "F =", text)
    text = re.sub(r"\bp Z\b", "p =", text)

    return text


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("", "")
    # Note: Using " +" instead of "[ \t]+" to preserve tabs for table columns
    text = re.sub(r" +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_extracted_evidence_text(text: str, domain: str = "general") -> str:
    text = clean_text(text)
    text = clean_ocr_artifacts(text)

    replacements = {
        "VO 2max": "VO2max",
        "VO 2peak": "VO2peak",
        "VO ₂max": "VO2max",
        "VO ₂peak": "VO2peak",
        "ln (": "ln(",
        "Ln (": "ln(",
        "Age 2": "Age2",
        "age 2": "age2",
        "L. min−1": "L·min−1",
        "mL. kg−1.min−1": "mL·kg−1·min−1",
        "kg −1": "kg−1",
        "min −1": "min−1",
        "body fat %": "body fat%",
        "BF %": "BF%",
        "Cyc le": "Cycle",
        "algor ithms": "algorithms",
        "method- ology": "methodology",
        "allome- tric": "allometric",
        "physiolog- ical": "physiological",
        "incor - porating": "incorporating",
        "log- transformed": "log-transformed",
        "out- of- sample": "out-of-sample",
        "all- cause": "all-cause",
        "fat- free": "fat-free",
        "whole- body": "whole-body",
        "body- composition": "body-composition",
        "X- ray": "X-ray",
        "meta- analysis": "meta-analysis",
        "cross- validation": "cross-validation",
        "peer- reviewed": "peer-reviewed",
        "non- linear": "non-linear",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"([A-Za-z])-\s+([A-Za-z])", r"\1\2", text)
    text = re.sub(r"ln\(mass\(\s*g\s*\)\)", "ln(mass)", text, flags=re.IGNORECASE)
    text = re.sub(r"ln\(mass\s*\(\s*g\s*\)\)", "ln(mass)", text, flags=re.IGNORECASE)
    text = re.sub(r"\bVO\s*\(?\s*2max\s*\)?", "VO2max", text)

    if domain == "exercise_physiology":
        text = re.sub(r"\bln\(masskg\)\b", "ln(mass kg)", text, flags=re.IGNORECASE)
        text = re.sub(r"\bln\(fffkg\)\b", "ln(FFM kg)", text, flags=re.IGNORECASE)
        text = re.sub(r"\blnfvc\b", "ln(FVC)", text, flags=re.IGNORECASE)
        text = re.sub(r"\blnwc\b", "ln(WC)", text, flags=re.IGNORECASE)

    text = re.sub(r" +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Table extraction (preserved from original)
# ---------------------------------------------------------------------------

@dataclass
class DocChunk:
    source_name: str
    chunk_id: int
    text: str


def truncate_cell(value, max_chars: int = MAX_TABLE_CHARS_PER_CELL) -> str:
    if value is None:
        return ""
    s = str(value).strip().replace("\n", " ")
    s = re.sub(r" {2,}", " ", s)
    if len(s) > max_chars:
        s = s[: max_chars - 3] + "..."
    return s


def table_to_block(table: List[List[str]], table_num: int, page_num: int) -> str:
    if not table:
        return ""

    trimmed_rows = table[:MAX_TABLE_ROWS]
    lines = [
        "[TABLE_START]",
        f"Page: {page_num}",
        f"Table number on page: {table_num}",
    ]
    for row in trimmed_rows:
        row = row[:MAX_TABLE_COLS]
        cleaned = [truncate_cell(cell) for cell in row]
        lines.append("\t".join(cleaned))
    lines.append("[TABLE_END]")
    return "\n".join(lines).strip()


def normalise_table_rows(rows: list[list[str]]) -> list[list[str]]:
    cleaned = []
    for row in rows:
        cleaned_row = []
        for cell in row:
            cell = "" if cell is None else str(cell)
            cell = clean_ocr_artifacts(cell)
            cell = re.sub(r" +", " ", cell).strip()
            cleaned_row.append(cell)
        if any(cell for cell in cleaned_row):
            cleaned.append(cleaned_row)
    return cleaned


# pdfplumber's default strategy looks for ruled lines in both directions. Many
# journal tables are set booktabs-style, with horizontal rules only and no
# vertical ones, and those extract as nothing at all. A second pass using the
# text-alignment strategy recovers them. It is noisier, so its output is
# cleaned, split at table captions, and scored below the ruled extraction, and
# the existing prose and reference-list filters reject what is not a table.
_TEXT_TABLE_SETTINGS = {
    "vertical_strategy": "text",
    "horizontal_strategy": "text",
}

# Tolerant of a caption whose leading character was lost to column splitting,
# which happens when the caption sits flush against a column boundary.
_TABLE_CAPTION_RE = re.compile(r"(?:^|\s)T?able\s+(\d+|[A-Z])\b[.:]?", re.I)


# A caption starts a line; "see Table 1" in a sentence does not. Matching
# anywhere in the page text made every narrative page that referred to a table
# look as though it contained one.
# Journals punctuate captions in several ways, and a caption must be
# recognised in all of them or the table arrives unlabelled:
#   "Table 1. Caption"   "Table 1: Caption"   "Table 1 | Caption"  (Frontiers)
#   "Table 1 Caption"    "Table1 Caption"     "TABLE I. Caption"
# Requiring a full stop recognised only the first two, and missed captions in
# five of fifteen papers in a real corpus.
_CAPTION_LINE_RE = re.compile(
    r"(?m)^[ \t]*(?:(Supplementary)\s*)?T\s?able\s*"
    r"([0-9]+[a-z]?|[IVXL]+|[A-Z])\s*(?:[.:|\u2013\u2014-]|\s)\s*\S",
    re.I,
)

# A text-strategy extraction is only believed when it is dense with numbers.
# Measured on a journal article: real tables scored 0.52-0.94, narrative pages
# 0.11-0.40. Raise MIN_TABLE_NUMERIC_RATIO if prose still leaks through.
MIN_TABLE_NUMERIC_RATIO = 0.45
MIN_TABLE_NUMERIC_CELLS = 12


def _page_table_captions(page_text: str) -> list[str]:
    """Table captions in reading order, from captions that begin a line."""
    if not page_text:
        return []
    captions = []
    for match in _CAPTION_LINE_RE.finditer(page_text):
        supplementary, number = match.group(1), match.group(2)
        if supplementary:
            label = f"Supplementary Table {number}".strip() if number else "Supplementary Table"
        elif number:
            label = f"Table {number}"
        else:
            continue
        if label not in captions:
            captions.append(label)
    return captions


def _table_numeric_ratio(rows: list[list[str]]) -> tuple[float, int]:
    """Proportion of non-empty cells that look numeric, and the raw count."""
    nonempty = sum(1 for row in rows for cell in row if str(cell).strip())
    numeric = _count_numeric_cells(rows)
    return (numeric / nonempty if nonempty else 0.0), numeric


def _looks_like_real_table(rows: list[list[str]]) -> bool:
    ratio, numeric = _table_numeric_ratio(rows)
    return ratio >= MIN_TABLE_NUMERIC_RATIO and numeric >= MIN_TABLE_NUMERIC_CELLS


_NUMERIC_TOKEN_RE = re.compile(r"^[-+(\u2212]?\d[\d,.]*\)?%?$")
_INTERVAL_DASHES = {"\u2013", "\u2014", "-", "\u2212", "to"}


def _split_table_line(line: str) -> list[str]:
    """
    Turn one line of a table into cells.

    pdfplumber's text strategy derives column edges from character gaps, which
    splits numbers down the middle ("-4.42" becomes "-4." and "42"). The page's
    own text lines keep every number intact, so rows are built from those
    instead: leading words form the row label, and each number becomes a cell,
    with "0.32 - 0.90" kept together as one interval.
    """
    tokens = line.split()
    if not tokens:
        return []
    # A caption is a sentence, not a row of cells: keep it in one piece so the
    # table keeps its name.
    if _TABLE_CAPTION_RE.match(line.strip()):
        return [line.strip()]
    first_number = next(
        (i for i, t in enumerate(tokens) if _NUMERIC_TOKEN_RE.match(t)), None
    )
    if first_number is None:
        return [line.strip()]
    cells: list[str] = []
    label = " ".join(tokens[:first_number]).strip()
    if label:
        cells.append(label)
    rest = tokens[first_number:]
    i = 0
    while i < len(rest):
        if (
            i + 2 < len(rest)
            and rest[i + 1] in _INTERVAL_DASHES
            and _NUMERIC_TOKEN_RE.match(rest[i])
            and _NUMERIC_TOKEN_RE.match(rest[i + 2])
        ):
            cells.append(f"{rest[i]} \u2013 {rest[i + 2]}")
            i += 3
        else:
            cells.append(rest[i])
            i += 1
    return cells


def _rows_from_page_text(page_text: str) -> list[list[str]]:
    """Build candidate table rows from a page's text lines."""
    rows = []
    for line in (page_text or "").split("\n"):
        if not line.strip():
            continue
        cells = _split_table_line(line)
        if cells:
            rows.append(cells)
    return rows


def _clean_text_strategy_rows(rows: list[list[str]]) -> list[list[str]]:
    """Drop prose rows and wholly empty columns from a text-strategy table."""
    kept: list[list[str]] = []
    for row in rows:
        cells = [str(c).strip() if c is not None else "" for c in row]
        nonempty = [c for c in cells if c]
        if len(nonempty) < 2:
            if len(nonempty) == 1 and _TABLE_CAPTION_RE.match(nonempty[0]):
                # The caption is a single cell by design; it names the table.
                kept.append(cells)
                continue
            # A lone short cell with no numbers is a section heading such as
            # "Breaststroke Predictors". Dropping those left the blocks of a
            # stacked table unidentifiable.
            if (
                len(nonempty) == 1
                and len(nonempty[0].split()) <= 8
                and not re.search(r"\d", nonempty[0])
            ):
                kept.append(cells)
            continue
        # Keep caption rows even though they read as prose: they carry the
        # table's identity, and dropping them left every table anonymous.
        joined = " ".join(c for c in cells if c).strip()
        if _TABLE_CAPTION_RE.match(joined):
            kept.append(cells)
            continue
        if row_is_prose_like(cells):
            continue
        kept.append(cells)
    if not kept:
        return []
    width = max(len(r) for r in kept)
    kept = [r + [""] * (width - len(r)) for r in kept]
    keep_cols = [i for i in range(width) if any(r[i] for r in kept)]
    return [[r[i] for i in keep_cols] for r in kept]


def _split_rows_on_captions(rows: list[list[str]]) -> list[list[list[str]]]:
    """Split one page-wide extraction into separate tables at caption rows."""
    groups: list[list[list[str]]] = []
    current: list[list[str]] = []
    for row in rows:
        joined = " ".join(c for c in row if c).strip()
        if _TABLE_CAPTION_RE.match(joined) and current:
            groups.append(current)
            current = [row]
        else:
            current.append(row)
    if current:
        groups.append(current)
    return [g for g in groups if len(g) >= 2]


def extract_tables_pdfplumber(path: Path) -> list[dict]:
    tables = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                except Exception:
                    page_text = ""
                captions = _page_table_captions(page_text)

                # --- Pass 1: ruled lines (unchanged behaviour) ---
                try:
                    raw_tables = page.extract_tables()
                except Exception:
                    raw_tables = []
                table_counter = 0
                for tbl in raw_tables or []:
                    if not tbl:
                        continue
                    rows = normalise_table_rows(tbl)
                    if len(rows) < 2:
                        continue
                    table_counter += 1
                    tables.append({
                        "source": "pdfplumber",
                        "page": page_index,
                        "table_index": table_counter,
                        "label": infer_table_label({"rows": rows}),
                        "page_captions": captions,
                        "rows": rows,
                    })

                # --- Pass 2: the page's own text lines, for unruled tables ---
                text_tables = [_rows_from_page_text(page_text)]
                for tbl in text_tables or []:
                    if not tbl:
                        continue
                    cleaned = _clean_text_strategy_rows(tbl)
                    if len(cleaned) < 2 or not _looks_like_real_table(cleaned):
                        continue
                    groups = _split_rows_on_captions(cleaned)
                    caption_index = 0
                    for group in groups:
                        first_row = " ".join(c for c in group[0] if c).strip()
                        starts_with_caption = bool(_TABLE_CAPTION_RE.match(first_row))
                        # A caption-led group is kept even when it is thin on
                        # numbers. A table whose caption and column headers sit
                        # at the foot of one page, with the body overleaf, was
                        # otherwise dropped entirely, and the quantile columns
                        # then reached the reviewer unlabelled.
                        if not starts_with_caption and not _looks_like_real_table(group):
                            continue
                        label = infer_table_label({"rows": group})
                        if starts_with_caption:
                            if not label and caption_index < len(captions):
                                label = captions[caption_index]
                            caption_index += 1
                        elif not label:
                            # A table running on from the previous page. Give it
                            # a label of its own so deduplication, which keeps
                            # one unlabelled table per page, cannot discard it.
                            label = f"Table (continued, page {page_index})"
                        table_counter += 1
                        tables.append({
                            "source": "pdfplumber_text",
                            "page": page_index,
                            "table_index": table_counter,
                            "label": label,
                            "page_captions": captions,
                            "rows": group,
                        })
    except Exception:
        pass
    return tables


def extract_tables_camelot(path: Path) -> list[dict]:
    tables = []
    try:
        import camelot
    except Exception:
        return tables

    for flavor in ("lattice", "stream"):
        try:
            found = camelot.read_pdf(str(path), pages="all", flavor=flavor)
        except Exception:
            continue

        for i, tbl in enumerate(found, start=1):
            try:
                df = tbl.df
            except Exception:
                continue

            if df is None or df.empty:
                continue

            rows = normalise_table_rows(df.fillna("").values.tolist())
            if len(rows) < 2:
                continue

            page_value = getattr(tbl, "page", None)
            tables.append({
                "source": f"camelot_{flavor}",
                "page": page_value,
                "table_index": i,
                "label": infer_table_label({"rows": rows}),
                "rows": rows,
            })
    return tables


def infer_table_label(table: dict) -> str | None:
    rows = table.get("rows", [])
    head_text = " ".join(" ".join(str(c) for c in row if str(c).strip()) for row in rows[:6])
    m = re.search(r"\b(Table\s+[A-Za-z0-9]+[a-z]?)\b", head_text, flags=re.I)
    if m:
        label = m.group(1)
        label = re.sub(r" +", " ", label).strip()
        return label
    # A caption can lose its leading character to a column split ("able 1.").
    m = re.search(r"(?:^|\s)able\s+([0-9]+|[A-Z])\b", head_text)
    if m:
        return f"Table {m.group(1)}"
    return None


def row_is_prose_like(row: list[str]) -> bool:
    text = " ".join(str(c) for c in row if str(c).strip()).strip()
    if not text:
        return False
    word_count = len(text.split())
    digit_count = len(re.findall(r"\d", text))
    return word_count >= 10 and digit_count < 2


def _is_reference_list(rows: list[list[str]]) -> bool:
    """Detect when a 'table' extraction is actually a reference list."""
    all_text = " ".join(
        " ".join(str(c) for c in row) for row in rows
    ).lower()
    # Reference lists contain many author-year patterns and journal names
    ref_markers = len(re.findall(
        r"(?:\(\d{4}\)|\b\d{4}\b[,.\s]|et\s+al|j\.\s+\w+\s+\w+|psychol\.|sci\.|exerc\.)",
        all_text,
    ))
    # If more than 8 reference-like markers, this is a reference list
    if ref_markers > 8:
        return True
    # Also check for "REFERENCES" heading
    if re.search(r"\breferences\b", all_text[:200]):
        return True
    return False


def _count_numeric_cells(rows: list[list[str]]) -> int:
    """Count cells that contain meaningful numeric content."""
    count = 0
    for row in rows:
        for cell in row:
            s = str(cell).strip()
            # A numeric cell contains at least one digit and is not pure prose
            if s and re.search(r"\d", s) and len(s.split()) <= 5:
                count += 1
    return count


def score_table_candidate(table: dict) -> int:
    rows = table.get("rows", [])
    nonempty_cells = sum(1 for row in rows for cell in row if str(cell).strip())
    nrows = len(rows)
    ncols = max((len(r) for r in rows), default=0)

    prose_like_rows = sum(1 for row in rows if row_is_prose_like(row))
    prose_ratio = prose_like_rows / max(nrows, 1)
    numeric_cells = _count_numeric_cells(rows)

    # --- Reject garbage extractions ---
    # Too few meaningful cells (e.g., pdfplumber returning just "12" and "5")
    if nonempty_cells < 6:
        return -9999
    # Too few numeric cells for a data table
    if numeric_cells < 3 and nrows > 1:
        return -9999
    # Reference list false positive
    if _is_reference_list(rows):
        return -9999

    if table.get("source") == "camelot_stream":
        if prose_ratio > 0.15:
            return -9999
        if ncols < 3 and nonempty_cells < 15:
            return -9999

    # General prose-ratio filter for all sources
    if prose_ratio > 0.5:
        return -9999

    score = nonempty_cells + 5 * nrows + 3 * ncols + 2 * numeric_cells

    source = table.get("source", "")
    if source == "pdfplumber":
        score += 15
    elif source == "camelot_lattice":
        score += 10
    # Below the ruled extraction, so a properly ruled table always wins when
    # both passes find the same one, but above nothing at all.
    elif source == "pdfplumber_text":
        score += 8
    elif source == "camelot_stream":
        score += 5

    label = infer_table_label(table)
    if label:
        score += 20
        table["label"] = label

    return score


def deduplicate_and_select_best_tables(table_candidates: list[dict]) -> list[dict]:
    filtered = []
    for table in table_candidates:
        score = score_table_candidate(table)
        if score > 0:
            if not table.get("label"):
                table["label"] = infer_table_label(table)
            filtered.append(table)

    grouped: dict[tuple, list[dict]] = {}
    unlabelled_by_page: dict[int, list[dict]] = {}

    for table in filtered:
        page = table.get("page")
        label = table.get("label")
        if label:
            key = (page, label)
            grouped.setdefault(key, []).append(table)
        else:
            if page is not None:
                unlabelled_by_page.setdefault(page, []).append(table)

    selected = []
    for _, candidates in grouped.items():
        best = max(candidates, key=score_table_candidate)
        selected.append(best)

    selected_pages = {t.get("page") for t in selected}
    for page, candidates in unlabelled_by_page.items():
        if page not in selected_pages and candidates:
            best = max(candidates, key=score_table_candidate)
            selected.append(best)

    selected.sort(key=lambda t: (t.get("page") or 9999, t.get("label") or ""))
    return selected


def table_dict_to_block(table: dict) -> str:
    lines = [
        "[TABLE_START]",
        f"Source: {table.get('source', 'unknown')}",
        f"Page: {table.get('page', 'unknown')}",
    ]
    label = table.get("label")
    if label:
        lines.append(f"Label: {label}")

    for row in table.get("rows", []):
        cleaned = [truncate_cell(cell) for cell in row]
        lines.append("\t".join(cleaned))

    lines.append("[TABLE_END]")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Marker integration
# ---------------------------------------------------------------------------

def find_marker_output(pdf_path: Path) -> Optional[Path]:
    """
    Look for a Marker-generated markdown file alongside the PDF.

    Marker typically outputs:  <stem>/<stem>.md  or  <stem>.marker.md
    We check several conventions.
    """
    stem = pdf_path.stem
    parent = pdf_path.parent

    candidates = [
        parent / f"{stem}.marker.md",
        parent / f"{stem}_marker.md",
        parent / stem / f"{stem}.md",
        parent / "marker_output" / f"{stem}.md",
        parent / "marker_output" / stem / f"{stem}.md",
    ]
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    return None


def read_marker_output(marker_path: Path) -> str:
    """Read and lightly clean Marker markdown output."""
    text = marker_path.read_text(encoding="utf-8", errors="ignore")
    text = clean_extracted_evidence_text(text, domain=OUTPUT_DOMAIN)
    return text


# ---------------------------------------------------------------------------
# Document readers
# ---------------------------------------------------------------------------

def read_pdf(path: Path) -> Tuple[str, List[Tuple[int, str]]]:
    """
    Read a PDF. If a Marker output file exists alongside it, use that for
    the main text (better layout recovery) while still extracting tables
    via pdfplumber/camelot.
    """
    marker_path = find_marker_output(path)

    # Always extract tables with pdfplumber/camelot regardless of Marker
    plumber_tables = extract_tables_pdfplumber(path)
    camelot_tables = extract_tables_camelot(path)
    best_tables = deduplicate_and_select_best_tables(plumber_tables + camelot_tables)

    table_blocks = []
    for t in best_tables:
        p_num = t.get("page")
        block = table_dict_to_block(t)
        table_blocks.append((p_num, block))

    if marker_path:
        print(f"  Using Marker output: {marker_path}")
        marker_text = read_marker_output(marker_path)
        # Append table blocks that may not be in Marker output
        table_section = ""
        if table_blocks:
            table_parts = ["\n\n[Extracted tables (pdfplumber/camelot)]"]
            for _, block in table_blocks:
                table_parts.append(block)
            table_section = "\n".join(table_parts)
        return (marker_text + table_section).strip(), table_blocks

    # Fallback: pypdf text extraction with interleaved tables
    reader = PdfReader(str(path))
    table_map: Dict[int, List[str]] = {}
    for t in best_tables:
        p_num = t.get("page")
        block = table_dict_to_block(t)
        if p_num is not None:
            table_map.setdefault(p_num, []).append(block)

    pages = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as e:
            text = f"[Could not extract text from page {i}: {e}]"

        cleaned_page_text = clean_extracted_evidence_text(text, domain=OUTPUT_DOMAIN)
        page_parts = [f"\n\n[Page {i}]\n{cleaned_page_text}"]

        if i in table_map:
            page_parts.append("\n[Extracted tables]")
            page_parts.extend(table_map[i])

        pages.append("\n".join(part for part in page_parts if part.strip()))

    return "\n".join(pages).strip(), table_blocks


def read_txt(path: Path) -> Tuple[str, List[Tuple[int, str]]]:
    return clean_extracted_evidence_text(
        path.read_text(encoding="utf-8", errors="ignore"),
        domain=OUTPUT_DOMAIN,
    ), []


def read_docx(path: Path) -> Tuple[str, List[Tuple[int, str]]]:
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    return clean_extracted_evidence_text(text, domain=OUTPUT_DOMAIN), []


def read_csv_file(path: Path, max_rows: int = 120) -> Tuple[str, List[Tuple[int, str]]]:
    rows = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i >= max_rows:
                break
            rows.append("\t".join("" if x is None else str(x) for x in row))
    text = f"[CSV preview: first {min(len(rows), max_rows)} rows]\n" + "\n".join(rows)
    return clean_extracted_evidence_text(text, domain=OUTPUT_DOMAIN), []


def read_xlsx(path: Path, max_rows_per_sheet: int = 80) -> Tuple[str, List[Tuple[int, str]]]:
    wb = load_workbook(filename=str(path), data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"\n[Sheet: {ws.title}]")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows_per_sheet:
                break
            out.append("\t".join("" if v is None else str(v) for v in row))
    return clean_extracted_evidence_text("\n".join(out), domain=OUTPUT_DOMAIN), []


# ---------------------------------------------------------------------------
# Marginal line numbers
# ---------------------------------------------------------------------------
# Manuscripts under review carry line numbers down the margin, and the PDF text
# layer puts them inside the sentences: "The number of clusters (K) 174 was
# chosen based on the elbow heuristic". Nothing downstream expects that. On one
# submitted manuscript it made every quotation spanning a line break fail the
# citation check -- three correct quotations reported as unverifiable, the
# reliability banner fired, and good concerns pushed down the evidence table --
# and it fed line 129 to the model as a count of races. Published papers do not
# have them, which is why the corpus never showed this.
#
# Stripping every trailing integer would eat real data: a table row ending in a
# count looks identical to a line number. So the numbers have to be recognised
# as a sequence: only a run of trailing integers that climbs by small steps
# through the document is removed, and only when there are enough of them for
# the pattern to be a numbering scheme rather than a coincidence.
_TRAILING_NUMBER_RE = re.compile(r"^(?P<body>.*\S)[ \t]+(?P<number>\d{1,4})[ \t]*$")
MIN_LINE_NUMBERS = 25          # fewer than this is not a numbering scheme
MAX_LINE_NUMBER_STEP = 4       # allows for blank or unnumbered lines between


def fold_mathematical_letters(text: str) -> str:
    """
    Map Unicode Mathematical Alphanumeric letters onto their ASCII equivalents.

    Manuscripts written in Word set variables in this range: the italic p of
    "p < 0.05" is U+1D45D, not "p", and every pattern looking for ASCII misses
    it silently.
    """
    if not text:
        return text
    out = []
    for char in text:
        if 0x1D400 <= ord(char) <= 0x1D7FF:
            folded = unicodedata.normalize("NFKC", char)
            out.append(folded if folded else char)
        else:
            out.append(char)
    return "".join(out)


def strip_marginal_line_numbers(text: str) -> Tuple[str, int]:
    """
    Remove marginal line numbers from extracted text.

    Returns the cleaned text and how many were removed. Text that is not line
    numbered comes back unchanged, and a trailing number that does not belong
    to the run -- a table cell, a year, a count -- is left alone.
    """
    if not text:
        return text, 0
    lines = text.split("\n")
    candidates = []
    for index, line in enumerate(lines):
        found = _TRAILING_NUMBER_RE.match(line)
        if found:
            candidates.append((index, int(found.group("number")), found))

    # Which candidates belong to a climbing run. A restart (the number drops)
    # begins a new run, so a document numbered per page is handled too.
    in_run = set()
    run: List[int] = []
    last = None
    for position, (_, value, _found) in enumerate(candidates):
        if last is not None and 0 < value - last <= MAX_LINE_NUMBER_STEP:
            run.append(position)
        else:
            if len(run) >= 3:
                in_run.update(run)
            run = [position]
        last = value
    if len(run) >= 3:
        in_run.update(run)

    if len(in_run) < MIN_LINE_NUMBERS or len(in_run) < len(candidates) * 0.5:
        return text, 0

    removed = 0
    kept: List[Tuple[int, int]] = []
    for position in sorted(in_run):
        index, value, found = candidates[position]
        lines[index] = found.group("body")
        kept.append((index, value))
        removed += 1

    removed += _strip_abutting_line_numbers(lines, kept)
    return "\n".join(lines), removed


def _strip_abutting_line_numbers(lines: List[str],
                                 kept: List[Tuple[int, int]]) -> int:
    """
    Remove a line number that is glued to the text with no space before it.

    _TRAILING_NUMBER_RE requires whitespace between the body and the number,
    which is right: without it, "n=173" would lose its value. But a line
    ending on an operator gets the number attached directly, and on
    RPAN-2026-0184 the women's 200 m sample extracted as

        ... and women's 200m (n =116
        173). Swimmers were deidentified

    where 116 is the marginal number for that line and 173 is the sample size
    wrapping onto the next. 702 numbers were stripped from that document and
    this was the one that was not, so two generations reported the sample as
    "116 173" and the citation check endorsed it, because that is what the text
    said. An extraction fault the checks cannot see is the worst kind.

    Only a value MISSING from an otherwise complete run is removed, only at the
    end of a line, and only on a line lying between the two survivors that
    bracket it. A figure that happens to equal a line number is untouched,
    because every other value in the run is accounted for.
    """
    if len(kept) < MIN_LINE_NUMBERS:
        return 0
    removed = 0
    for (index_a, value_a), (index_b, value_b) in zip(kept, kept[1:]):
        if value_b - value_a != 2:
            continue
        missing = value_a + 1
        pattern = re.compile(rf"(?<=[^\s\d]){missing}[ \t]*$")
        for index in range(index_a, min(index_b + 1, len(lines))):
            line, count = pattern.subn("", lines[index], count=1)
            if count:
                lines[index] = line.rstrip()
                removed += 1
                break
    return removed


# Set by load_document so the report can say what was done to the text before
# anything read it. A silent repair is the failure mode this pipeline keeps
# finding in itself.
LAST_EXTRACTION_NOTES: List[str] = []


# ---------------------------------------------------------------------------
# Reading damaged tables from the page image
# ---------------------------------------------------------------------------
# pdfplumber recovers the numbers in a table more often than it recovers the
# shape. Two measured examples: a Kruskal-Wallis table kept every value but
# scattered "p-Value" and "Freedom" onto orphan lines and split "Men's / 100m /
# Freestyle" away from the rows it labels, so reports could only refer to "the
# first men's block"; and a Bayesian results table broke "Parabolic" into
# "Paraboli" and "c" on separate lines. A model with a vision encoder read both
# pages correctly, invented nothing, and preserved a garbled cell ("s0") rather
# than tidying it away -- which is the behaviour that makes it usable, because
# the garbled cell turned out to be what the paper prints.
#
# The transcription is added BESIDE the text-layer block, never in place of it,
# and labelled Source: vision. Two reasons. A reader can see which reading a
# concern rests on; and because the transcription joins the source text, the
# citation check verifies quotations against it like any other evidence. A
# reading with nothing behind it is what this pipeline spends its time removing.
VISION_TABLES = os.environ.get("QWEN_VISION_TABLES", "").strip().lower() in (
    "1", "true", "yes", "on")
VISION_DPI = int(os.environ.get("QWEN_VISION_DPI", "200"))

# VISION_TABLES is the process default, set by --vision at start-up. A single
# review can override it, so the browser can offer vision per run in the same
# way it offers a reasoning mode, without restarting the server.
_VISION_OVERRIDE: Optional[bool] = None


def vision_enabled() -> bool:
    return VISION_TABLES if _VISION_OVERRIDE is None else _VISION_OVERRIDE


@contextmanager
def vision(enabled: Optional[bool]):
    """Force vision on or off for the duration of one review."""
    global _VISION_OVERRIDE
    previous = _VISION_OVERRIDE
    _VISION_OVERRIDE = enabled
    try:
        yield
    finally:
        _VISION_OVERRIDE = previous
MAX_VISION_PAGES = int(os.environ.get("QWEN_VISION_MAX_PAGES", "6"))

VISION_TABLE_PROMPT = (
    "Transcribe every table on this page exactly as printed. Preserve the "
    "column headers, the row labels and every numeric cell, one row per line "
    "with columns separated by ' | '. Repeat the row label on every row it "
    "applies to. Do not summarise, do not interpret, and do not correct "
    "anything that looks wrong: if a cell is malformed, transcribe it as it "
    "appears. If there is no table on the page, say 'No table on this page' "
    "and nothing else."
)

# A single letter or two stranded on its own line is a word broken by the
# extractor. A line with no digit at all, inside a numeric table, is header or
# label debris. A token like "s0" is a cell that lost or gained a character.
_ORPHAN_FRAGMENT_RE = re.compile(r"^\s*[A-Za-z]{1,2}\s*$")
_MIXED_TOKEN_RE = re.compile(r"(?<![\w.])(?:[A-Za-z]\d|\d[A-Za-z])(?![\w.])")


def table_looks_damaged(block_text: str) -> str:
    """
    Why a table block looks structurally broken, or "" if it does not.

    Deliberately conservative: each test comes from damage actually seen, and
    a table that merely looks unusual is left alone. Rendering a page and
    asking a model to read it costs a minute; doing it for every table would
    cost more than it returns.
    """
    lines = [line for line in block_text.splitlines()
             if line.strip() and not line.startswith(("[TABLE", "Source:", "Page:", "Label:"))]
    if len(lines) < 3:
        return ""
    orphans = [line for line in lines if _ORPHAN_FRAGMENT_RE.match(line)]
    if orphans:
        return f"{len(orphans)} word(s) broken across lines"
    mixed = _MIXED_TOKEN_RE.findall(block_text)
    if mixed:
        return f"malformed cell(s) such as {mixed[0]!r}"
    digitless = [line for line in lines if not re.search(r"\d", line)]
    if len(digitless) >= max(3, len(lines) * 0.3):
        return (f"{len(digitless)} of {len(lines)} rows carry no numbers, so "
                f"headers and labels are detached from their data")
    return ""


def _render_page_png(pdf_path: Path, page_number: int, dpi: int) -> bytes:
    """One 1-based page as PNG bytes, or b"" if it cannot be rendered."""
    try:
        import io
        import pypdfium2 as pdfium
    except ImportError:
        return b""
    try:
        document = pdfium.PdfDocument(str(pdf_path))
        try:
            if not 1 <= page_number <= len(document):
                return b""
            image = document[page_number - 1].render(scale=dpi / 72).to_pil()
            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            return buffer.getvalue()
        finally:
            document.close()
    except Exception:                                       # noqa: BLE001
        return b""


def _page_of_block(block_text: str) -> Optional[int]:
    found = re.search(r"^Page:\s*(\d+)", block_text, re.M)
    return int(found.group(1)) if found else None


def rescue_damaged_tables(pdf_path: Path,
                          table_blocks: List[Tuple[int, str]]
                          ) -> Tuple[List[Tuple[int, str]], List[str]]:
    """
    Re-read structurally damaged tables from the page image.

    Returns the table blocks with any transcriptions appended, and notes for
    the report header. Never raises: vision is an enhancement, and a review
    that fails for want of a picture is worse than one that never had it.
    """
    if not vision_enabled():
        return table_blocks, ["vision was off, so no table was re-read from "
                              "its page image"]
    if not table_blocks:
        return table_blocks, ["vision was on, but no table was extracted for "
                              "it to check"]
    try:
        from llm_backend import describe_image, vision_available
        if not vision_available():
            return table_blocks, ["table images were not read: no llama-server "
                                  "backend in use"]
    except Exception:                                       # noqa: BLE001
        return table_blocks, []

    wanted: Dict[int, str] = {}
    for _page, block in table_blocks:
        if "Source: vision" in block:
            continue
        reason = table_looks_damaged(block)
        page_number = _page_of_block(block)
        if reason and page_number and page_number not in wanted:
            wanted[page_number] = reason
    if not wanted:
        return table_blocks, ["vision was on and no extracted table looked "
                              "damaged, so none was re-read"]

    added: List[Tuple[int, str]] = []
    notes: List[str] = []
    for page_number in sorted(wanted)[:MAX_VISION_PAGES]:
        png = _render_page_png(pdf_path, page_number, VISION_DPI)
        if not png:
            continue
        transcription = describe_image(png, VISION_TABLE_PROMPT)
        if not transcription or "no table on this page" in transcription.lower():
            continue
        added.append((page_number,
                      "[TABLE_START]\n"
                      f"Source: vision\nPage: {page_number}\n"
                      "Note: read from the page image because the text layer "
                      f"was damaged ({wanted[page_number]}). The text-layer "
                      "version of this table is also present.\n"
                      f"{transcription}\n[TABLE_END]"))
    if added:
        notes.append(
            f"{len(added)} table(s) re-read from the page image because the "
            f"text layer was damaged (page(s) "
            f"{', '.join(str(p) for p, _ in added)}); both readings are in the "
            f"evidence, labelled Source: vision"
        )
    else:
        # Asked for and got nothing: almost always llama-server running without
        # a projector. Said out loud, because a review that quietly fell back
        # to the text layer looks identical to one that never asked.
        notes.append(
            f"{len(wanted)} damaged table(s) could not be re-read from the "
            f"page image; the server may be running without --mmproj. The "
            f"text-layer reading was used."
        )
    return table_blocks + added, notes


def load_document(path: Path) -> Tuple[str, List[Tuple[int, str]]]:
    text, tables = _read_document(path)
    LAST_EXTRACTION_NOTES.clear()
    text, removed = strip_marginal_line_numbers(text)
    if removed:
        LAST_EXTRACTION_NOTES.append(
            f"{removed} marginal line number(s) removed before analysis; the "
            f"manuscript is line numbered for review"
        )
    if path.suffix.lower() == ".pdf":
        try:
            tables, notes = rescue_damaged_tables(path, tables)
            LAST_EXTRACTION_NOTES.extend(notes)
        except Exception:                                   # noqa: BLE001
            pass
    return text, tables


def _read_document(path: Path) -> Tuple[str, List[Tuple[int, str]]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix in {".txt", ".md"}:
        return read_txt(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix == ".csv":
        return read_csv_file(path)
    if suffix in {".xlsx", ".xlsm"}:
        return read_xlsx(path)
    raise ValueError(f"Unsupported file type: {path.name}")


def split_text(text: str, max_chars: int = MAX_SECTION_CHARS) -> List[str]:
    if not text:
        return ["[No extractable text found]"]
    if len(text) <= max_chars:
        return [text]
    paras = text.split("\n\n")
    chunks: List[str] = []
    current: List[str] = []
    current_len = 0
    for para in paras:
        para = para.strip()
        if not para:
            continue
        add_len = len(para) + 2
        if current and current_len + add_len > max_chars:
            chunks.append("\n\n".join(current).strip())
            current = [para]
            current_len = len(para)
        else:
            current.append(para)
            current_len += add_len
    if current:
        chunks.append("\n\n".join(current).strip())
    return chunks if chunks else ["[No extractable text found]"]


def collect_input_paths(paths: Iterable[str]) -> List[Path]:
    supported = {".pdf", ".txt", ".md", ".docx", ".csv", ".xlsx", ".xlsm"}
    results: List[Path] = []
    for raw in paths:
        p = Path(raw).expanduser()
        if p.is_dir():
            for child in sorted(p.rglob("*")):
                if child.is_file() and child.suffix.lower() in supported:
                    results.append(child)
        elif p.is_file() and p.suffix.lower() in supported:
            results.append(p)
    return results


# ---------------------------------------------------------------------------
# LLM interaction
# ---------------------------------------------------------------------------

def apply_chat_template_compat(tokenizer, user_text: str) -> str:
    """
    Apply chat template in a model-agnostic way.
    Works for Qwen and non-Qwen models (e.g., Ouro).

    ``enable_thinking`` was hard-coded to False here, which silently outranked
    the per-review reasoning choice: every prompt in the pipeline was built
    with thinking off, whatever the browser asked for, so the mode was
    recorded in the header and never applied. It now follows the resolved
    setting, which is still False unless a review or the environment asks
    otherwise.
    """
    messages = [{"role": "user", "content": user_text}]

    if getattr(tokenizer, "chat_template", None) is None:
        return user_text

    try:
        return tokenizer.apply_chat_template(
            messages,
            tokenize=False,
            add_generation_prompt=True,
            enable_thinking=thinking_enabled(),
        )
    except TypeError:
        return tokenizer.apply_chat_template(
            messages,
            tokenize=False,
            add_generation_prompt=True,
        )

def make_default_sampler(temperature: float = None):
    """The standard sampler, or a warmer one for a repeat synthesis pass."""
    return make_sampler(TEMPERATURE if temperature is None else temperature,
                        top_p=TOP_P, top_k=TOP_K,
                        repetition_penalty=REPETITION_PENALTY,
                        presence_penalty=PRESENCE_PENALTY)


# A repeat pass at the same low temperature mostly repeats itself, which makes
# the union worth nothing. Later passes are sampled warmer so they explore a
# different part of the space; the first pass stays at TEMPERATURE and remains
# the base report.
REPEAT_PASS_TEMPERATURE = float(os.environ.get("REPEAT_PASS_TEMPERATURE", "0.7"))


def clean_model_output(text: str) -> str:
    # Reasoning is disabled via the chat template, but strip any that leaks
    # through so a <think> span can never reach a review report.
    text = strip_reasoning(text)
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("```markdown", "").replace("```text", "").replace("```", "")
    text = re.sub(r" +", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# System prompt
# ---------------------------------------------------------------------------

SYSTEM_STYLE = """You are a rigorous academic peer review assistant.

Core rules:
- Work strictly from the provided material.
- Use British English spelling.
- Distinguish clearly between:
  1. Directly supported by the text
  2. Reasonable inference
  3. Speculation / missing information
- Do not invent sample sizes, analyses, results, measures, design features, claims, contradictions, artefacts, or errors.
- Prioritise internal validity, measurement quality, causal assumptions, statistical reasoning, modelling choices, transparency, and reporting.
- Be critical but fair.
- Use concise markdown headings and bullet points.
- Do NOT use LaTeX or dollar-sign math notation.
- If tabular material is present, treat it as available evidence unless clearly unreadable.
- If a table is imperfectly extracted, say "partly extracted" or "some entries may be ambiguous", not "missing", unless it is genuinely absent.
- Ignore minor OCR artefacts in variable names (e.g., WHT.5R, WHT$5R, WHT^0.5R all refer to Waist/Height^0.5). Do not critique these as notation errors.
- Be aware that OCR sometimes extracts '=' as 'Z' (e.g., r Z 0.35 means r = 0.35). Interpret accordingly.
- EQUATION OCR RULE: Ignore garbled equations, subscripts, or superscripts (e.g., vconsi, u(3)0, LEV2bout). These are known PDF extraction failures. Do NOT claim equations are unreadable, missing, or unverifiable due to formatting.
- PEDANTRY RULE: Ignore trivial typographical inconsistencies in reference lists, citations, or software version numbers (e.g., MLwiN version 2.1 vs 2.10). Do not mention them in the review.
- PARADIGM RULE: MCMC is a computational method, not a paradigm. If MCMC, MLwiN, Gibbs sampling, or DIC are mentioned without explicit Bayesian terminology (e.g., priors, posteriors, Bayes factors), you MUST treat the analysis as STRICTLY FREQUENTIST using MCMC for estimation. Do not hedge, do not describe the framework as 'unclear', and do not critique the absence of Bayesian diagnostics.

Generality rule:
- Apply the same standards across frequentist, Bayesian, predictive, and causal analyses.
- Do not assume a method-specific problem unless the extracted material supports it.

Claim-strength ladder:
- Prefer "could be clarified" over "is wrong" unless the text directly supports the stronger claim.
- Prefer "reported briefly" or "not fully detailed" over "unreported" when something is mentioned but not elaborated.
- Prefer "interpretation may need clarification" over "data entry artefact" unless there is direct evidence of corruption.
- Prefer "could not verify from the extracted material" over "missing" when extraction is partial.

Manifest-awareness rule:
- An evidence manifest is provided listing what the extraction pipeline has confirmed is present (e.g., confidence intervals, p-values, model fit statistics).
- Do NOT claim a statistic type is "absent", "missing", or "not reported" if the manifest says it is present.
- If the manifest says "Confidence intervals reported: True", you may still note that CIs are not provided for every individual test, but you must NOT make blanket claims that CIs are missing.
- Frame any gap as "CIs are reported for some analyses but not all" or "CIs for [specific test] could not be identified", not "CIs are not reported".
- Apply the same principle to standard errors, p-values, effect sizes, and model fit statistics.

Design-verification rule:
- For observational or survey data that explicitly describes a survey design (e.g., Health Survey for England, NHANES, stratified random probability sample), you may raise survey-specific concerns such as survey weights, clustering adjustments, or medication sensitivity analyses.
- For other observational or experimental studies that are NOT survey-based (e.g., a study recruiting swimmers from clubs, a clinical trial, a lab experiment), do NOT raise survey-specific concerns. Survey weights and clustering adjustments are only relevant to complex survey designs.
- In all cases, distinguish between "verify whether X was handled" and "X was omitted". The absence of a description in the extracted material does not confirm the absence of the procedure in the study.

Paper-specificity rule:
- Every criticism, concern, or suggestion you raise must be specific to the actual paper being reviewed, not a generic checklist item.
- Do NOT include generic scrutiny points (e.g., "Verify whether missing data mechanisms were addressed", "Clarify the handling of missing data") unless the extracted material gives specific reason to suspect that missing data is a problem for this particular paper (e.g., varying sample sizes across analyses, or explicit mention of exclusions).
- Do NOT suggest survey weights, clustering adjustments, or population-representativeness concerns unless the paper describes a population-level survey design.
- Do NOT suggest "verify whether interactions were tested" unless the paper's claims specifically depend on interaction effects.
- Before including any "verify whether..." point, ask: does the extracted material provide specific evidence that this is a concern for THIS paper? If not, omit it.

Numeric provenance rule:
- Academic papers frequently cite statistics from OTHER studies for context (e.g., "In a previous study by X et al., N=405 decisions were analysed").
- Do NOT conflate numbers from cited/referenced studies with the current study's data.
- Do NOT claim a "discrepancy" between a number from a cited study and a number from the current study.
- Before identifying any numeric inconsistency, verify that both numbers refer to the SAME dataset/analysis within the current study.
- If a number appears near "et al.", a year in parentheses, or "previous study/larger dataset", it likely belongs to a different study.
"""

COMMON_DIAGNOSTIC_ALIASES = """Examples of reported diagnostics and checks across paradigms:
- frequentist/model-based: VIF, residual plots, heteroscedasticity checks, AIC, BIC, confidence intervals, sensitivity analyses, multiple imputation, cross-validation
- Bayesian: priors, prior predictive checks, posterior predictive checks, R-hat, effective sample size, divergences, LOO, WAIC, Bayes factors, prior sensitivity
- prediction/ML: calibration, discrimination, ROC/AUC, Brier score, internal validation, external validation, optimism correction
- causal/design: DAGs, assumptions, preregistration, protocol deviations, missing-data mechanisms
- MCMC as estimation (non-Bayesian): convergence diagnostics, burn-in, DIC, variance components with standard errors, Wald-type tests
- distributional/GAMLSS: worm plots, randomised quantile residuals, GAIC, distribution family comparison
"""


# ---------------------------------------------------------------------------
# Review prompts (method-aware)
# ---------------------------------------------------------------------------

def review_chunk(model, tokenizer, chunk: DocChunk, method_expectations: str = "",
                 manifest_summary: str = "") -> str:
    context_block = ""
    if method_expectations:
        context_block += f"\n{method_expectations}\n"
    if manifest_summary:
        context_block += f"\nEvidence manifest (what has been extracted from this file):\n{manifest_summary}\n"

    user_text = f"""
{SYSTEM_STYLE}

{COMMON_DIAGNOSTIC_ALIASES}
{context_block}
Review this manuscript chunk.

Source: {chunk.source_name}
Chunk: {chunk.chunk_id}

Return markdown under exactly these headings:
## Supported points
## Reasonable inferences
## Missing or unclear
## Statistical / methodological concerns
## Table-specific notes
## Questions raised by this chunk

Instructions:
- MCMC is an estimation method, not a paradigm. If the text mentions MCMC, MLwiN, or DIC but does not explicitly mention "priors" or "posteriors", it is a frequentist analysis. Do not evaluate it against Bayesian reporting standards and do not state that the framework is unclear.
- Where tables are present, use them as evidence.
- If a table appears partially extracted, say so explicitly.
- Do not claim a table, coefficient, equation, or diagnostic is missing if it is present in the chunk.

Internal-consistency checks. These are among the most useful things you can find, because they are verifiable from the text alone rather than matters of opinion. Check for, and report, any of the following that this chunk actually shows:
- A numeric statement that contradicts itself. For example, a value offered as an exception to a threshold that does not in fact breach it ("all were <99%, with the exception of X (90.85%)" - 90.85 is not an exception to <99, so either the value or the direction of the inequality is wrong).
- A sentence that begins by describing one variable and draws a conclusion about a different one, or a comparison whose stated variable does not match the variable in the accompanying table.
- A number in the narrative that disagrees with the corresponding cell in a table, including a coefficient, an interval, a sample size or a total that does not add up.
- A claim about the model that the coefficient table contradicts: a covariate the text says was adjusted for that does not appear among that model's predictors, or a term present in one model but silently absent from another that is described the same way.
- A quantity described as controlled for, held constant, or included, where the reported model for that subgroup does not list it.
Report each with the exact quoted text and, where relevant, the conflicting table value. If the chunk shows no such inconsistency, say nothing about this; do not manufacture one.
- If the evidence manifest above says a statistic type is present (e.g., "Confidence intervals reported: True"), do not make blanket claims that CIs are missing. You may note that CIs are not provided for a specific test in this chunk, but frame the gap narrowly.
- For observational or survey data, distinguish between "verify whether X was handled" and "X was omitted". Use "Verify whether..." unless the text explicitly confirms the omission.
- Do NOT raise survey-specific concerns (survey weights, clustering adjustments) unless the paper describes a population-level survey design.
- Do NOT raise generic scrutiny points (missing data mechanisms, interaction testing) unless this specific chunk provides evidence that the concern is relevant.
- Do not use LaTeX or dollar-sign notation.
- Be concise and avoid repetition.
- Do not ask for information already clearly reported in the chunk.
- Do not convert obvious PDF/OCR extraction artefacts into substantive methodological concerns.
- Do not critique garbled equations, subscripts, or superscripts (e.g., vconsi, u(3)0). Treat them as text extraction limits, not author errors.
- Do not complain about trivial typographical discrepancies in software versioning or references (e.g., MLwiN 2.1 vs 2.10).
- Do not describe a parameterisation choice or scaling convention as an error unless the text directly contradicts itself.
- Do not infer data entry artefacts from a small coefficient with a large test statistic alone.
- If an equation or model form is shown in the chunk, do not say the functional form is absent or unclear.
- Distinguish between "not reported", "reported briefly", and "reported but not fully interpretable from this chunk".
- If the text refers to an equation, figure, table, appendix item, or model specification that is not clearly recoverable from extraction, do not call it missing.
- Instead say that it appears to be present in the PDF but could not be evaluated confidently from the extracted material.
- Match your diagnostic expectations to the method classification provided above. Do not request diagnostics from the wrong framework.
- Do not request VIF for single-predictor models.

Material:
\"\"\"
{chunk.text}
\"\"\"
""".strip()
    prompt = apply_chat_template_compat(tokenizer, user_text)
    sampler = make_default_sampler()
    out = generate(
        model, tokenizer, prompt=prompt,
        max_tokens=SECTION_MAX_TOKENS, sampler=sampler, verbose=False
    )
    return clean_model_output(out)


# ---------------------------------------------------------------------------
# Which version of this file is actually running
# ---------------------------------------------------------------------------
# server.py imports this module once at start-up, so edits made afterwards do
# not take effect until it is restarted. Several reviews were produced by a
# process running superseded code with nothing in the output to show it. The
# report now carries the fingerprint of the code that produced it, and says so
# when the file on disk has moved on.

def _fingerprint_of(*paths) -> str:
    """
    A short hash over the files that decide what a review looks like.

    This covered review_pipeline.py alone, so a change confined to server.py --
    where the synthesis loop, the validation order and the report header all
    live -- left the stamp unchanged, and the header then claimed a version the
    running code no longer matched. llm_backend.py was the same gap one step
    further out: it decides the sampler, the thinking mode and the reasoning
    accounting. The service script fingerprints all three files; this now
    agrees with it.
    """
    digest = hashlib.sha1()
    for path in paths:
        try:
            digest.update(Path(path).read_bytes())
        except Exception:
            return "unknown"
    return digest.hexdigest()[:8]


PIPELINE_PATH = Path(__file__).resolve()
SERVER_PATH = PIPELINE_PATH.parent / "server.py"
BACKEND_PATH = PIPELINE_PATH.parent / "llm_backend.py"
# The deterministic checks and the design registry decide what a report says as
# much as the prompts do, so they belong in the fingerprint. Without this a
# pattern could change and two reports would still claim to be comparable.
CHECKS_PATH = PIPELINE_PATH.parent / "manuscript_checks.py"
DESIGN_PATH = PIPELINE_PATH.parent / "design_expectations.py"
PIPELINE_VERSION = _fingerprint_of(PIPELINE_PATH, SERVER_PATH, BACKEND_PATH,
                                   CHECKS_PATH, DESIGN_PATH)


def stale_module_warning() -> str:
    """A warning when the running code is older than the file on disk."""
    current = _fingerprint_of(PIPELINE_PATH, SERVER_PATH, BACKEND_PATH,
                              CHECKS_PATH, DESIGN_PATH)
    if current in ("unknown", PIPELINE_VERSION):
        return ""
    return (
        f"WARNING: this review was produced by pipeline version "
        f"{PIPELINE_VERSION}, but the code on disk is now version {current}. "
        f"The server was started before that change, so recent edits are NOT "
        f"in effect here. Restart it with ./scripts/qwen_service.sh restart "
        f"and run the review again."
    )


# ---------------------------------------------------------------------------
# Guard: refuse to review this pipeline's own output
# ---------------------------------------------------------------------------
# Feeding a review or an evidence appendix back in produces a report that looks
# entirely normal and is about the wrong document: the manifest picks up the
# appendix's own prose, so flags such as "Randomisation described" come out true
# because the words appear in the report's furniture, and the reviewer complains
# that diagnostics are missing when it is only reading a summary of them.
DERIVED_INPUT_MARKERS = (
    "# Local peer-review report",
    "# Evidence appendix",
    "Evidence manifest for:",
    "## Chunk previews",
    "Detected model/result blocks",
    "Block counts:",
)


def detect_derived_input(text: str) -> List[str]:
    """Markers showing the text is output from this pipeline, not a manuscript."""
    if not text:
        return []
    return [marker for marker in DERIVED_INPUT_MARKERS if marker in text]


# ---------------------------------------------------------------------------
# Read each model's predictor list straight out of the coefficient tables
# ---------------------------------------------------------------------------
_SKIP_ROW_LABELS = re.compile(
    r"^\s*(?:r2|r\^?2|r-?squared|observations?|n|quantile|estimates?|ci|"
    r"probability|intercept|note|source|page|label)\b",
    re.I,
)
_POLY_QUALIFIER = re.compile(r"\b(?:quadratic|cubic|quartic|linear|poly\w*|spline)\b", re.I)
# In a stacked coefficient table each model block closes with its fit
# statistic. Treating that as the terminator stops a following descriptive
# table being swallowed into the last model.
_MODEL_TERMINATOR = re.compile(
    r"^\s*(?:r2|r\^?2|r-?squared|observations?|looic|loo|waic|aic|bic|deviance|dic)\b",
    re.I,
)
_TRAILING_UNIT = re.compile(r"\b(?:cm|mm|kg|m|s|yrs?|years?|pct|percent)\s*$", re.I)
# A model section in a stacked coefficient table is conventionally labelled
# "<something> Predictors" or "<something> Model". Requiring that keeps
# descriptive tables out: their wrapped row labels ("Body", "Sitting", "Bi-")
# otherwise became models of their own and buried the real signal.
_SECTION_HEADING = re.compile(
    r"^[A-Z][A-Za-z /-]{2,40}\b(?:predictors?|model|models|equation)\s*$", re.I
)


def _normalise_predictor(label: str) -> str:
    """Fold a coefficient row label to a comparable predictor name."""
    text = label.strip()
    text = re.sub(r"^\s*l[no]g?\s*\(", "(", text, flags=re.I)   # ln( / log(
    text = _POLY_QUALIFIER.sub(" ", text)
    text = re.sub(r"\[[^\]]*\]|\([%\s]*\)", " ", text)          # units
    text = re.sub(r"[\[\]{}()]", " ", text)
    text = re.sub(r"\b\d+\b", " ", text)                        # term index
    text = re.sub(r"[^A-Za-z\s-]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = _TRAILING_UNIT.sub("", text).strip()   # "... breadth cm" -> "... breadth"
    return text.lower()


def extract_model_predictors(
    table_blocks: List[Tuple[int, str]]
) -> "Dict[str, List[str]]":
    """
    Map each model heading in the coefficient tables to its predictor list.

    Reading this off the tables mechanically, rather than asking the model to
    notice it, is the point: a covariate present in three of four models and
    absent from the fourth is exactly the kind of asymmetry a reviewer should
    see, and exactly the kind a prose summary loses.
    """
    models: "Dict[str, List[str]]" = {}
    current = None
    for _, block_text in table_blocks:
        for raw_line in block_text.splitlines():
            cells = [c.strip() for c in raw_line.split("\t")]
            nonempty = [c for c in cells if c]
            if not nonempty:
                continue
            if len(nonempty) == 1:
                candidate = nonempty[0]
                if _SECTION_HEADING.match(candidate) and not re.search(r"\d", candidate):
                    current = candidate
                    models.setdefault(current, [])
                continue
            if current is None:
                continue
            label = nonempty[0]
            if _MODEL_TERMINATOR.match(label):
                current = None      # this model's block has ended
                continue
            if _SKIP_ROW_LABELS.match(label):
                continue
            if not any(re.search(r"\d", c) for c in nonempty[1:]):
                continue
            name = _normalise_predictor(label)
            if name and name not in models[current]:
                models[current].append(name)
    return {k: v for k, v in models.items() if v}


def summarise_model_predictors(table_blocks: List[Tuple[int, str]]) -> str:
    """A plain statement of what each model contains, and what it omits."""
    models = extract_model_predictors(table_blocks)
    if len(models) < 2:
        return ""

    lines = [
        "Predictor lists read directly from the coefficient tables. These are "
        "what each model actually contains, not what the narrative says it "
        "contains; where the two disagree, the table is the evidence:",
    ]
    for name, predictors in models.items():
        lines.append(f"- {name}: {', '.join(predictors)}")

    union: List[str] = []
    for predictors in models.values():
        for p in predictors:
            if p not in union:
                union.append(p)

    asymmetries = []
    for predictor in union:
        present = [n for n, ps in models.items() if predictor in ps]
        absent = [n for n, ps in models.items() if predictor not in ps]
        if present and absent:
            asymmetries.append(
                f"- {predictor}: present in {', '.join(present)}; "
                f"ABSENT from {', '.join(absent)}"
            )

    if asymmetries:
        lines.append("")
        lines.append(
            "Predictors that appear in some of these models but not others. "
            "Check each against what the manuscript claims was adjusted for or "
            "held constant:"
        )
        lines.extend(asymmetries)
    return "\n".join(lines)


# A cell holding a number's integer part with the fraction stranded on the
# next row: "13." above, "5" below, for a value of 13.5. Multi-line cells in a
# wide descriptive table do this, and the value is then unreadable. Left
# unflagged, a reviewer invented "13.79s" from such a fragment and raised a
# data-integrity concern about it.
_ORPHAN_STEM_RE = re.compile(r"^-?\d+\.$")


# A PDF whose word spacing is lost, or whose columns are interleaved into
# single lines, yields text like "modelwhereWrepresentsthey-axisintercept".
# Measured across a fifteen-paper corpus, clean documents scored 0.00% and
# degraded ones 3.1-16.3%, so 1% separates them with room to spare.
RUN_ON_TOKEN_THRESHOLD = 0.01


def document_extraction_quality(text: str) -> str:
    """A warning when the document's own text did not extract cleanly."""
    tokens = re.findall(r"[A-Za-z]{2,}", text or "")
    if len(tokens) < 200:
        return ""
    run_on = sum(1 for t in tokens if len(t) > 20) / len(tokens)
    if run_on < RUN_ON_TOKEN_THRESHOLD:
        return ""
    return (
        f"WARNING: this document did not extract cleanly. {run_on * 100:.1f}% of "
        "words have run together, which happens when a PDF loses its word "
        "spacing or when two columns are merged into single lines. Sentences "
        "may be interleaved from different columns and numbers may be joined to "
        "words. Treat absence as unproven: if something appears to be missing, "
        "say that it could not be recovered from the extracted text, not that "
        "the authors omitted it. Report this as an extraction limit."
    )


# Some journal PDFs split the word: this corpus's Frontiers paper extracts its
# captions as "T able 2", so every table in it was announced-but-unlabelled and
# the missing-tables warning could never fire.
_TABLE_NUMBER_RE = re.compile(
    r"(?mi)^[ \t]*(?:supplementary\s*)?t\s?able\s*([0-9]+|[IVXL]+)\b"
)


def _table_numbers(text: str) -> set:
    """Distinct table numbers announced at the start of a line."""
    return {m.group(1).lower() for m in _TABLE_NUMBER_RE.finditer(text or "")}


def missing_tables_warning(text: str, table_blocks: List[Tuple[int, str]]) -> str:
    """A warning when the paper announces tables the extractor did not find."""
    # Counting caption LINES over-reports badly: sub-labels such as "Table 3a"
    # and wrapped in-text references made a paper with three tables look like
    # fourteen, and the warning then fired on 13 of 15 papers. Distinct table
    # numbers are the honest measure.
    announced = _table_numbers(text)
    found = set()
    for _, block_text in table_blocks:
        found |= _table_numbers(block_text)
    missing = sorted(announced - found)
    if not announced or not missing:
        return ""
    return (
        f"WARNING: the document announces {len(announced)} numbered table(s) but "
        f"{len(missing)} of them could not be extracted "
        f"(missing: {', '.join('Table ' + m for m in missing)}). Any table not "
        "shown below was not read at all. Do not conclude that a value, "
        "coefficient or diagnostic is unreported merely because it is absent "
        "here; say that the table could not be extracted. Report this as an "
        "extraction limit."
    )


def table_fragmentation_warning(block_text: str) -> str:
    """A warning line when a table's numbers have been split across rows."""
    stems = 0
    for line in block_text.splitlines():
        for cell in line.split("\t"):
            if _ORPHAN_STEM_RE.match(cell.strip()):
                stems += 1
    if stems < 3:
        return ""
    return (
        f"WARNING: this table did not extract cleanly. {stems} cells hold only "
        "the integer part of a number, with the digits after the decimal point "
        "stranded on a neighbouring row, so values such as 13.5 appear as "
        "\"13.\" and \"5\". Use this table for the variables it names and for "
        "nothing else. Do not quote, compare, or compute with any number in it, "
        "and do not raise a data-integrity concern about a value that looks "
        "implausible here: the implausibility is this extraction, not the paper."
    )


# A paper states its total sample once, in prose, and then reports a per-analysis
# N in every table. When a table's N exceeds the stated total, something is
# wrong -- a mislabelled total, a different sampling frame, or a reporting error
# -- and it is worth an author's attention. Three consecutive runs of this
# pipeline over one paper found the discrepancy once: the model had both numbers
# in a single sentence on the third run and did not subtract them. An
# observation that reduces to arithmetic should be computed, not hoped for.
_TOTAL_SAMPLE_UNITS = (
    r"(?:participants?|adults?|patients?|subjects?|respondents?|individuals?|"
    r"children|adolescents?|infants?|volunteers?|cases|athletes?|players?|"
    r"cyclists?|swimmers?|runners?|students?|boys|girls|women|men)"
)
_TOTAL_SAMPLE_NUM = r"(\d{1,3}(?:,\d{3})+|\d+)"
_TOTAL_SAMPLE_RE = re.compile(
    r"(?:total\s+)?(?:sample|cohort|data\s*set|dataset)\s+"
    r"(?:of|comprises|comprised|comprising|consisted\s+of|was)\s+"
    + _TOTAL_SAMPLE_NUM + r"\s*" + _TOTAL_SAMPLE_UNITS
    + r"|total[\s:(]+n\s*[=:]\s*" + _TOTAL_SAMPLE_NUM
    + r"|total\s+of\s+" + _TOTAL_SAMPLE_NUM + r"\s*" + _TOTAL_SAMPLE_UNITS
    + r"|" + _TOTAL_SAMPLE_NUM + r"\s*" + _TOTAL_SAMPLE_UNITS
    + r"\s+(?:were\s+|was\s+)?(?:recruited|enrolled|included|took\s+part|"
      r"completed|participated|volunteered)",
    re.I,
)


# An N label, then the numbers that follow it on the same line. Table rows read
# "N 40,348 41,015 41,015 41,015": one label, one value per column.
_N_LABEL_RE = re.compile(r"(?:^|(?<=[\s\t(]))[Nn]\s*(?:[=:]|\bZ\b)?\s*")
_BIG_INT_RE = re.compile(r"(?:\d{1,3},)+\d{3}$|^\d{4,}$")


_NUMBER_WORDS = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5, "six": 6, "seven": 7,
    "eight": 8, "nine": 9, "ten": 10, "eleven": 11, "twelve": 12,
    "thirteen": 13, "fourteen": 14, "fifteen": 15, "sixteen": 16,
    "seventeen": 17, "eighteen": 18, "nineteen": 19, "twenty": 20,
    "thirty": 30, "forty": 40, "fifty": 50, "sixty": 60, "seventy": 70,
    "eighty": 80, "ninety": 90,
}
_NUM_WORD = "(?:" + "|".join(sorted(_NUMBER_WORDS, key=len, reverse=True)) + "|hundred)"
_NUM_WORD_PHRASE = _NUM_WORD + r"(?:[\s-]+(?:and[\s-]+)?" + _NUM_WORD + r")*"
# Sports-science papers routinely open "Twenty-seven recreational cyclists
# completed...". Written out, the sample size was invisible to the digit
# patterns above, so the consistency check stayed silent on papers it could
# otherwise have checked.
_TOTAL_WORD_RE = re.compile(
    r"\b(" + _NUM_WORD_PHRASE + r")(?:\s+[a-z][a-z-]*){0,2}\s+" + _TOTAL_SAMPLE_UNITS
    # A breakdown often sits between the noun and the verb:
    # "... swimmers (male [n=202]; female [n=161]) participated in the study".
    + r"(?:\s*\([^)]{0,90}\))?"
    + r"\s+(?:were\s+|was\s+)?(?:recruited|enrolled|included|took\s+part|"
      r"completed|participated|volunteered)",
    re.I,
)


def _words_to_int(phrase: str) -> Optional[int]:
    """Read a written-out cardinal up to 999. Returns None if it does not parse."""
    total = 0
    current = 0
    seen = False
    for word in re.split(r"[\s-]+", phrase.lower()):
        if not word or word == "and":
            continue
        if word == "hundred":
            if not current:
                return None
            current *= 100
            seen = True
            continue
        value = _NUMBER_WORDS.get(word)
        if value is None:
            return None
        seen = True
        # "twenty seven" accumulates; "one hundred twelve" already multiplied.
        if value >= 20 and current and current % 100 == 0:
            total += current
            current = value
        else:
            current += value
    return (total + current) if seen else None


def stated_total_sample(text: str) -> Optional[int]:
    """The total sample size the document states in prose, if it states one."""
    candidates = []
    for match in _TOTAL_SAMPLE_RE.finditer(text or ""):
        raw = next((g for g in match.groups() if g), None)
        if not raw:
            continue
        try:
            candidates.append(int(raw.replace(",", "")))
        except ValueError:
            continue
    for match in _TOTAL_WORD_RE.finditer(text or ""):
        value = _words_to_int(match.group(1))
        if value:
            candidates.append(value)
    if not candidates:
        return None
    # A paper may state its total more than once. Where the statements
    # disagree, the largest is the most generous reading and so the one least
    # likely to raise a false discrepancy.
    return max(candidates)


def _reported_sample_sizes(text: str) -> List[int]:
    """Every N value labelled as such, from table rows or prose."""
    found: List[int] = []
    for line in (text or "").splitlines():
        for match in _N_LABEL_RE.finditer(line):
            for token in line[match.end():].split():
                token = token.strip("()[]{}.,;:")
                if not _BIG_INT_RE.match(token):
                    break
                try:
                    found.append(int(token.replace(",", "")))
                except ValueError:
                    break
    return found


def _footnote_marker_artefact(value: int, text: str) -> bool:
    """
    True when a sample size looks like a value with a footnote marker attached.

    A table header reading "N = 177" with a superscript 1 extracts as "1771",
    and on RJSP-2024-1475 that produced a report telling the reviewer to ask
    the authors why their sample of 191 contained 1,771 people -- while the
    same report's extraction limits explained the artefact four headings
    earlier. The signature is specific: drop the last digit and the remainder
    is a number the document also states on its own.
    """
    stem = str(value)[:-1]
    if len(stem) < 2:
        return False
    return re.search(rf"(?<!\d){re.escape(stem)}(?!\d)", text) is not None


def sample_size_warning(text: str, table_blocks: List[Tuple[int, str]]) -> str:
    """A warning when a reported N exceeds the stated total sample."""
    total = stated_total_sample(text)
    if not total:
        return ""
    reported: List[int] = list(_reported_sample_sizes(text))
    for _, block_text in table_blocks or []:
        reported.extend(_reported_sample_sizes(block_text))
    exceeding = sorted({n for n in reported if n > total})
    exceeding = [n for n in exceeding if not _footnote_marker_artefact(n, text)]
    if not exceeding:
        return ""
    shown = ", ".join(f"{n:,}" for n in exceeding[:6])
    if len(exceeding) > 6:
        shown += ", ..."
    return (
        f"WARNING: the document states a total sample of {total:,}, but "
        f"{len(exceeding)} reported sample size(s) exceed it ({shown}; largest "
        f"exceeds the total by {max(exceeding) - total:,}). Either the stated "
        "total refers to a subset, a different sampling frame was used for "
        "those analyses, or there is a reporting error. Check the figures in "
        "the manuscript before raising it: this compares numbers as extracted, "
        "and a table header can glue a footnote marker to a value."
    )


_CONCERN_RE = re.compile(r"^\s*[-*]?\s*\**\s*Concern:?\**\s*:?\s*(.+)$", re.I)
_SEVERITY_RE = re.compile(r"^\s*[-*]?\s*\**\s*Severity:?\**\s*:?\s*(\w+)", re.I)
_EVIDENCE_RE = re.compile(r"^\s*[-*]?\s*\**\s*Evidence:?\**\s*:?\s*(.+)$", re.I)
_CHECK_RE = re.compile(r"^\s*[-*]?\s*\**\s*Check:?\**\s*:?\s*(.+)$", re.I)
# The synthesis usually puts each field on its own line, but not always: a
# thinking-mode run wrote "Concern: X. Severity: Substantive. Evidence: ..."
# on one line, and the action table then carried the whole paragraph as the
# item and defaulted every row to High because no line began with Severity.
# These find the fields wherever on the line they fall.
_INLINE_FIELD_RE = re.compile(
    r"\s*(?:Severity|Evidence|Why it matters|Confidence|Reason)\s*\**\s*:", re.I)
_ANY_SEVERITY_RE = re.compile(r"Severity\s*\**\s*:\s*\**\s*(\w+)", re.I)
_HEADING_RE = re.compile(r"^#{1,3}\s*(.+?)\s*$")
# Verdicts a journal reviewer would not write. The prompt asks for the
# consequence instead; this catches the cases where it does not comply.
_OVERCLAIM_RE = re.compile(
    r"\b(fundamental(?:ly)?\s+(?:error|flawed?)|fatal(?:ly)?\s+flaw(?:ed)?|"
    r"renders?\s+(?:\w+\s+){1,4}(?:invalid|meaningless|worthless|unusable)|"
    r"completely\s+undermines?|entirely\s+invalid|scientifically\s+worthless)\b",
    re.I,
)


def _report_sections(report_text: str) -> Dict[str, List[str]]:
    """The report split into its headed sections, in order."""
    sections: Dict[str, List[str]] = {}
    current = ""
    for line in (report_text or "").splitlines():
        heading = _HEADING_RE.match(line)
        if heading:
            current = heading.group(1).strip().lower()
            sections.setdefault(current, [])
            continue
        if current:
            sections[current].append(line)
    return sections


def _concern_groups(lines: List[str]) -> List[Tuple[int, int]]:
    """Line ranges, one per concern, from the first Concern: line to the next."""
    starts = [i for i, line in enumerate(lines) if _CONCERN_RE.match(line)]
    groups = []
    for index, start in enumerate(starts):
        end = starts[index + 1] if index + 1 < len(starts) else len(lines)
        groups.append((start, end))
    return groups


_EVIDENCE_PREFIX_RE = re.compile(
    r"^\s*(?:the\s+)?(?:file\s+|evidence\s+|chunk\s+)?"
    r"(?:summary|summaries|synopsis|synopses|overview|manifest|notes|extraction|appendix)\s+"
    r"(?:states|notes|says|indicates|confirms|flags|raises|reports|shows|suggests|"
    r"highlights|identifies|lists|mentions)\s*:?\s*",
    re.I,
)


def _content_tokens(text: str) -> set:
    """Content words, folded for plurals so "season" matches "seasons"."""
    tokens = set()
    for word in _strip_punctuation(_normalise_for_match(text)).split():
        if len(word) <= 2:
            continue
        if len(word) > 4 and word.endswith("ies"):
            word = word[:-3] + "y"
        elif len(word) > 4 and word.endswith("es") and not word.endswith("ses"):
            word = word[:-2]
        elif len(word) > 3 and word.endswith("s") and not word.endswith("ss"):
            word = word[:-1]
        tokens.add(word)
    return tokens


def evidence_echoes_concern(group_text: str) -> bool:
    """
    True when a concern's Evidence line merely restates the concern.

    Asked for an Evidence line it cannot source, the model has begun filling
    the slot with the concern's own sentence, prefixed with "The file synopsis
    states:". That is not weak evidence, it is no evidence, and it needs no
    comparison against the manuscript to detect.
    """
    lines = group_text.splitlines()
    concern = ""
    evidence = ""
    for line in lines:
        found = _CONCERN_RE.match(line)
        if found and not concern:
            concern = found.group(1)
            continue
        found = _EVIDENCE_RE.match(line)
        if found and not evidence:
            evidence = _EVIDENCE_PREFIX_RE.sub("", found.group(1).strip())
    if not concern or not evidence:
        return False
    concern_tokens = _content_tokens(concern)
    evidence_tokens = _content_tokens(evidence)
    if len(evidence_tokens) < 5 or not concern_tokens:
        return False
    contained = len(evidence_tokens & concern_tokens) / len(evidence_tokens)
    return contained >= 0.8


def evidence_echo_problems(report_text: str) -> List[str]:
    """Report-level findings for concerns whose evidence restates them."""
    sections = _report_sections(report_text)
    concerns = sections.get("directly supported concerns", [])
    problems = []
    for begin, end in _concern_groups(concerns):
        chunk = concerns[begin:end]
        if not evidence_echoes_concern("\n".join(chunk)):
            continue
        title = _CONCERN_RE.match(chunk[0]).group(1).strip().rstrip("*").strip()
        problems.append(
            f"The Evidence line restates the concern rather than citing the "
            f"manuscript, so the concern is unsupported: \"{title[:100]}\""
        )
    return problems


# Findings rotate between runs. The same paper, the same code and the same
# model, two and a half hours apart, produced six defensible concerns across two
# runs with no overlap at all -- and the second run lost a verified inconsistency
# the first had found. A single pass gives roughly half of what the model can
# see, and which half is arbitrary. Running the synthesis more than once and
# taking the union recovers the rest.
REVIEW_PASSES = max(1, int(os.environ.get("REVIEW_PASSES", "1") or 1))

# Passes repeat the synthesis, and in a thinking review the synthesis is the
# stage that reasons. Three passes there is seven thinking generations before
# any retries: one measured hybrid run took eighty minutes with three, so three
# passes would be hours. Whether extra passes and reasoning add anything to each
# other has never been tested -- they may well be two ways of buying the same
# thing -- so a thinking review takes one pass unless this says otherwise, and
# the expensive combination has to be asked for rather than arrived at.
THINKING_PASSES = max(1, int(os.environ.get("QWEN_THINKING_PASSES", "1") or 1))

_MERGEABLE_SECTIONS = (
    ("directly supported concerns", _CONCERN_RE),
    ("verification prompts", _CHECK_RE),
)


def _item_groups(lines: List[str], marker: re.Pattern) -> List[Tuple[int, int]]:
    starts = [i for i, line in enumerate(lines) if marker.match(line)]
    return [(start, starts[i + 1] if i + 1 < len(starts) else len(lines))
            for i, start in enumerate(starts)]


# Measured on a run where three validated passes produced two concerns that
# were plainly the same finding in different words -- the narrow null interval
# for the Bayes Factors -- and both reached the report and the action table.
# Across that run's items the duplicate pair scored 0.56 and the closest
# genuinely distinct pair 0.39, so the threshold sits between them rather than
# at a round number chosen in advance.
SAME_ITEM_OVERLAP = 0.45


def _same_item(first: str, second: str) -> bool:
    """Two items are the same finding when their wording largely coincides."""
    left, right = _content_tokens(first), _content_tokens(second)
    if not left or not right:
        return False
    overlap = len(left & right) / min(len(left), len(right))
    return overlap >= SAME_ITEM_OVERLAP


def merge_reports(reports: List[str], stats: Optional[dict] = None) -> str:
    """
    Union the concerns and checks of several synthesis passes.

    The first pass is the base report and keeps its synopsis, strengths and
    closing sections: merging prose would produce a document no pass actually
    wrote. Only the two list sections where rotation was observed are unioned,
    and an item is added only when no item already present says the same thing.
    Additions are marked so the reader can see the report is a union.

    Sameness is lexical, so two passes raising one question in wholly different
    words ("included season or year as a covariate" against "account for the
    lockdown and non-lockdown seasons") will still both survive. That remains
    the safe direction to fail: a duplicated question costs the reader a few
    seconds where a swallowed finding costs the review. The threshold is set
    from measured overlaps rather than chosen in advance -- see
    SAME_ITEM_OVERLAP.
    """
    reports = [r for r in reports if r and r.strip()]
    if stats is not None:
        stats.setdefault("passes", len(reports))
        stats.setdefault("added", 0)
    if len(reports) < 2:
        return reports[0] if reports else ""

    base_lines = reports[0].splitlines()
    for name, marker in _MERGEABLE_SECTIONS:
        bounds = None
        for index, line in enumerate(base_lines):
            heading = _HEADING_RE.match(line)
            if not heading:
                continue
            if heading.group(1).strip().lower().startswith(name):
                bounds = [index + 1, len(base_lines)]
            elif bounds and bounds[1] == len(base_lines):
                bounds[1] = index
                break
        if not bounds:
            continue
        start, stop = bounds
        body = base_lines[start:stop]
        existing = ["\n".join(body[b:e]) for b, e in _item_groups(body, marker)]
        additions: List[str] = []
        for other in reports[1:]:
            other_body = _report_sections(other).get(name, [])
            for begin, end in _item_groups(other_body, marker):
                chunk = other_body[begin:end]
                title = marker.match(chunk[0]).group(1)
                if any(_same_item(title, marker.match(e.splitlines()[0]).group(1))
                       for e in existing if marker.match(e.splitlines()[0])):
                    continue
                if any(_same_item(title, marker.match(a.splitlines()[0]).group(1))
                       for a in additions if marker.match(a.splitlines()[0])):
                    continue
                text = "\n".join(line for line in chunk if line.strip())
                additions.append(text + "\n* Found only in a later pass.")
                if stats is not None:
                    stats["added"] = stats.get("added", 0) + 1
        if additions:
            body = [line for line in body if line.strip()]
            body.append("")
            body.extend("\n".join(additions).splitlines())
            body.append("")
            base_lines = base_lines[:start] + body + base_lines[stop:]
    return "\n".join(base_lines)


def count_report_items(report_text: str) -> Tuple[int, int]:
    """Concerns and verification prompts currently in a report."""
    sections = _report_sections(report_text)
    concerns = sections.get("directly supported concerns", [])
    checks = sections.get("verification prompts", [])
    return (len(_concern_groups(concerns)),
            len(_item_groups(checks, _CHECK_RE)))


def report_reliability_banner(report_text: str) -> str:
    """
    A single line at the top when no concern carries a located quotation.

    The per-concern confidence lines are accurate but easy to miss: a reader
    has to find and tally them. When not one concern reaches High, that is a
    fact about the whole report and belongs where it will be seen first.

    What it is a fact ABOUT is the wording, not the findings. On four
    consecutive manuscripts the label demoted a concern that a reviewer
    independently raised, and on RJSP-2019-0453 this banner told the reader to
    treat as mere leads a report whose concerns two referees raised and the
    authors acted on. The banner therefore says what it measures and stops
    instructing the reader what to conclude from it.
    """
    levels = re.findall(r"^\s*[-*]?\s*\**\s*Confidence:?\**\s*:?\s*(\w+)",
                        report_text or "", re.M | re.I)
    if not levels or any(level.lower() == "high" for level in levels):
        return ""
    return (
        "> **No concern in this report carries a quotation that could be "
        "located in the extracted manuscript.** That is a statement about how "
        "the concerns are worded and evidenced, not about whether they are "
        "right: a correct concern often arrives paraphrased. Check each one "
        "against the paper before accepting or dismissing it.\n\n"
    )


def concern_confidence(group_text: str, source_text: str) -> Tuple[str, str]:
    """
    Confidence in one concern, computed rather than asked for.

    A model's own confidence statement is a token prediction, not a
    measurement: the run that inverted this paper's medication finding
    asserted it without hedging, and three runs cited the pipeline's own
    summary while sounding certain. What the reader actually wants to know --
    does this rest on the manuscript or on something softer -- is already
    decidable from the checks the pipeline performs.
    """
    if evidence_echoes_concern(group_text):
        return ("Low", "the evidence line restates the concern rather than "
                       "citing the manuscript")
    if _SELF_CITATION_RE.search(group_text):
        return ("Low", "the evidence cites this pipeline's own summary rather "
                       "than the manuscript")
    haystack = _normalise_for_match(source_text)
    haystack_punct = _strip_punctuation(haystack)
    haystack_nospace = re.sub(r"\s+", "", haystack_punct)
    quotes = [q.strip() for q in _quoted_spans(group_text)
              if len(q.strip().split()) >= 4]
    # Every quotation must hold, not merely one of them. A concern pairing a
    # genuine quotation with an invented one previously read High on the
    # strength of the first: this paper's effect-size concern quoted "just over
    # half a point per bout", which is in the abstract, alongside a fabricated
    # second quotation, and was reported as High.
    unverified = [q for q in quotes
                  if _missing_fragment(q, haystack, haystack_punct,
                                       haystack_nospace) is not None]
    if unverified and len(unverified) < len(quotes):
        return ("Low", "one quotation was located in the manuscript and "
                       "another could not be")
    if unverified:
        return ("Low", "the quoted evidence could not be located in the "
                       "extracted manuscript")
    if quotes:
        return ("High", "every quotation was located in the manuscript")
    source_numbers = set(
        re.findall(r"\d+\.\d+", _normalise_numeric_artefacts(source_text))
    )
    cited = re.findall(r"(?<![\d.])\d+\.\d+(?!\d)", group_text)
    if cited and any(n in source_numbers for n in cited):
        return ("High", "the cited values were located in the extracted tables")
    return ("Moderate", "the concern is an inference; no verbatim quotation or "
                        "table value supports it directly")


def annotate_concern_confidence(report_text: str, source_text: str) -> str:
    """Append a computed Confidence line to each directly supported concern."""
    lines = (report_text or "").splitlines()
    bounds = None
    for index, line in enumerate(lines):
        heading = _HEADING_RE.match(line)
        if not heading:
            continue
        name = heading.group(1).strip().lower()
        if name.startswith("directly supported concerns"):
            bounds = [index + 1, len(lines)]
        elif bounds and bounds[1] == len(lines):
            bounds[1] = index
            break
    if not bounds:
        return report_text
    start, stop = bounds
    body = lines[start:stop]
    if any("Confidence:" in line for line in body):
        return report_text
    out: List[str] = []
    groups = _concern_groups(body)
    if not groups:
        return report_text
    covered = groups[0][0]
    out.extend(body[:covered])
    for begin, end in groups:
        chunk = body[begin:end]
        level, reason = concern_confidence("\n".join(chunk), source_text)
        while chunk and not chunk[-1].strip():
            chunk.pop()
        chunk.append(f"* Confidence: {level} — {reason}.")
        chunk.append("")
        out.extend(chunk)
    return "\n".join(lines[:start] + out + lines[stop:])


# Concerns were once ordered by a severity label the synthesis wrote for
# itself. Across the reports we measured it used one of its three permitted
# values for 82% of concerns, and on the single occasion it did discriminate it
# placed its best-supported concern below its weakest. It was the only
# self-assessed grade left in a report otherwise built on computed checks, so it
# is gone. What remains is where each item's evidence actually stands, which the
# citation check establishes mechanically.
# "Verified" and "Unverified" read as verdicts on whether a concern is right.
# They record only whether a quotation resolved, and the two came apart on four
# consecutive manuscripts. The labels now name what was measured.
_EVIDENCE_RANK = {
    "high": ("Quoted", 0),
    "moderate": ("Reasoned", 1),
    "low": ("Unquoted", 2),
}
_CONFIDENCE_RE = re.compile(r"^\s*[-*]?\s*\**\s*Confidence:?\**\s*:?\s*(\w+)", re.I)


def format_action_list(report_text: str) -> str:
    """
    A priority table assembled from the report's own severity labels.

    Rendered from the concerns already written rather than generated again:
    an earlier template asked for the same material twice and produced two
    headings over one list of items. The wording here is the concern's own, so
    the table cannot drift from the body it summarises.
    """
    sections = _report_sections(report_text)
    rows: List[Tuple[int, str, str]] = []

    concerns = sections.get("directly supported concerns", [])
    for begin, end in _concern_groups(concerns):
        chunk = concerns[begin:end]
        title = _CONCERN_RE.match(chunk[0]).group(1).strip().rstrip("*").strip()
        # Cut at the next field when the model ran them together on one line.
        inline = _INLINE_FIELD_RE.search(title)
        if inline:
            title = title[:inline.start()].strip().rstrip(".").strip()
        confidence = ""
        for line in chunk:
            found = _CONFIDENCE_RE.match(line)
            if found:
                confidence = found.group(1).lower()
                break
        # An unparsed confidence line falls to the most cautious label the
        # preamble defines. The old default was "Unverified", a word the
        # renamed vocabulary no longer explains to the reader.
        label, order = _EVIDENCE_RANK.get(confidence, ("Unquoted", 2))
        rows.append((order, label, title))

    for line in sections.get("verification prompts", []):
        found = _CHECK_RE.match(line)
        if found:
            check = found.group(1).strip().rstrip("*").strip()
            inline = _INLINE_FIELD_RE.search(check)
            if inline:
                check = check[:inline.start()].strip().rstrip(".").strip()
            # "Check: whether X" reads badly as a table row on its own.
            reason = re.match(r"\s*Reason\s*\**\s*:", check, re.I)
            if not reason and check[:1].islower():
                check = check[:1].upper() + check[1:]
            rows.append((3, "Question", check))

    if not rows:
        return ""
    rows.sort(key=lambda row: row[0])
    out = ["\n\n# Items by evidence", "",
           "Ordered by how each item is worded and evidenced, which the "
           "citation check establishes mechanically. Quoted means every "
           "quotation in it was located in the manuscript; Reasoned means it "
           "argues from the evidence rather than quoting it; Unquoted means a "
           "quotation could not be located or the evidence cited this "
           "pipeline's own summary; Question marks a check to settle rather "
           "than an established fault. **This orders provenance, not "
           "correctness.** An Unquoted item may be right and a Quoted one "
           "wrong; on four manuscripts in the evaluation set an Unquoted "
           "concern was one a human reviewer raised independently. The wording "
           "is each item's own, and how much each matters is your judgement.", "",
           "| Evidence | Item |", "| --- | --- |"]
    for _, label, title in rows:
        title = re.sub(r"\s+", " ", title).replace("|", "/")
        if len(title) > 150:
            title = title[:147].rstrip() + "..."
        out.append(f"| {label} | {title} |")
    return "\n".join(out) + "\n"


def overclaim_problems(report_text: str) -> List[str]:
    """Verdict phrasing the register rule asks the model to avoid."""
    problems = []
    for match in _OVERCLAIM_RE.finditer(report_text or ""):
        problems.append(
            f"Verdict phrasing rather than a stated consequence: "
            f"\"{match.group(0)}\". A journal review would say what follows "
            f"from the problem instead."
        )
    return list(dict.fromkeys(problems))


# Word writes these into a PDF when a cross-reference cannot be resolved. An
# unpublished manuscript submitted with them is telling every reader that its
# table and figure references are broken -- 21 of them in one submission this
# pipeline reviewed, and the review did not mention it once. It is the first
# thing a human reviewer would raise and it needs no judgement to find.
_SUBMISSION_ARTEFACT_RE = re.compile(
    r"Error!\s*(?:Reference source not found|Bookmark not defined|"
    r"No text of specified style|Unknown switch argument)"
    r"|\?\?\?\s*(?:ref|cite)|\\ref\{|\\cite\{|\[CITATION\]|\bTODO\b|\bTBD\b"
    r"|\bXXX\b|\blorem ipsum\b",
    re.I,
)


def submission_integrity_warning(text: str) -> str:
    """
    A warning when the document carries unresolved references or placeholders.

    The earlier version matched a bare "XXX" and so reported the review
    blinding: on RJSP-2021-1229 it flagged a redacted grant number and contact
    address as seventeen unresolved placeholders and instructed the reviewer to
    raise them as an editorial concern, then the report described a blinded
    submission as an unfinished draft. Redactions in funding, ethics, contact
    and self-citation slots are what a journal asks authors for. The detection
    now lives in manuscript_checks.placeholder_warning(), which reports
    authoring artefacts, ignores the blinding, and does not instruct.
    """
    findings = mchk.placeholder_warning(text or "")
    if not findings:
        return ""
    return "WARNING: " + " ".join(f"{f.label} {f.detail}".strip() for f in findings)


# Measured over the sixteen-paper reference set: the totals ran 0 to 10 display
# items with a median of 4, and the highest single counts were 6 tables and 7
# figures. The thresholds sit clear of those maxima, so an ordinary article
# never trips them. A manuscript declaring 17 tables and 24 figures does.
MAX_USUAL_TABLES = 8
MAX_USUAL_FIGURES = 10
MAX_USUAL_DISPLAY_ITEMS = 12

_FIGURE_CAPTION_RE = re.compile(r"(?mi)^[ \t]*(?:\d+\s+)?fig(?:ure)?\.?\s*(\d+)\s*[.:]")


def display_item_warning(text: str) -> str:
    """A warning when a manuscript declares an unusual number of tables and figures."""
    tables = _table_numbers(text)
    figures = set(_FIGURE_CAPTION_RE.findall(text or ""))
    total = len(tables) + len(figures)
    if (len(tables) <= MAX_USUAL_TABLES and len(figures) <= MAX_USUAL_FIGURES
            and total <= MAX_USUAL_DISPLAY_ITEMS):
        return ""
    return (
        f"WARNING: the document declares {len(tables)} numbered table(s) and "
        f"{len(figures)} numbered figure(s), {total} display items in total. "
        "That is well above what a research article of this kind normally "
        "carries, and journals commonly cap them. Consider whether some belong "
        "in supplementary material, and note that a reader cannot hold this "
        "many in view at once. Raise it as an editorial concern."
    )


def format_consistency_check(text: str,
                             table_blocks: List[Tuple[int, str]]) -> str:
    """A report section for deterministic findings, appended after synthesis."""
    findings = [w for w in (sample_size_warning(text, table_blocks),
                            submission_integrity_warning(text),
                            display_item_warning(text)) if w]

    # Checks measured over twenty documents before integration: nothing fired
    # on any of the sixteen published papers, and every finding on the four
    # manuscripts was true. See section 12 of the evaluation record.
    blocks = [block for _, block in (table_blocks or [])]
    extra = mchk.run_all(text or "", blocks)
    findings.extend(f"{f.label} {f.detail}".strip()
                    + (f" (page {f.page})" if f.page is not None else "")
                    for f in extra)

    if not findings:
        return ""
    lines = ["\n\n# Data consistency check", "",
             "Computed directly from the extracted text, independently of the "
             "model's reading. These are arithmetic and pattern checks, not "
             "judgements: confirm each against the manuscript before raising "
             "it.", ""]
    for warning in findings:
        lines.append(f"* {warning[len('WARNING: '):] if warning.startswith('WARNING: ') else warning}")
    return "\n".join(lines) + "\n"


def format_expected_elements(text: str,
                             design_class: Optional["de.DesignClass"] = None) -> str:
    """
    The elements a study of this design is normally expected to report, and
    which of them could not be found.

    This is the one place the pipeline raises an omission. Everything else it
    does requires a quotation, and an absence has none, which is why concerns
    of this kind never reached a report. An absence check also fails in the
    opposite direction from a quotation check -- a missed synonym accuses
    rather than misses -- so every line prints the terms that were searched.
    """
    if design_class is None:
        return ""
    findings = de.check_expected_elements(text or "", design_class)
    section = de.format_expected_elements_section(findings)
    return ("\n\n" + section) if section else ""


MAX_PROMPT_TABLE_ROWS = 45
MAX_PROMPT_TABLE_CHARS = 9000


_SELF_CITATION_RE = re.compile(
    r"\b(?:the\s+)?(?:file\s+|evidence\s+|chunk\s+)?"
    r"(?:summary|summaries|synopsis|synopses|overview|manifest|notes|"
    r"extraction|appendix)\s+"
    r"(?:states|notes|says|indicates|confirms|flags|raises|reports|shows|suggests|highlights)\b",
    re.I,
)

_QUOTE_CHAR_RE = re.compile(r"[\u201c\u201d\"]")


def _quoted_spans(text: str) -> List[str]:
    """
    Quoted passages, found by pairing quote characters in order.

    A regex of the form "([^"]+)" pairs the CLOSING quote of one passage with
    the OPENING quote of the next whenever the first is too short to match,
    inventing a span out of the words between two real quotations. Pairing
    positionally avoids that.
    """
    positions = [m.start() for m in _QUOTE_CHAR_RE.finditer(text)]
    spans = []
    for i in range(0, len(positions) - 1, 2):
        content = text[positions[i] + 1 : positions[i + 1]]
        if "\n" in content or not (12 <= len(content) <= 400):
            continue
        spans.append(content)
    return spans


SHORT_QUOTE_MIN = 4


def _short_quoted_spans(text: str) -> List[str]:
    """
    Quoted passages too short for _quoted_spans, kept only when numeric.

    The 12-character floor there exists so that a stray pair of quote marks
    around one word is not treated as a citation. That is the right default for
    prose and the wrong one for figures.
    """
    positions = [m.start() for m in _QUOTE_CHAR_RE.finditer(text)]
    spans = []
    for i in range(0, len(positions) - 1, 2):
        content = text[positions[i] + 1: positions[i + 1]]
        if "\n" in content or not (SHORT_QUOTE_MIN <= len(content) < 12):
            continue
        if re.search(r"\d", content):
            spans.append(content)
    return spans


def _normalise_numeric_artefacts(text: str) -> str:
    """
    Repair decimal separators that PDF extraction mangles.

    Several publishers set the decimal point in a font whose glyph extracts as
    "$" or a middle dot, so 0.795 arrives as "0$795" and, where the digits are
    split, as "0 $795". The citation check compared the report's numbers
    against the raw extraction and therefore reported every correct decimal in
    such a paper as absent: one Elsevier paper contained 314 of them. Fold the
    separators back before any comparison. Only a separator between two digits
    is touched, so an index name such as WHT$5R is left alone.
    """
    return re.sub(r"(?<=\d)[ \t]*[$\u00b7\u2219\u00b7][ \t]*(?=\d)", ".", text)


def _normalise_for_match(text: str) -> str:
    """Fold the differences that stop a true quotation matching its source."""
    text = _normalise_numeric_artefacts(text)
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u2019", "'").replace("\u2018", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    for dash in ("\u2013", "\u2014", "\u2212"):
        text = text.replace(dash, "-")
    text = re.sub(r"\s+", " ", text)
    return text.strip().lower()


def _strip_punctuation(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^0-9a-z ]+", " ", text)).strip()


_ELLIPSIS_RE = re.compile(r"\s*(?:\.\s?\.\s?\.|\u2026)\s*")


def _fragment_located(fragment: str, hay: str, hay_punct: str,
                      hay_nospace: str) -> bool:
    """One contiguous fragment, tried at three levels of tolerance."""
    needle = _normalise_for_match(fragment)
    if not needle:
        return True
    if needle in hay:
        return True
    stripped = _strip_punctuation(needle)
    if stripped and stripped in hay_punct:
        return True
    return bool(stripped) and re.sub(r"\s+", "", stripped) in hay_nospace


def _missing_fragment(quoted: str, hay: str, hay_punct: str,
                      hay_nospace: str) -> Optional[str]:
    """
    The first part of a quotation that is not in the source, or None.

    A quotation may legitimately elide with "...". Requiring the whole span to
    be contiguous reported such a quotation as absent even when both halves
    were faithful, and -- worse -- said nothing about which half was invented.
    Each fragment is checked separately, and the one that fails is named.
    """
    fragments = [f.strip() for f in _ELLIPSIS_RE.split(quoted) if f.strip()]
    for fragment in fragments:
        if len(fragment.split()) < 3 and len(fragments) > 1:
            continue  # too short to locate meaningfully on its own
        if not _fragment_located(fragment, hay, hay_punct, hay_nospace):
            return fragment
    return None


def verify_report_citations(report_text: str, source_text: str) -> List[str]:
    """
    Check the report's citations mechanically.

    Two rules have proved unreliable when left to the prompt alone: quote the
    manuscript rather than this pipeline's own summary, and quote it exactly.
    One run obeyed both and the next did not, so they are checked here instead
    of hoped for. Nothing is rewritten; the findings are appended to the report
    so a reader can see which citations stand up.
    """
    problems: List[str] = []
    haystack = _normalise_for_match(source_text)
    # Precomputed once: these were rebuilt per quotation, which is O(n*m) over
    # a whole manuscript. Behaviour is unchanged.
    haystack_punct = _strip_punctuation(haystack)
    haystack_nospace = re.sub(r"\s+", "", haystack_punct)

    for match in _SELF_CITATION_RE.finditer(report_text):
        line = report_text[: match.start()].split("\n")[-1] + match.group(0)
        problems.append(
            f"Cites this pipeline's own summary rather than the manuscript: "
            f"...{line.strip()[-110:]}"
        )

    # A numeric quotation is checked however short it is. "n = 116 173" is
    # eleven characters and three words, so it passed under both floors on
    # RPAN-2026-0184 -- and it was fabricated: the marginal line number 116 had
    # been stripped correctly, and the manuscript reads n = 173. Numbers are
    # the quotations most worth checking and the easiest to invent.
    spans = list(_quoted_spans(report_text))
    spans += [s for s in _short_quoted_spans(report_text) if s not in spans]
    for quoted in spans:
        quoted = quoted.strip()
        if len(quoted.split()) < 4 and not re.search(r"\d", quoted):
            continue
        # Tolerances live in _missing_fragment: punctuation, spacing (extraction
        # splits "strati fied" and fuses "i tw a sW C"), and elision with "...".
        missing = _missing_fragment(quoted, haystack, haystack_punct,
                                    haystack_nospace)
        if missing is None:
            continue
        if missing.strip() != quoted.strip():
            problems.append(
                f"Quotation only partly found: \"{quoted[:80]}\" -- this part "
                f"is not in the manuscript: \"{missing[:80]}\""
            )
        else:
            # A quotation that cannot be located is not the same as evidence
            # that is not there. On three consecutive manuscripts the report's
            # most substantive concern was demoted to Unverified because the
            # model had joined a row label to a value from another cell, while
            # every number it cited was present. Say which.
            note = mchk.numeric_fallback(quoted, source_text)
            if note:
                problems.append(
                    f"Quotation not found in the manuscript: \"{quoted[:110]}\" "
                    f"-- {note}"
                )
            else:
                problems.append(
                    f"Quotation not found in the manuscript: \"{quoted[:110]}\""
                )

    # Numbers are quoted as readily as words and fabricated more easily: a
    # value invented from a badly extracted cell reads as authoritative and is
    # the kind of thing that gets sent to an author.
    # \d*\. so that a manuscript writing ".001" is matched by a report writing
    # "0.001": on the basketball manuscript the model reported the gamma prior
    # faithfully and the check called the number absent.
    source_digits = re.findall(r"\d*\.\d+", _normalise_numeric_artefacts(source_text))
    source_numbers = set(source_digits)
    # A report that rounds an extracted value is not fabricating it. Saying
    # "approximately 0.13" of a table's 0.134, or calling 0.221 - 0.211 a
    # difference of 0.01, was reported as three absent numbers on the cycling
    # paper; every one of them was a faithful rounding. Accept a report number
    # when some source value rounds to it at the precision the report used.
    source_values = []
    for token in source_digits:
        try:
            source_values.append(float(token))
        except ValueError:
            continue
    # A verification prompt asks a question and offers examples; the numbers in
    # "e.g., 0.95, 0.99, or a ROPE" are hypothetical thresholds, not citations,
    # and were reported as absent on Eustace et al. 2025.
    citable = re.split(r"^#+\s*Verification prompts\s*$", report_text, flags=re.M | re.I)[0]
    for number in dict.fromkeys(re.findall(r"(?<![\d.])\d+\.\d+(?!\d)", citable)):
        if number in source_numbers:
            continue
        places = len(number.split(".")[1])
        if any(f"{value:.{places}f}" == number for value in source_values):
            continue
        problems.append(
            f"Number {number} does not appear in the manuscript. If it is a "
            f"value you calculated, say so; if it was read from a table, "
            f"re-check it against the extracted text."
        )

    # Collapse duplicates while keeping order.
    seen = set()
    unique = []
    for p in problems:
        if p not in seen:
            seen.add(p)
            unique.append(p)
    return unique


def mark_unverified_quotations(report_text: str, source_text: str) -> str:
    """
    Remove quotation marks from spans that are not in the manuscript.

    The false claim in these reports is one of provenance, not content: the
    wording is the model's, but the marks assert it is the paper's. Removing
    them leaves an ordinary paraphrase, which is a legitimate thing for a
    reviewer to write, and adds nothing that was not already there.

    Substituting the nearest verbatim sentence was the alternative and is
    worse: it introduces a claim rather than withdrawing one, and a wrongly
    chosen sentence passes the citation check precisely because it did come
    from the source. A visible problem would become an invisible one.

    No corroboration gate is applied. Content-word overlap against the source
    was measured on a report where nine of nine quotations failed: an invented
    control scored 0.12 and a verbatim span 1.00, but the nine real spans ran
    from 0.17 to 0.57 with no gap between paraphrase and invention. The tag is
    therefore neutral about which it is, and the Citation check keeps the full
    record either way.
    """
    if not report_text:
        return report_text
    haystack = _normalise_for_match(source_text)
    haystack_punct = _strip_punctuation(haystack)
    haystack_nospace = re.sub(r"\s+", "", haystack_punct)

    # Line by line, so an Evidence line carrying two unverified spans is tagged
    # once at its end rather than interrupted mid-sentence twice. Quoted spans
    # never cross a line break, so pairing within a line is equivalent to
    # pairing across the report and is not thrown off by a stray mark
    # elsewhere.
    lines = report_text.split("\n")
    changed = False
    for number, line in enumerate(lines):
        positions = [m.start() for m in _QUOTE_CHAR_RE.finditer(line)]
        drop = []
        for index in range(0, len(positions) - 1, 2):
            open_at, close_at = positions[index], positions[index + 1]
            content = line[open_at + 1:close_at].strip()
            if not (12 <= len(content) <= 400) or len(content.split()) < 4:
                continue
            if _missing_fragment(content, haystack, haystack_punct,
                                 haystack_nospace) is None:
                continue
            drop.extend((open_at, close_at))
        if not drop:
            continue
        kept = [char for position, char in enumerate(line)
                if position not in set(drop)]
        lines[number] = ("".join(kept).rstrip()
                         + " *[not verbatim - see Citation check]*")
        changed = True
    return "\n".join(lines) if changed else report_text


# ---------------------------------------------------------------------------
# Asking questions of a reviewed manuscript
# ---------------------------------------------------------------------------
# The recurring failure in the evaluation record is that a reviewer cannot
# cheaply check a concern, so the evidence label ends up doing the checking for
# them and does it badly: on four consecutive manuscripts it demoted a concern a
# human reviewer independently raised. A lookup that answers "where does the
# paper say this?" turns each of those into a two-second check.
#
# The danger is exactly the one this project exists to measure. An answer has
# none of the machinery a report has, so a naive chat produces confident,
# unverifiable claims about a manuscript. Three rules follow, and they are
# enforced here rather than requested in the prompt: the model sees only
# passages drawn from this manuscript, every answer is put through the same
# citation check as a report, and the check's result travels with the answer.

QA_MAX_TOKENS = int(os.environ.get("QWEN_QA_MAX_TOKENS", "700"))
QA_BLOCK_CHARS = 1100
QA_MIN_BUDGET = 6000
# Room for the question and the chat scaffolding around the passages.
QA_SLACK_CHARS = 2000
QA_REPORT_CHARS = 12000

# Retrieval is a fallback, not the design. A manuscript in this field extracts
# to around 20,000 tokens and the context window is 32,768, so the whole paper
# usually fits in one prompt -- and the best retrieval is no retrieval, because
# a passage that was never selected cannot be quoted. Only a long document is
# narrowed, and the prompt then says which case it is: an absence in a complete
# document is evidence, an absence in a selection is not.


def qa_passage_budget(report_chars: int = 0) -> int:
    """Characters of manuscript one question may carry."""
    override = os.environ.get("QWEN_QA_PASSAGE_CHARS")
    if override:
        return int(override)
    budget = prompt_char_budget(reserve_tokens=QA_MAX_TOKENS)
    budget -= len(QA_SYSTEM_PARTIAL) + QA_SLACK_CHARS
    budget -= min(report_chars, QA_REPORT_CHARS)
    return max(QA_MIN_BUDGET, budget)

_QA_STOPWORDS = {
    "the", "a", "an", "of", "for", "to", "in", "and", "or", "was", "were",
    "is", "are", "be", "been", "with", "that", "this", "these", "those", "as",
    "by", "on", "at", "from", "it", "we", "our", "their", "they", "what",
    "which", "where", "when", "how", "why", "did", "does", "do", "any", "there",
    "about", "paper", "study", "authors", "manuscript", "report", "you",
}


def _qa_terms(question: str) -> set:
    words = re.findall(r"[A-Za-z][A-Za-z0-9\-]+", (question or "").lower())
    return {w for w in words if w not in _QA_STOPWORDS and len(w) > 2}


def _qa_blocks(text: str) -> List[Tuple[Optional[int], str]]:
    """Split the extracted manuscript into passages, keeping page numbers."""
    blocks: List[Tuple[Optional[int], str]] = []
    page: Optional[int] = None
    buffer: List[str] = []
    size = 0

    def flush():
        nonlocal buffer, size
        if buffer:
            blocks.append((page, "\n".join(buffer).strip()))
        buffer, size = [], 0

    for line in (text or "").splitlines():
        marker = re.match(r"\s*\[Page (\d+)\]", line)
        if marker:
            flush()
            page = int(marker.group(1))
            continue
        buffer.append(line)
        size += len(line) + 1
        if size >= QA_BLOCK_CHARS and not line.strip():
            flush()
    flush()
    return [(p, b) for p, b in blocks if b]


BM25_K1 = 1.5
BM25_B = 0.75


def _bm25_scores(terms: set, blocks: List[Tuple[Optional[int], str]]) -> List[float]:
    """
    BM25 over the blocks of one manuscript.

    Counting how many question terms a block contains, which is what this did
    first, weighs "cluster" and "S_Dbw" the same in a paper about clustering:
    the common word contributes nothing but dominates the count, and a long
    block wins by containing more words. Inverse document frequency and length
    normalisation are the standard fix and need no model.
    """
    tokenised = [re.findall(r"[a-z0-9][a-z0-9\-]+", body.lower()) for _p, body in blocks]
    lengths = [len(tokens) or 1 for tokens in tokenised]
    average = sum(lengths) / len(lengths)
    total = len(blocks)

    scores = [0.0] * total
    for term in terms:
        containing = sum(1 for tokens in tokenised if term in tokens)
        if not containing:
            continue
        idf = math.log(1 + (total - containing + 0.5) / (containing + 0.5))
        for index, tokens in enumerate(tokenised):
            frequency = tokens.count(term)
            if not frequency:
                continue
            norm = 1 - BM25_B + BM25_B * lengths[index] / average
            scores[index] += idf * frequency * (BM25_K1 + 1) / (frequency + BM25_K1 * norm)
    return scores


def select_passages(question: str, text: str,
                    budget: Optional[int] = None) -> List[Tuple[Optional[int], str]]:
    """
    The passages of the manuscript most likely to answer a question, or all of
    them when the manuscript fits.
    """
    if budget is None:
        budget = qa_passage_budget()
    terms = _qa_terms(question)
    blocks = _qa_blocks(text)
    if not blocks:
        return []
    if len(text) <= budget:
        return blocks
    if not terms:
        return blocks[:3]

    scored = [(score, index, page, body)
              for (score, (index, (page, body)))
              in zip(_bm25_scores(terms, blocks), enumerate(blocks))
              if score > 0]
    scored.sort(key=lambda row: -row[0])

    chosen, used = [], 0
    for _score, index, page, body in scored:
        if used + len(body) > budget:
            continue
        chosen.append((index, page, body))
        used += len(body)
        if used >= budget:
            break
    if not chosen:
        # Terms that match nothing is not the same as an empty manuscript.
        # Returning the opening passages lets the model answer "not in the
        # passages I can see", which is the honest answer and the one the
        # system prompt asks for; returning nothing would make the caller
        # invent an error message instead.
        used = 0
        for index, (page, body) in enumerate(blocks):
            if used + len(body) > budget:
                break
            chosen.append((index, page, body))
            used += len(body)
    chosen.sort(key=lambda row: row[0])
    return [(page, body) for _index, page, body in chosen]


QA_SYSTEM_HEAD = """You answer questions about one manuscript, for a reviewer who is \
checking it. You are a lookup instrument, not a second opinion.

Rules, all of them absolute:
- Answer only from the passages below. They are the manuscript.
- Quote the manuscript for every factual claim, reproducing the wording exactly.
- If the passages do not contain the answer, say so plainly and say what section \
or wording would settle it. Never fill a gap by inference and never guess at \
what a paper of this kind usually says.
{absence}
- If asked about the review rather than the paper, quote the review and say \
that is what you are quoting.
- Do not judge the paper, recommend a decision, or suggest improvements unless \
asked, and then only from what the passages support.
- Be brief. A located quotation with its page is a complete answer."""

_ABSENCE_PARTIAL = ("- The passages below are a selection from a long document. "
                    "The absence of something from them is not evidence of "
                    "absence in the manuscript: say \"not in the passages I can "
                    "see\", not \"the manuscript does not report it\".")
_ABSENCE_COMPLETE = ("- The passages below are the whole manuscript, in order. "
                     "If something is not in them it is not in the paper, and "
                     "you may say so - but say it of what the text contains, "
                     "never of what a figure might show, since figures are not "
                     "in the text.")

QA_SYSTEM_PARTIAL = QA_SYSTEM_HEAD.format(absence=_ABSENCE_PARTIAL)
QA_SYSTEM_COMPLETE = QA_SYSTEM_HEAD.format(absence=_ABSENCE_COMPLETE)
QA_SYSTEM = QA_SYSTEM_PARTIAL          # kept for callers that want the default


def answer_manuscript_question(model, tokenizer, question: str, text: str,
                               report_text: str = "",
                               history: Optional[List[dict]] = None,
                               ) -> Tuple[str, List[str]]:
    """
    Answer one question about a reviewed manuscript, and check the answer.

    Returns the answer and the citation-check problems found in it. The caller
    is expected to show both: an answer whose quotations did not resolve is the
    one a reviewer most needs to know about.
    """
    budget = qa_passage_budget(len(report_text or ""))
    passages = select_passages(question, text, budget)
    complete = len(text) <= budget
    if not passages:
        return ("No manuscript text is available for this review, so there is "
                "nothing to look in.", [])
    del history  # reserved: each question is answered from the manuscript alone

    rendered = []
    for page, body in passages:
        where = f"[Page {page}]" if page is not None else "[Page unknown]"
        rendered.append(f"{where}\n{body}")
    passage_block = "\n\n---\n\n".join(rendered)

    report_block = ""
    if report_text:
        report_block = (
            "\n\nThe review this pipeline produced, for questions about the "
            "review itself. It is not the manuscript, and quoting it is not "
            "evidence about the paper:\n\"\"\"\n"
            + report_text[:12000] + "\n\"\"\"\n"
        )

    user_text = f"""{QA_SYSTEM_COMPLETE if complete else QA_SYSTEM_PARTIAL}

Passages from the manuscript:
\"\"\"
{passage_block}
\"\"\"{report_block}

Question: {question}"""

    prompt = apply_chat_template_compat(tokenizer, user_text)
    sampler = make_default_sampler()
    out = generate(model, tokenizer, prompt=prompt, max_tokens=QA_MAX_TOKENS,
                   sampler=sampler, verbose=False)
    answer = clean_model_output(out)

    problems = verify_report_citations(answer, text)
    if report_text:
        # A quotation taken from the review is not a failed manuscript
        # quotation. Relabel rather than drop it: the reviewer still needs to
        # know the answer is quoting this pipeline and not the paper.
        kept = []
        for problem in problems:
            quoted = re.search(r'"([^"]{8,})"', problem)
            if quoted and quoted.group(1)[:80] in report_text:
                kept.append(
                    "Quoted from the review rather than the manuscript: "
                    f"\"{quoted.group(1)[:80]}\""
                )
            else:
                kept.append(problem)
        problems = kept
    return answer, problems


def format_answer_check(answer: str, problems: List[str]) -> str:
    """A short verification footer for one answer."""
    if problems:
        lines = ["", "_Check:_"]
        lines += [f"- {p}" for p in problems]
        return "\n".join(lines)
    if not _quoted_spans(answer or ""):
        # The check only inspects quotations of 12 characters or more, so a
        # two-word quotation is not "nothing quoted"; it is nothing checkable.
        if _QUOTE_CHAR_RE.search(answer or ""):
            return ("\n_Check: the quotations in this answer are too short for "
                    "the citation check to verify._")
        return ("\n_Check: this answer quotes nothing, so nothing in it was "
                "verified against the extracted text._")
    return ("\n_Check: every quotation was located in the extracted text, "
            "which is not the same as the manuscript. Check any numeric value "
            "against the PDF._")


# The last heading the synthesis is asked to write. A report that does not
# reach it stopped early, whatever else it contains.
FINAL_REPORT_HEADING = "Overall confidence"


def report_looks_truncated(report_text: str) -> str:
    """
    Say why a synthesis looks unfinished, or return an empty string.

    A thinking-mode synthesis spent 9,800 of its 10,000 tokens reasoning and
    stopped four words into a sentence. What arrived was two headings and half
    a strength, formatted perfectly, with no concerns section at all -- and
    every check downstream passed it, because half a report is still valid
    markdown. Truncation has to be detected structurally as well as from the
    server's own finish reason, since not every backend reports one.
    """
    if not report_text.strip():
        return "the synthesis returned nothing"
    if not re.search(rf"^#+\s*{FINAL_REPORT_HEADING}", report_text, re.M | re.I):
        return (f"it never reached its final '{FINAL_REPORT_HEADING}' heading, "
                f"so the synthesis stopped partway through")
    tail = [line for line in report_text.strip().splitlines() if line.strip()]
    last = tail[-1].strip() if tail else ""
    if last and not last.startswith(("#", "|", ">")) and last[-1] not in ".!?:)\"'":
        return f"its last line breaks off mid-sentence: {last[-60:]!r}"
    return ""


def format_citation_check(problems: List[str], report_text: str = "") -> str:
    """
    The citation check, or an honest account of why there was nothing to check.

    A clean check used to read the same whether every quotation had been
    verified or the report contained none at all. One thinking-mode review
    quoted nothing whatsoever -- every line of evidence was a paraphrase -- and
    the report congratulated itself on having no unverifiable quotations. That
    is the reverse of the truth: a report with no quotations has nothing
    anchored to the manuscript at all.
    """
    if not problems:
        quoted = len(_quoted_spans(report_text)) if report_text else None
        if quoted == 0:
            return (
                "\n\n# Citation check\n\n"
                "* **This report quotes the manuscript nowhere.** There were no "
                "quotations to check, which is not the same as every quotation "
                "being verified: no statement in this report is anchored to the "
                "paper's own words.\n"
            )
        return (
            "\n\n# Citation check\n\n"
            "* Every quotation in this report was located in the extracted "
            "manuscript text, and no citation refers to the pipeline's own summary.\n"
            # The clean branch is the one a reviewer takes assurance from, so it
            # carries the same limit as the branch that reports problems.
            "* **This establishes a match to the extracted text, not to the "
            "manuscript.** A value corrupted in extraction is confirmed here "
            "rather than caught. Check numeric values against the PDF, and "
            "read the Extraction lines in the header.\n"
        )
    lines = [
        "\n\n# Citation check",
        "",
        "Checked mechanically against the extracted manuscript text, and "
        "against the report's own phrasing. Treat the following with caution: "
        "a quotation that cannot be found may be a paraphrase in quotation "
        "marks, and a citation of the pipeline's own summary is not evidence "
        "from the paper.",
        "",
        # A reviewer reads this section and reasonably takes a clean result as
        # assurance. It is narrower than that, and the limit is invisible from
        # inside: the model, this check and the question box all read the same
        # extracted text, so a value extraction corrupted is confirmed by all
        # three. One manuscript's sample size of 173 arrived as "116 173" with a
        # marginal line number glued to it, and every mechanism agreed.
        "**This establishes that a quotation matches the extracted text, not "
        "that it matches the manuscript.** Extraction can corrupt a value, and "
        "this check would confirm the corrupted version. Read the Extraction "
        "lines in the header, and check numeric values -- sample sizes, "
        "coefficients, degrees of freedom, p-values, interval limits -- against "
        "the PDF before using them. Anything that exists only in a graphical "
        "figure was never read at all.",
        "",
    ]
    for problem in problems:
        lines.append(f"* {problem}")
    return "\n".join(lines) + "\n"


def tables_for_prompt(table_blocks: List[Tuple[int, str]],
                      source_text: str = "") -> str:
    """
    Render extracted tables for a synthesis prompt.

    The synthesis stages previously saw only the manifest's list of table
    labels, never the values. That is why a report could ask for a credible
    interval that was sitting in the evidence appendix, and then declare the
    same interval unavailable. The numbers have to be in the prompt for a
    claim about them to be worth anything.
    """
    notes = [
        note for note in (
            document_extraction_quality(source_text),
            missing_tables_warning(source_text, table_blocks),
            sample_size_warning(source_text, table_blocks),
        ) if note
    ]
    if not table_blocks:
        return "\n\n".join(notes)
    parts: List[str] = list(notes)
    predictor_summary = summarise_model_predictors(table_blocks)
    if predictor_summary:
        parts.append(predictor_summary)
    total = 0
    for page_num, block_text in table_blocks:
        lines = [ln for ln in block_text.splitlines() if ln.strip()]
        truncated = False
        if len(lines) > MAX_PROMPT_TABLE_ROWS:
            lines = lines[:MAX_PROMPT_TABLE_ROWS]
            truncated = True
        rendered = "\n".join(lines)
        if truncated:
            rendered += "\n... (table truncated for length)"
        warning = table_fragmentation_warning(block_text)
        if warning:
            rendered = warning + "\n" + rendered
        if total + len(rendered) > MAX_PROMPT_TABLE_CHARS:
            parts.append("... (further tables omitted for length)")
            break
        total += len(rendered)
        parts.append(rendered)
    return "\n\n".join(parts)


# Chunk notes were concatenated into the file-level prompt without limit. A
# 127-page submission -- a 22-page article bound with 105 pages of
# supplementary tables, 73% of the document by characters -- produced a 45,904
# token prompt against a 32,768 token server, and the review died with a raw
# traceback partway through. Budget the prompt against the window the server
# actually has, and say what was left out.
PROMPT_CHARS_PER_TOKEN = 3.2      # deliberately pessimistic for mixed prose
GENERATION_HEADROOM_TOKENS = 3000  # room for the answer plus template overhead


def prompt_char_budget(reserve_tokens: int = 0) -> int:
    """Characters a prompt may use before the server would refuse it."""
    try:
        context = available_context()
    except Exception:
        context = 32768
    # A thinking run generates reasoning inside the same window, so the prompt
    # must give that room back or the extra allowance would push the request
    # past the context the server has.
    try:
        thinking_reserve = thinking_token_allowance()
    except Exception:
        thinking_reserve = 0
    usable = (context - GENERATION_HEADROOM_TOKENS - thinking_reserve
              - max(0, reserve_tokens))
    return max(4000, int(usable * PROMPT_CHARS_PER_TOKEN))


def fit_to_budget(text: str, budget: int, unit: str = "section") -> Tuple[str, str]:
    """
    Trim text to a character budget on a boundary, and report what went.

    Returns the kept text and a note naming the loss, empty when nothing was
    dropped. Trimming from the end keeps the article body, which comes first,
    and sheds the supplementary material, which does not.
    """
    if not text or len(text) <= budget:
        return text, ""
    parts = re.split(r"(?=\n#+ )|(?=\n\[Page \d+\])", text)
    kept: List[str] = []
    used = 0
    for part in parts:
        if used + len(part) > budget:
            break
        kept.append(part)
        used += len(part)
    if not kept:                      # no usable boundary; cut bluntly
        kept = [text[:budget]]
        used = budget
    dropped = len(text) - used
    note = (
        f"WARNING: the document is too long for the model's context window. "
        f"{dropped:,} of {len(text):,} characters of {unit} notes "
        f"({dropped / len(text):.0%}) were omitted from this synthesis, taken "
        "from the end of the document. In a submission that bundles "
        "supplementary material after the article, the omitted part is "
        "normally that appendix -- but confirm it, and treat anything absent "
        "from this review as unexamined rather than unreported."
    )
    return "".join(kept), note


def synthesize_file_review(model, tokenizer, file_name: str, combined_chunk_review: str,
                           method_expectations: str = "", manifest_summary: str = "",
                           tables_text: str = "") -> str:
    budget = prompt_char_budget(reserve_tokens=FILE_SYNTHESIS_MAX_TOKENS)
    budget -= len(tables_text) + len(method_expectations) + len(manifest_summary) + 4000
    combined_chunk_review, overflow_note = fit_to_budget(
        combined_chunk_review, max(4000, budget), "chunk"
    )
    context_block = ""
    if overflow_note:
        print(f"[context] {overflow_note}", flush=True)
        context_block += f"\n{overflow_note}\n"
    if method_expectations:
        context_block += f"\n{method_expectations}\n"
    if manifest_summary:
        context_block += f"\nEvidence manifest:\n{manifest_summary}\n"
    if tables_text:
        context_block += (
            "\nTables extracted from this manuscript. These are the paper's own "
            "reported values; treat them as primary evidence and quote from them "
            "where a concern turns on a number:\n"
            f"{tables_text}\n"
        )

    user_text = f"""
{SYSTEM_STYLE}

{COMMON_DIAGNOSTIC_ALIASES}
{context_block}
Create a concise file-level review summary from the chunk notes below.

File: {file_name}

Required headings:
# File synopsis
# Strongly supported strengths
# Strongly supported concerns
# Reporting limits or ambiguities
# Questions raised by this file

Rules:
- If a study uses MCMC or MLwiN without explicit Bayesian terminology (priors, posteriors), treat it as strictly frequentist. Do not suggest the paradigm is unclear and do not mention missing Bayesian elements.
- Do not include complaints about OCR-garbled equations or trivial software version discrepancies.
- Base the summary on the chunk notes and on the extracted tables above.
- Where a table above reports an estimate, interval or fit statistic, use it. Never describe a value as unreported, unavailable or not extractable if it appears in those tables.
- Compare what the tables show against what the narrative claims. A covariate the text says was adjusted for, but which is absent from that model's predictor list, is a reportable inconsistency, as is a model that includes a term its counterparts omit.
- Maximum 4 bullets under any heading.
- Prefer directly supported points over generic reviewer concerns.
- Do not say a table, coefficient, equation, or diagnostic is missing if the chunk notes contain it.
- If something is present but brief, say "reported briefly" or "not fully detailed", not "unreported".
- Do not use LaTeX or dollar-sign notation.
- Do not convert obvious PDF/OCR extraction artefacts into substantive concerns.
- Do not call a reporting choice a contradiction unless two directly conflicting statements are present.
- Do not infer data entry artefacts from small coefficients, exponent notation, or scaling choices alone.
- Finish cleanly after the final heading.
- Distinguish between "not present in the manuscript" and "not recoverably extracted".
- If an item is referenced in the manuscript text but not clearly accessible from extraction, describe it as "present but not fully accessible in the extracted material".
- Treat extraction failure as a review limitation, not automatically as a manuscript fault.
- Match diagnostic expectations to the method classification. Do not request diagnostics from the wrong paradigm.
- Consult the evidence manifest to avoid claiming something is absent when it is recorded as present.

Chunk notes:
\"\"\"
{combined_chunk_review}
\"\"\"
""".strip()
    prompt = apply_chat_template_compat(tokenizer, user_text)
    sampler = make_default_sampler()
    out = generate(
        model, tokenizer, prompt=prompt,
        max_tokens=FILE_SYNTHESIS_MAX_TOKENS, sampler=sampler, verbose=False
    )
    return clean_model_output(out)


def synthesize_report(
    model,
    tokenizer,
    file_summaries: List[Tuple[str, str]],
    all_manifests: Optional[List[EvidenceManifest]] = None,
    tables_text: str = "",
    temperature: float = None,
) -> str:
    """
    Create one integrated internal critical-appraisal memo.

    This version deliberately separates:
    - directly supported concerns,
    - verification prompts,
    - extraction limits.

    The aim is to reduce generic reviewer boilerplate and repeated concerns
    across overlapping headings.
    """
    joined: List[str] = []
    for name, summary in file_summaries:
        joined.append(f"# File: {name}\n{summary}")
    joined_text = "\n\n".join(joined)

    manifest_block = ""
    if all_manifests:
        manifest_parts: List[str] = []
        for m in all_manifests:
            part = m.summary_text()
            if hasattr(m, "model_block_summary"):
                part += "\n\nDetected model/result blocks:\n" + m.model_block_summary()
            manifest_parts.append(part)
        manifest_block = "\n\nEvidence manifests:\n" + "\n\n".join(manifest_parts)

    tables_block = ""
    if tables_text:
        tables_block = (
            "\n\nTables extracted from the manuscript(s). These are the paper's own "
            "reported values. Treat them as primary evidence, quote from them when a "
            "concern turns on a number, and check the narrative against them:\n"
            f"{tables_text}\n"
        )

    user_text = f"""
{SYSTEM_STYLE}

{COMMON_DIAGNOSTIC_ALIASES}
{manifest_block}
{tables_block}

Create one integrated internal critical-appraisal memo from the file summaries below.

Use exactly these headings:

# Overall synopsis
# Major strengths
# Directly supported concerns
# Verification prompts
# Extraction limits
# Overall confidence

Core purpose:
- Write as an internal memo to support human judgement.
- Do not write as a final reviewer report to be pasted directly to authors or editors.
- Prioritise the few paper-specific issues most worth scrutiny.
- Do not produce a checklist-style review.

Strict evidence hierarchy:
- "Directly supported concerns" are concerns clearly grounded in the supplied summaries, tables, or evidence manifest.
- "Verification prompts" are plausible checks motivated by specific evidence, but not confirmed flaws.
- "Extraction limits" are limitations of the extraction or summarisation, not manuscript faults.

Required format:
- Under "Major strengths", each bullet must start with "Strength:" and include "Evidence:". The same evidence rules below apply: quote the manuscript, a table cell or a numeric value, or state the specific feature without quotation marks.
- Under "Directly supported concerns", each bullet must start with "Concern:" and include "Evidence:" and "Why it matters:".
- Do not rank or grade your own concerns. State each one and its evidence; how much each matters is the reviewer's judgement, not yours.
- Under "Verification prompts", each bullet must start with "Check:" and include "Reason:".
- Under "Extraction limits", each bullet must start with "Limit:".
- An impossible or nonsensical value in the extracted material is not
  automatically an extraction fault. A prior given as a uniform distribution on
  a single point, an estimate outside its own confidence interval, or a
  percentage set summing to 163% may be exactly what the manuscript prints. Where
  you cannot tell an extraction artefact from the manuscript's own error, put it
  under "Verification prompts" as a check to settle, not under "Extraction
  limits". A reviewer reads the checks; the limits section is where this
  pipeline records its own shortcomings, and a real fault filed there is lost.
- Do not include concerns that cannot be linked to specific evidence in the supplied material.

Evidence must come from the manuscript:
- Every "Evidence:" line must quote the manuscript, a table cell, or a numeric value, with its page or table where known.
- A quotation must reproduce the source exactly, character for character. If you cannot reproduce the wording exactly, describe it in your own words WITHOUT quotation marks. Never place quotation marks around a paraphrase, a reconstruction, or a plausible-sounding phrase: an invented quotation in a review is worse than no quotation.
- Never cite this pipeline's own intermediate output as evidence. Phrases such as "the file summary notes", "the evidence summary flags", "the manifest indicates" are NOT evidence: they describe a summary of the paper, not the paper. If the only support for a concern is such a phrase, the concern belongs under "Verification prompts", not here.
- If you cannot produce a verbatim quotation or a specific number for a concern, move it to "Verification prompts" or drop it.
- A passage marked [PREVIEW CUT SHORT - DO NOT QUOTE FROM HERE] is an orientation aid, not a source. It stops mid-sentence and often mid-word, so a quotation taken from one names a fragment that does not exist in the paper. Use it to find where to look; quote the manuscript text or a table instead.

Register:
- Describe a problem by its consequence, not with a verdict. Do not write "fundamental error", "fatal flaw", "invalid", "meaningless", "worthless" or "completely undermines". Write what follows from it: "may introduce optimistic bias because ...", "makes the reported comparison difficult to interpret because ...".
- Where the manuscript itself acknowledges a limitation you are raising, say so in the same bullet. A concern that ignores the authors' own discussion of it reads as though the discussion was not consulted.

Claims must not outrun the evidence:
- Do not assert in the synopsis or in a concern anything you are simultaneously asking to be checked under "Verification prompts". If you are unsure whether a feature exists (an interaction term, a covariate, a correction), say so once, in "Verification prompts" only.
- State what the reported model actually contains, judged from the coefficient tables where they are available, rather than what a method of this kind usually contains.

Length limits:
- Overall synopsis: maximum 5 bullets.
- Major strengths: maximum 5 bullets.
- Directly supported concerns: 2-5 bullets.
- Verification prompts: 2-4 bullets.
- Extraction limits: maximum 2 bullets.
- Overall confidence: maximum 3 bullets.
- Do not fill sections mechanically. If only one or two strong points exist, use only one or two.
- Do not make the memo artificially brief. If the evidence contains several distinct paper-specific issues, include them. The aim is not minimalism; the aim is prioritised, evidence-grounded scrutiny.

Method-alignment rules:
- Do not assume the paper uses only one analysis type.
- If the evidence manifests or model/result blocks indicate multiple distinct models, describe the paper as combining those models.
- Prefer model-level evidence over a single paper-level label when they conflict.
- If a specific table or model block clearly names the model type, treat that as strong evidence for the analysis family.
- Do not describe the whole paper as "unclassified" if model/result blocks clearly identify named models.
- The paper-level method classification is only a summary heuristic; when specific model blocks provide clearer evidence, follow the model blocks.
- Treat studies using MCMC/MLwiN/DIC without explicit Bayesian terms as frequentist. Do not hedge, do not suggest the framework is unclear, and do not raise Bayesian-specific concerns.
- Match diagnostic expectations to the classified method and the actual analysis used.

Concern-selection rules:
- Prefer one paper-specific concern over several generic methodological checks.
- Omit plausible but low-priority concerns if stronger paper-specific concerns are available.
- A concern should be included only if it changes how a reviewer, editor, or author would inspect the paper.
- Do not repeat the same concern across different headings.
- Do not include generic checklist items such as "clarify missing data mechanisms", "verify whether interactions were tested", or "check diagnostics" unless specific evidence in this paper motivates them.
- For comparative-index or model-comparison papers, prioritise how "best", "superior", "selected", or "preferred" was operationalised when the text ranks models, predictors, indices, or methods.
- For observational or survey data, distinguish between "verify whether X was handled" and "X was omitted".
- Raise survey-specific checks such as survey weights, clustering, or stratification only when the paper clearly describes a population-level survey or complex sampling frame.
- Do not apply survey-design checks to studies recruiting from clubs, clinics, teams, laboratories, or convenience settings.
- Do not raise medication-adjustment checks unless medication use is directly relevant to the outcome or exposure and is mentioned in the evidence.

Missingness and reporting rules:
- Do not claim that a table, coefficient, equation, diagnostic, statistic, or model detail is missing if it is present in the file summaries, the evidence manifest, or the extracted tables above.
- Before writing any "Extraction limits" bullet, check the extracted tables. Do not record as an extraction limit a value that appears there. Similarly, do not raise a "Check:" asking for a number the tables already give; answer it from the table instead, and report the answer as a concern only if the number warrants one.
- If the evidence manifest says a statistic type is present, for example "Confidence intervals reported: True", do not make blanket claims that it is missing.
- You may say that a statistic is "reported for some analyses but not visible for [specific analysis] in the extracted material" if that is accurate.
- Apply the same principle to standard errors, p-values, effect sizes, model fit statistics, equations, and diagnostics.
- If reporting is present but limited, say "reported briefly" or "not fully detailed", not "unreported".
- If the manuscript text points to an item but extraction is inadequate, say it appears to be present but could not be fully evaluated from the extracted material.
- Treat extraction limitations as limitations of this review unless the evidence supports a stronger claim.

Interpretation rules:
- Do not overclaim.
- State clearly when a point is directly supported versus inferred.
- Do not convert obvious PDF/OCR extraction artefacts into substantive concerns.
- Do not infer data-entry artefacts from small coefficients, exponent notation, scaling choices, or OCR symbols alone.
- Do not call a reporting choice a contradiction unless two directly conflicting statements are present.
- Do not use LaTeX or dollar-sign notation.
- Do not directly address authors. Prefer "Check whether...", "Verify whether...", or "Consider whether...".
- Do not instruct the editor to accept, reject, or request revision.
- Finish cleanly after the final heading.

For model-comparison or prediction papers, explicitly scrutinise:
- how competing models were compared;
- whether reported improvements are practically meaningful;
- whether the paper distinguishes prediction, explanation, and normalisation;
- whether back-transformation from log models is handled appropriately.

Outcome-summarising rules:
- In the overall synopsis, do not compress multiple outcomes into a single directional sentence unless the supplied summaries show the same direction for all of them.
- If different outcomes point in different directions, summarise them separately in plain language.
- If one outcome shows no clear difference, do not merge it into a sentence implying the same direction as other outcomes.

File summaries:
\"\"\"
{joined_text}
\"\"\"
""".strip()

    prompt = apply_chat_template_compat(tokenizer, user_text)
    sampler = make_default_sampler(temperature)
    out = generate(
        model,
        tokenizer,
        prompt=prompt,
        max_tokens=SYNTHESIS_MAX_TOKENS,
        sampler=sampler,
        verbose=False,
    )
    return clean_model_output(out)


def validate_report_against_evidence(
    model,
    tokenizer,
    report_text: str,
    file_summaries: List[Tuple[str, str]],
    programmatic_corrections: Optional[List[str]] = None,
) -> str:
    """
    Validate and recalibrate the integrated report against extracted evidence.

    This pass should remove overstatements, repeated generic concerns, and
    claims that contradict the evidence summaries. It should not add new
    criticisms.
    """
    joined: List[str] = []
    for name, summary in file_summaries:
        joined.append(f"# File: {name}\n{summary}")
    evidence_text = "\n\n".join(joined)

    correction_block = ""
    if programmatic_corrections:
        correction_block = (
            "\n\nPROGRAMMATIC CORRECTIONS "
            "(these have been verified automatically and MUST be applied):\n"
            + "\n".join(f"- {c}" for c in programmatic_corrections)
            + "\n"
        )

    user_text = f"""
You are checking an internal critical-appraisal memo for factual consistency,
proportionality, specificity, and non-repetition against extracted evidence summaries.

Task:
Revise the report only where it:
- incorrectly claims that information is missing, absent, contradictory, erroneous, or unreported;
- overstates a plausible verification issue as a confirmed flaw;
- repeats the same concern in multiple sections;
- includes generic checklist concerns not specifically motivated by this paper;
- applies diagnostics or expectations from the wrong method family.

Do not introduce new criticisms.

{COMMON_DIAGNOSTIC_ALIASES}
{correction_block}

Evidence hierarchy:
- A directly supported concern must be grounded in specific evidence from the summaries, tables, or manifest.
- A verification prompt must be motivated by specific evidence but must not be phrased as a confirmed flaw.
- An extraction limit must be framed as a limitation of the review, not a manuscript fault.

Important correction rules:
- If the evidence mentions MCMC, Gibbs sampling, MLwiN, or DIC without explicit Bayesian terms, treat it as a frequentist model. Remove hedging that the framework is unclear and remove Bayesian-specific diagnostic requests.
- If tables, coefficients, standard errors, confidence intervals, fit statistics, equations, diagnostics, or model details are present in the evidence summaries, do not describe them as missing, absent, or unverifiable.
- If the evidence manifest says a statistic type is present, for example "Confidence intervals reported: True", remove blanket claims that this statistic type is missing.
- You may revise to: "reported for some analyses but not visible for [specific analysis] in the extracted material" if this is supported.
- If something is present but concise, revise wording to "reported briefly", "not fully detailed", or "could not be fully evaluated from the extracted material".
- If a model form appears in the evidence summaries, do not say the functional form is undefined.
- If the evidence summaries refer to an equation, figure, table, appendix, or supplementary item, do not say the item is missing unless the evidence explicitly indicates absence.
- Treat extraction failure as a review limitation rather than evidence that the manuscript omitted the item.

Specificity and priority rules:
- Remove generic boilerplate concerns that could apply to almost any paper unless they are clearly motivated by specific evidence.
- If the same concern appears under multiple headings, keep the best version once and remove or merge the others.
- Prefer paper-specific concerns over general audit prompts.
- Do not fill headings mechanically.
- If the report contains more concerns than are well supported, remove the weakest ones.
- Keep no more than three directly supported concerns and no more than three verification prompts.
- If a concern is plausible but not directly shown, reclassify it as a verification prompt.

Survey and observational-study rules:
- For observational or survey data, revise "The study did not adjust for X" to "Verify whether X was applied" unless the evidence explicitly confirms the omission.
- Remove survey-specific concerns such as survey weights, clustering, or stratification if the paper does not describe a population-level survey or complex sampling frame.
- Do not apply survey-design checks to studies recruiting from clubs, clinics, teams, laboratories, or convenience settings.
- Medication adjustment should be raised only where medication is directly relevant to the outcomes or exposures and is mentioned in the evidence.

Method-specific rules:
- For Bayesian papers, do not request every Bayesian diagnostic by default. If R-hat is reported but ESS is not visible, say "ESS could not be verified in the extracted material", not "MCMC diagnostics are missing".
- For GAMLSS or distributional models, diagnostic checks such as worm plots or randomised quantile residuals may be verification prompts, but do not state they were omitted unless the evidence confirms omission.
- For comparative-index, predictor-ranking, model-selection, or "best model" papers, prioritise whether the ranking criterion is well defined and interpretable.
- For papers ranking predictors by correlations, F-values, information criteria, or fit statistics, avoid generic multiplicity language unless the review specifically explains why multiplicity affects the paper's stated claim.

Style rules:
- Preserve valid critical points.
- Replace overly strong wording with proportionate wording.
- Do not use LaTeX or dollar-sign notation.
- Do not convert OCR artefacts into methodological concerns.
- Do not describe table numbering as inconsistent unless the evidence summaries themselves contain conflicting labels.
- Do not treat notation or terminology conventions as major concerns unless they create genuine inferential ambiguity.
- Return the revised report only.
- Apply all PROGRAMMATIC CORRECTIONS listed above.

Report:
\"\"\"
{report_text}
\"\"\"

Evidence summaries:
\"\"\"
{evidence_text}
\"\"\"
""".strip()

    prompt = apply_chat_template_compat(tokenizer, user_text)
    sampler = make_default_sampler()
    out = generate(
        model,
        tokenizer,
        prompt=prompt,
        max_tokens=VALIDATION_MAX_TOKENS,
        sampler=sampler,
        verbose=False,
    )
    return clean_model_output(out)
def correct_review_contradictions(
    report_text: str,
    combined_text: str,
    table_blocks: list[str],
) -> str:
    """
    Light post-processing pass to soften review claims that directly
    contradict extracted evidence.
    """
    text_low = combined_text.lower()
    tables_low = "\n\n".join(table_blocks).lower()
    full_text = text_low + "\n\n" + tables_low
    revised = report_text

    # Confidence intervals present in tables/text
    if any(k in full_text for k in [
        "95% ci", "confidence interval", "lower bound", "upper bound", " lower ", " upper "
    ]):
        revised = revised.replace(
            "confidence intervals are absent from the extracted text",
            "confidence interval information appears in the extracted material, although fuller reporting may still be desirable",
        )
        revised = revised.replace(
            "does not contain confidence intervals",
            "contains some confidence interval information",
        )
        revised = revised.replace(
            "lacks confidence intervals",
            "does not always present confidence intervals as prominently as it could",
        )
    # Strong table visibility rule:
    # if structured table evidence contains estimates, bounds, and inferential columns,
    # do not allow the review to claim that the main numerical values are not visible.
    table_has_main_numeric_columns = (
        "estimate" in full_text
        and "lower" in full_text
        and "upper" in full_text
        and ("effect size" in full_text or "\tp\t" in full_text or " p " in full_text)
    )

    if table_has_main_numeric_columns:
        replacements = [
            (
                "specific numerical values for confidence intervals and individual data points are not fully recoverable from the extracted text chunks",
                "substantial numerical values, including interval bounds and effect estimates, are recoverable from the extracted tables",
            ),
            (
                "specific numerical values for trimmed means and confidence interval bounds are not visible in the extracted text chunks",
                "trimmed means and confidence interval bounds are visible in the extracted tables",
            ),
            (
                "the actual numerical data are not fully extracted in the provided text",
                "substantial numerical table data are present in the extracted material",
            ),
            (
                "specific numerical values for confidence intervals are not visible in the extracted text",
                "confidence interval values are visible in the extracted tables",
            ),
            (
                "specific numerical values are not visible in the extracted text",
                "specific numerical values are visible in the extracted tables",
            ),
            (
                "not fully recoverable from the extracted text chunks",
                "recoverable from the extracted tables",
            ),
            (
                "not visible in the extracted text chunks",
                "visible in the extracted tables",
            ),
            (
                "preventing a direct verification of the reported summary statistics",
                "allowing at least partial direct verification of the reported summary statistics",
            ),
            (
                "limiting the ability to verify the precision of the reported results",
                "still allowing substantial verification of the precision of the reported results from the extracted tables",
            ),
        ]
        
            # If extracted tables contain named model outputs with estimates and intervals,
    # do not overstate partial extraction or unverifiability.
    table_has_named_model_outputs = (
        (
            "log odds" in full_text
            or "odds ratio" in full_text
            or "coefficient" in full_text
            or "estimated mean" in full_text
        )
        and (
            "95% ci" in full_text
            or "95% hdi" in full_text
            or "credible interval" in full_text
            or "lower" in full_text
        )
    )

    if table_has_named_model_outputs:
        replacements = [
            (
                "are only partly extracted",
                "are extracted well enough to support partial verification",
            ),
            (
                "is only partly extracted",
                "is extracted well enough to support partial verification",
            ),
            (
                "only partly extracted",
                "substantially extracted",
            ),
            (
                "referenced without full data visibility",
                "available with substantial data visibility in the extracted material",
            ),
            (
                "hindering verification of the reported statistics",
                "still allowing substantial verification of the reported statistics",
            ),
            (
                "could not be fully verified",
                "could be partly verified from the extracted tables",
            ),
            (
                "the provided text chunks do not contain the actual data rows",
                "the extracted material contains substantial table rows",
            ),
        ]
        for old, new in replacements:
            revised = revised.replace(old, new)


        # Soften broader claims if they still survive
        revised = revised.replace(
            "the extracted text does not show the numerical results",
            "the extracted tables show substantial numerical results",
        )
        revised = revised.replace(
            "the extracted material does not show the numerical results",
            "the extracted material includes substantial numerical results in the tables",
        )

    # Model specification / equations present
    if any(k in full_text for k in [
        "the following equation",
        "equation was applied",
        "model equation",
        "ln(vo2max",
        "p = w",
        "work-time model",
        "work–time model",
    ]):
        revised = revised.replace(
            "specific model specifications and detailed diagnostics are not fully detailed in the extracted text",
            "the extracted text contains at least a partial model specification, though some details may still be brief",
        )
        revised = revised.replace(
            "model specification is absent",
            "model specification is at least partly present",
        )

    # Standard errors / SD-like reporting present
    if any(k in full_text for k in [
        "\tse\t", " se ", "standard error", "trimmed sd", "sd (0.2)"
    ]):
        revised = revised.replace(
            "standard errors are absent",
            "standard errors are at least partly present in the extracted material",
        )
        revised = revised.replace(
            "the extracted material does not contain standard errors",
            "the extracted material contains at least some standard error or SD-like information",
        )

    # Structured table evidence present: estimates, bounds, p-values, effect sizes
    table_has_ci_like = all(k in full_text for k in ["estimate", "lower", "upper"])
    table_has_inference_like = ("p" in full_text and "effect size" in full_text)

    if table_has_ci_like or table_has_inference_like:
        replacements = [
            (
                "the actual numerical data are not fully extracted in the provided text",
                "substantial numerical table data are present in the extracted material",
            ),
            (
                "preventing a direct verification of the reported summary statistics",
                "allowing at least partial direct verification of the reported summary statistics",
            ),
            (
                "the tables are not accessible",
                "the tables are at least partly accessible in the extracted material",
            ),
            (
                "does not contain confidence intervals",
                "contains confidence interval information in the extracted tables",
            ),
            (
                "confidence intervals are absent from the extracted text",
                "confidence interval information appears in the extracted tables and text",
            ),
            (
                "the extracted material does not contain standard errors",
                "the extracted material contains at least some uncertainty information in the tables",
            ),
        ]
        for old, new in replacements:
            revised = revised.replace(old, new)
                # If parsed text contains subgroup table summaries and explicit group-assignment text,
    # do not let the review overstate their absence.
    subgroup_descriptives_present = (
        ("table 1" in full_text and "±" in full_text and "group" in full_text.lower())
        or (
            "body weight" in full_text.lower()
            and "sprinter group" in full_text.lower()
            and "classics group" in full_text.lower()
        )
    )

    group_assignment_criteria_present = (
        "categorisation occurred due to various factors" in full_text.lower()
        or (
            "subjective coach classification" in full_text.lower()
            and "physical characteristics" in full_text.lower()
            and "previous race programme" in full_text.lower()
        )
    )

    if subgroup_descriptives_present:
        replacements = [
            (
                "specific means and standard deviations for each of the three rider groups are not listed in the provided text",
                "group-specific means and standard deviations are visible in the parsed text, although structured table extraction is limited",
            ),
            (
                "subgroup descriptive statistics are unavailable",
                "subgroup descriptive statistics are at least partly available in the parsed text",
            ),
            (
                "hindering direct comparison of group characteristics",
                "still allowing basic comparison of group characteristics from the parsed text",
            ),
        ]
        for old, new in replacements:
            revised = revised.replace(old, new)

    if group_assignment_criteria_present:
        replacements = [
            (
                "The specific criteria used to categorize riders into GC, Classics, or Sprinter groups are not defined",
                "The manuscript does describe rider categorisation criteria, although these criteria may still warrant clearer justification",
            ),
            (
                "the criteria for rider classification into specialization groups must be explicitly defined",
                "the criteria for rider classification into specialization groups should be justified and described as explicitly as possible",
            ),
            (
                "rider-classification criteria are undefined",
                "rider-classification criteria are described, though partly subjective",
            ),
        ]

        replacements.extend([
            (
                'The specific criteria used to assign riders to "GC," "sprinter," or "classics" groups are not defined in the extracted text',
                "The manuscript does describe rider categorisation criteria, although these criteria may still warrant clearer justification",
            ),
            (
                "the specific criteria used to assign riders",
                "the manuscript does describe the criteria used to assign riders, though the criteria may still be partly subjective",
            ),
            (
                "are not defined in the extracted text",
                "are described in the extracted text, though they may still warrant clearer justification",
            ),
        ])

        for old, new in replacements:
            revised = revised.replace(old, new)

    revised = revised.replace(
        "in the extracted material in the extracted text",
        "in the extracted material",
    )
    revised = revised.replace(
        "in the extracted text in the extracted material",
        "in the extracted material",
    )
    revised = revised.replace(
        "the extracted material in the extracted text",
        "the extracted material",
    )
    revised = revised.replace(
        "in the extracted material in the extracted material",
        "in the extracted material",
    )
    revised = revised.replace(
        "in the extracted text in the extracted text",
        "in the extracted text",
    )

    return revised

# ---------------------------------------------------------------------------
# Output writers
# ---------------------------------------------------------------------------

def write_report(output_dir: Path, input_paths: List[Path], report_text: str,
                 manifests: Optional[List[EvidenceManifest]] = None) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    stem = input_paths[0].stem if len(input_paths) == 1 else "multi_file_review"
    report_path = output_dir / f"{stem}_review_{stamp}.md"
    sources = "\n".join(f"- {p}" for p in input_paths)

    manifest_section = ""
    if manifests:
        manifest_lines = []
        for m in manifests:
            mc = m.method_class.value if m.method_class else "unclassified"
            manifest_lines.append(f"- **{m.source_name}**: method={mc}, tables={m.n_tables}, "
                                  f"model_spec={m.has_model_spec}, equations={m.has_equations}, "
                                  f"SEs={m.has_standard_errors}, variance_components={m.has_variance_components}, "
                                  f"fit_stats={m.has_model_fit_stats}")
        manifest_section = "\nEvidence summary:\n" + "\n".join(manifest_lines) + "\n"

    header = f"""# Local peer-review report

Generated: {datetime.now().isoformat(timespec="seconds")}
Model: {model_display_name()}
Pipeline: {PIPELINE_VERSION}

Input files:
{sources}
{manifest_section}
---

"""
    report_path.write_text(header + report_text, encoding="utf-8")
    return report_path


def write_evidence_appendix(
    output_dir: Path,
    input_paths: List[Path],
    per_file_tables: Dict[str, List[Tuple[int, str]]],
    per_file_chunks: Dict[str, List[str]],
    manifests: Optional[List[EvidenceManifest]] = None,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    stem = input_paths[0].stem if len(input_paths) == 1 else "multi_file_review"
    appendix_path = output_dir / f"{stem}_evidence_appendix_{stamp}.md"

    lines: List[str] = []
    lines.append("# Evidence appendix")
    lines.append("")
    lines.append(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    lines.append(f"Model: {model_display_name()}")
    lines.append(f"Pipeline: {PIPELINE_VERSION}")
    lines.append("")
    lines.append("## Input files")
    for p in input_paths:
        lines.append(f"- {p}")
    lines.append("")

    # Manifest summary section
    if manifests:
        lines.append("## Evidence manifests")
        lines.append("")
        for m in manifests:
            lines.append(f"### {m.source_name}")
            lines.append("")
            lines.append("```text")
            lines.append(m.summary_text())
            lines.append("")
            lines.append("Detected model/result blocks:")
            lines.append(m.model_block_summary())
            lines.append("```")
            lines.append("")

            # Block type breakdown
            type_groups: Dict[str, List[str]] = {}
            for b in m.blocks:
                preview = b.text[:120].replace("\n", " ") + ("..." if len(b.text) > 120 else "")
                if b.model_method_class:
                    preview = f"[{b.model_method_class.value}] {preview}"
                type_groups.setdefault(b.block_type.name, []).append(preview)

            for btype, previews in type_groups.items():
                lines.append(f"**{btype}** ({len(previews)} blocks)")
                lines.append("")
                for prev in previews[:5]:  # Show up to 5 previews per type
                    lines.append(f"- {prev}")
                if len(previews) > 5:
                    lines.append(f"- ... and {len(previews) - 5} more")
                lines.append("")

    for p in input_paths:
        name = p.name
        lines.append(f"# File: {name}")
        lines.append("")

        tables = per_file_tables.get(name, [])
        if tables:
            lines.append("## Extracted tables")
            lines.append("")
            for page_num, block in tables:
                lines.append(f"### Page {page_num}")
                lines.append("")
                lines.append("```text\n" + block + "\n```")
                lines.append("")
        else:
            # No structured tables from pdfplumber — check if inline tables
            # were recovered by detect_inline_tables (stored in manifest).
            manifest_tables = []
            if manifests:
                for m in manifests:
                    if m.source_name == name:
                        manifest_tables = [
                            b for b in m.blocks if b.block_type == BlockType.TABLE
                        ]
                        break

            lines.append("## Extracted tables")
            lines.append("")
            if manifest_tables:
                lines.append(
                    f"No structured tables extracted by PDF parser. "
                    f"{len(manifest_tables)} inline table(s) were recovered "
                    f"from prose text (shown in evidence manifest above)."
                )
            else:
                marker_path = find_marker_output(p) if p.suffix.lower() == ".pdf" else None
                if marker_path:
                    lines.append("No extracted table blocks recorded (Visual parser may have embedded them in chunks).")
                else:
                    lines.append("No extracted table blocks recorded.")
            lines.append("")

        chunks = per_file_chunks.get(name, [])
        if chunks:
            lines.append("## Chunk previews")
            lines.append("")
            for i, chunk_text in enumerate(chunks, start=1):
                preview = chunk_text[:MAX_APPENDIX_CHUNK_PREVIEW]
                if len(chunk_text) > MAX_APPENDIX_CHUNK_PREVIEW:
                    preview += "\n...[truncated preview]..."
                lines.append(f"### Chunk {i}")
                lines.append("")
                lines.append("```text\n" + preview + "\n```")
                lines.append("")
        else:
            lines.append("## Chunk previews")
            lines.append("")
            lines.append("No chunk previews available.")
            lines.append("")

    appendix_path.write_text("\n".join(lines), encoding="utf-8")
    return appendix_path
    
def correct_direction_of_effect_summaries(
    report_text: str,
    combined_text: str,
    table_blocks: list[str],
) -> str:
    """
    Conservative post-processing pass to reduce obvious direction-of-effect
    reversals in summaries when extracted evidence strongly supports a direction.

    This function is intentionally heuristic and only acts when:
    1. paired condition labels are clearly present, and
    2. the extracted text contains directional wording or table-style evidence.
    """
    full_text = (combined_text + "\n\n" + "\n\n".join(table_blocks)).lower()
    revised = report_text

    def has_any(*phrases: str) -> bool:
        return any(p.lower() in full_text for p in phrases)

    # Common paired labels seen in manuscripts
    paired_labels = [
        ("indoor", "outdoor"),
        ("control", "intervention"),
        ("control", "treatment"),
        ("placebo", "treatment"),
        ("placebo", "intervention"),
        ("male", "female"),
        ("pre", "post"),
        ("before", "after"),
        ("baseline", "follow-up"),
        ("baseline", "post"),
        ("rest", "exercise"),
        ("condition a", "condition b"),
    ]

    # Strong general “no difference” cues
    evidence_no_difference = has_any(
        "no difference was observed",
        "no significant difference was observed",
        "a reliable difference could not be established",
        "did not differ",
        "no effect was observed",
        "nonsignificant relative differences",
        "no statistically significant effect",
    )

    # General table-like structure cues
    has_estimate_structure = all(
        key in full_text for key in ["estimate", "lower", "upper"]
    )

    # Build generic directional evidence from paired labels
    for left, right in paired_labels:
        left_in_text = left in full_text
        right_in_text = right in full_text
        if not (left_in_text and right_in_text):
            continue

        evidence_right_higher = has_any(
            f"higher {right} compared with {left}",
            f"greater {right} compared with {left}",
            f"{right} was higher than {left}",
            f"{right} were higher than {left}",
            f"{right} condition was higher",
            f"{right} condition were higher",
        )

        evidence_left_higher = has_any(
            f"higher {left} compared with {right}",
            f"greater {left} compared with {right}",
            f"{left} was higher than {right}",
            f"{left} were higher than {right}",
            f"{left} condition was higher",
            f"{left} condition were higher",
        )

        # If table-style evidence exists and text explicitly states a direction,
        # do not allow the review to assert the opposite direction.
        if has_estimate_structure:
            if evidence_right_higher:
                revised = revised.replace(
                    f"higher {left} compared with {right}",
                    f"higher {right} compared with {left}",
                )
                revised = revised.replace(
                    f"greater {left} compared with {right}",
                    f"greater {right} compared with {left}",
                )
                revised = revised.replace(
                    f"{left} was higher than {right}",
                    f"{right} was higher than {left}",
                )
                revised = revised.replace(
                    f"{left} were higher than {right}",
                    f"{right} were higher than {left}",
                )

            if evidence_left_higher:
                revised = revised.replace(
                    f"higher {right} compared with {left}",
                    f"higher {left} compared with {right}",
                )
                revised = revised.replace(
                    f"greater {right} compared with {left}",
                    f"greater {left} compared with {right}",
                )
                revised = revised.replace(
                    f"{right} was higher than {left}",
                    f"{left} was higher than {right}",
                )
                revised = revised.replace(
                    f"{right} were higher than {left}",
                    f"{left} were higher than {right}",
                )

    # Soften over-strong summary language where evidence explicitly says no difference
    if evidence_no_difference:
        soft_replacements = [
            (
                "was significantly different",
                "was reported as different only where supported by the extracted material",
            ),
            (
                "showed a clear difference",
                "showed a difference only in some outcomes",
            ),
            (
                "demonstrated a clear effect",
                "demonstrated an effect only where supported by the extracted material",
            ),
            (
                "there was a clear difference",
                "there was a difference in some outcomes, while others showed no clear difference",
            ),
        ]
        for old, new in soft_replacements:
            revised = revised.replace(old, new)
            revised = revised.replace("in the extracted material in the extracted text", "in the extracted material")
            revised = revised.replace("in the extracted text in the extracted material", "in the extracted material")

    return revised
    
def apply_method_sensitive_critique_rules(
    report_text: str,
    manifests: list[EvidenceManifest],
) -> str:
    """
    Light post-processing pass to remove or soften critiques that do not
    match the detected methodological framework.
    """
    revised = report_text
    method_classes = {
        str(m.method_class.value if hasattr(m.method_class, "value") else m.method_class)
        for m in manifests
    }

    def replace_many(pairs: list[tuple[str, str]]) -> None:
        nonlocal revised
        for old, new in pairs:
            revised = revised.replace(old, new)
            
    if "bayesian_mixed_effects" in method_classes:
        replace_many([
            (
                "p-values",
                "posterior uncertainty summaries",
            ),
            (
                "fixed and random effects were not clearly distinguished",
                "the hierarchical structure could be reported more explicitly",
            ),
        ])

    if "frequentist_profile_model" in method_classes:
        replace_many([
            (
                "Missing Effect Sizes",
                "Limited Uncertainty Reporting for Modelled Parameters",
            ),
            (
                "Model Validation Gaps",
                "Profile Validation and Filtering Transparency",
            ),
            (
                "tests for multicollinearity between RPM and RPM^2",
                "clearer justification of the fitted polynomial profile and its robustness to sparse 1 Hz data",
            ),
            (
                "Variance Components",
                "Within-Rider Dependence and Profile Estimation",
            ),
            (
                "No variance components or random effects are reported, which is consistent with the frequentist classification but leaves the handling of repeated measures (6 sprints per rider) unexplained regarding within-subject variability.",
                "The analysis is not framed as a variance-components model; the more relevant issue is how repeated sprints per rider were handled when estimating each rider's profile and comparing groups.",
            ),
            (
                "residual diagnostics, homoscedasticity checks, or tests for multicollinearity",
                "clearer profile-validation reporting, including how filtering and sparse data affect the fitted relationships",
            ),
            (
                "Incomplete Model Specification",
                "Partial Reporting of Fitted Profile Parameters",
            ),
        ])

    if "frequentist_mixed_effects" in method_classes:
        replace_many([
            (
                "Levene's test",
                "appropriate mixed-model assumption checks",
            ),
            (
                "ANOVA assumptions",
                "mixed-model assumptions and structure",
            ),
        ])

    if "frequentist_simple_comparison" in method_classes:
        replace_many([
            (
                "variance components",
                "effect sizes and comparison precision",
            ),
            (
                "mixed-effects structure",
                "paired or unpaired comparison structure",
            ),
        ])

    if "frequentist_correlation" in method_classes:
        replace_many([
            (
                "regression diagnostics",
                "checks for outliers, non-linearity, and over-interpretation of association",
            ),
            (
                "predictive accuracy",
                "strength and uncertainty of association",
            ),
        ])

    if "frequentist_robust" in method_classes:
        replace_many([
            (
                "Lack of Variance Component Reporting",
                "Robust-Method Reporting Could Be More Specific",
            ),
            (
                "The absence of variance component reporting limits the interpretability of the repeated measures design, particularly regarding individual variation.",
                "The manuscript could more clearly explain how individual variation should be interpreted within the robust paired-testing framework.",
            ),
            (
                "Variance components",
                "Robust paired-structure reporting",
            ),
        ])

    if "bayesian_model" in method_classes:
        replace_many([
            (
                "the absence of p-values",
                "the need for clearer posterior uncertainty summaries",
            ),
            (
                "frequentist p-values",
                "posterior uncertainty summaries",
            ),
        ])

    if "frequentist_allometric" in method_classes:
        replace_many([
            (
                "Ambiguous Functional Form",
                "Allometric Model Clarification",
            ),
            (
                "does not provide the explicit equations or the exact functional form",
                "reports the equations, though some readers may benefit from clearer exposition of the allometric terms",
            ),
            (
                "the exact functional form needed to verify the allometric modeling choices",
                "clearer explanation of the allometric modeling choices and derived quantities",
            ),
            (
                "raising questions about the exact nature of the curvature and potential multicollinearity",
                "raising questions about the interpretation and stability of the curvature terms",
            ),
            (
                "The study relies on a predictive equation for maturity offset without measuring leg length",
                "The study relies on a predictive maturity-offset equation whose interpretation should be explained carefully",
            ),
        ])

    replace_many([
        ("in the extracted material in the extracted text", "in the extracted material"),
        ("in the extracted text in the extracted material", "in the extracted material"),
        ("the extracted material in the extracted text", "the extracted material"),
    ])

    return revised



def enforce_negative_constraints(report_text: str) -> str:
    """
    Final light-touch cleanup pass.

    Purpose:
    - soften absolute absence claims that are often too strong
    - avoid overclaiming based on extraction limitations
    - keep wording conservative rather than rewriting substance
    """
    revised = report_text

    replacements = [
        (
            "is absent from the extracted text",
            "is not clearly identifiable in the extracted text",
        ),
        (
            "are absent from the extracted text",
            "are not clearly identifiable in the extracted text",
        ),
        (
            "is absent from the extracted material",
            "is not clearly identifiable in the extracted material",
        ),
        (
            "are absent from the extracted material",
            "are not clearly identifiable in the extracted material",
        ),
        (
            "is missing from the extracted text",
            "is not clearly recovered in the extracted text",
        ),
        (
            "are missing from the extracted text",
            "are not clearly recovered in the extracted text",
        ),
        (
            "is missing from the extracted material",
            "is not clearly recovered in the extracted material",
        ),
        (
            "are missing from the extracted material",
            "are not clearly recovered in the extracted material",
        ),
        (
            "not reported",
            "not clearly reported in the extracted material",
        ),
        (
            "not provided",
            "not clearly provided in the extracted material",
        ),
        (
            "cannot be verified",
            "cannot be fully verified from the extracted material",
        ),
        (
            "preventing direct verification",
            "limiting direct verification",
        ),
        (
            "no evidence of",
            "no clearly visible evidence of",
        ),
        (
            "does not report",
            "does not clearly report in the extracted material",
        ),
        (
            "fails to report",
            "does not clearly report in the extracted material",
        ),
        (
            "is not described",
            "is not clearly described in the extracted material",
        ),
        (
            "are not described",
            "are not clearly described in the extracted material",
        ),
        (
            "is not discussed",
            "is not clearly discussed in the extracted material",
        ),
        (
            "is not mentioned",
            "is not visible in the extracted material",
        ),
        (
            "are not mentioned",
            "are not visible in the extracted material",
        ),
        (
            "is omitted",
            "is not visible in the extracted material",
        ),
        (
            "was omitted",
            "was not visible in the extracted material",
        ),
        (
            "were omitted",
            "were not visible in the extracted material",
        ),
        (
            "is lacking",
            "is not clearly present in the extracted material",
        ),
        (
            "are lacking",
            "are not clearly present in the extracted material",
        ),
        (
            "no mention of",
            "no clearly visible mention of",
        ),
        (
            "without any",
            "without clearly visible",
        ),
    ]

    for old, new in replacements:
        revised = revised.replace(old, new)

    revised = revised.replace(
        "the manuscript does not contain",
        "the extracted material does not clearly contain",
    )
    revised = revised.replace(
        "the manuscript lacks",
        "the extracted material does not clearly show",
    )
    revised = revised.replace(
        "there is no",
        "there is no clearly recovered",
    )

    return revised

# ---------------------------------------------------------------------------
# General-purpose chat and query modes
# ---------------------------------------------------------------------------

_DEFAULT_CHAT_SYSTEM = """You are a knowledgeable academic assistant running locally on Apple Silicon.

Core principles:
- Prioritise factual accuracy, verifiable evidence, and academic integrity above fluency or style.
- Accuracy first: prioritise factual and logical soundness over fluency or speed.
- Honesty and uncertainty: if evidence is lacking or mixed, state this clearly and avoid speculation or false confidence.
- Academic integrity: never invent data, quotations, or references.
- Criticality: evaluate claims rigorously, avoid unmerited praise, and acknowledge flaws and limits.

Evidence and reasoning:
- Clearly distinguish between (1) verifiable fact, (2) interpretation, and (3) speculation.
- Cite full, authentic references (author, year, title, outlet, DOI/ISBN) where relevant. Prefer peer-reviewed or primary sources. Never fabricate references.
- Be transparent: show assumptions and logical steps. For quantitative work, include workings or explain reasoning.

Style and format:
- Use a clear academic tone in British English without contractions.
- Write concisely but precisely.
- Structure responses with headings or steps where useful.
- Do not use LaTeX dollar-sign notation. Write mathematics in plain text (e.g., R-squared, beta_1, P(success | theta)).

Code and technical work:
- Provide full, reproducible code where possible.
- Explain parameters and rationale outside code.
- Maintain naming consistency and acknowledge any execution limits, suggesting realistic alternatives."""

_CHAT_MAX_TOKENS = 1200


def _build_chat_prompt(
    tokenizer,
    user_text: str,
    system_text: Optional[str] = None,
    history: Optional[List[dict]] = None,
) -> str:
    """Build a multi-turn chat prompt from message history."""
    messages: List[dict] = []
    if system_text:
        messages.append({"role": "system", "content": system_text})
    if history:
        messages.extend(history)
    messages.append({"role": "user", "content": user_text})

    if getattr(tokenizer, "chat_template", None) is None:
        return user_text

    try:
        return tokenizer.apply_chat_template(
            messages,
            tokenize=False,
            add_generation_prompt=True,
            enable_thinking=thinking_enabled(),
        )
    except TypeError:
        return tokenizer.apply_chat_template(
            messages,
            tokenize=False,
            add_generation_prompt=True,
        )


def run_query_mode(args) -> int:
    """Send a single query and print the response."""
    max_tok = args.max_tokens or _CHAT_MAX_TOKENS
    system_text = args.system or _DEFAULT_CHAT_SYSTEM

    print(f"Loading model: {MODEL_NAME}")
    model, tokenizer = load(MODEL_NAME)
    print(f"Backend: {current_backend()}")

    prompt = _build_chat_prompt(tokenizer, args.query, system_text=system_text)
    sampler = make_default_sampler()
    out = generate(
        model, tokenizer, prompt=prompt,
        max_tokens=max_tok, sampler=sampler, verbose=False,
    )
    print(clean_model_output(out))
    return 0


def run_chat_mode(args) -> int:
    """Interactive multi-turn chat loop."""
    max_tok = args.max_tokens or _CHAT_MAX_TOKENS
    system_text = args.system or _DEFAULT_CHAT_SYSTEM

    print(f"Loading model: {MODEL_NAME}")
    model, tokenizer = load(MODEL_NAME)
    print(f"Backend: {current_backend()}")

    print(f"\nChat mode active  (model: {MODEL_NAME})")
    print("Type your message and press Enter.  Commands:")
    print("  /quit or /exit  — end the session")
    print("  /clear          — reset conversation history")
    print("  /system <text>  — change the system prompt")
    print("  /tokens <n>     — change max response tokens")
    print()

    history: List[dict] = []
    sampler = make_default_sampler()

    while True:
        try:
            user_input = input("You: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\nGoodbye.")
            break

        if not user_input:
            continue

        # Commands
        if user_input.lower() in ("/quit", "/exit"):
            print("Goodbye.")
            break
        if user_input.lower() == "/clear":
            history.clear()
            print("[Conversation history cleared]")
            continue
        if user_input.lower().startswith("/system "):
            system_text = user_input[8:].strip()
            print(f"[System prompt updated: {system_text[:80]}{'...' if len(system_text) > 80 else ''}]")
            continue
        if user_input.lower().startswith("/tokens "):
            try:
                max_tok = int(user_input[8:].strip())
                print(f"[Max tokens set to {max_tok}]")
            except ValueError:
                print("[Invalid number]")
            continue

        prompt = _build_chat_prompt(
            tokenizer, user_input,
            system_text=system_text,
            history=history,
        )
        out = generate(
            model, tokenizer, prompt=prompt,
            max_tokens=max_tok, sampler=sampler, verbose=False,
        )
        response = clean_model_output(out)
        print(f"\nAssistant: {response}\n")

        # Keep history for multi-turn context
        history.append({"role": "user", "content": user_input})
        history.append({"role": "assistant", "content": response})

    return 0


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def main() -> int:
    global MODEL_NAME
    global OUTPUT_DOMAIN

    parser = argparse.ArgumentParser(
        description="Local manuscript review pipeline (llama.cpp GGUF or MLX), "
                    "with optional general chat mode",
    )
    parser.add_argument("inputs", nargs="*", help="One or more files or folders (for review mode)")
    parser.add_argument(
        "--output-dir",
        default=str(Path(__file__).resolve().parent.parent / "reports"),
        help="Directory for markdown reports",
    )
    parser.add_argument(
        "--model",
        default=MODEL_NAME,
        help="Model alias (e.g. qwen38, 35b, 27b, gemma4), a path to a "
             ".gguf file, or an MLX repo id. Run --list-models for the "
             f"full list. Default: {model_display_name(MODEL_NAME)}",
    )
    parser.add_argument(
        "--allow-derived-input",
        action="store_true",
        help="Review a file even when it looks like this pipeline's own output "
             "(a previous review or evidence appendix). Off by default.",
    )
    parser.add_argument(
        "--list-models",
        action="store_true",
        help="List the model aliases accepted by --model, then exit.",
    )
    parser.add_argument(
        "--backend",
        choices=["llama-server", "llama-cpp", "mlx"],
        default=None,
        help="Force a backend. By default a .gguf model uses llama-server "
             "and anything else uses mlx.",
    )
    parser.add_argument(
        "--llama-url",
        default=None,
        help="Base URL of the llama-server instance "
             "(default: http://127.0.0.1:8081).",
    )
    parser.add_argument(
        "--domain",
        default="general",
        help="Domain hint for OCR cleaning (e.g., exercise_physiology)",
    )
    parser.add_argument(
        "--chat",
        action="store_true",
        help="Start an interactive chat session with the model",
    )
    parser.add_argument(
        "--query",
        type=str,
        default=None,
        help="Send a single query to the model and print the response",
    )
    parser.add_argument(
        "--max-tokens",
        type=int,
        default=None,
        help="Override max tokens for --chat or --query mode (default: 1200)",
    )
    parser.add_argument(
        "--system",
        type=str,
        default=None,
        help="Custom system prompt for --chat or --query mode",
    )

    args = parser.parse_args()

    if args.list_models:
        print(list_model_aliases())
        return 0

    MODEL_NAME = resolve_model_alias(args.model)
    if args.llama_url:
        os.environ["LLAMA_SERVER_URL"] = args.llama_url
    if args.backend:
        set_backend(args.backend)

    # --- Chat mode ---
    if args.chat:
        return run_chat_mode(args)

    # --- Single query mode ---
    if args.query:
        return run_query_mode(args)

    # --- Review mode (default) ---
    OUTPUT_DOMAIN = args.domain
    if not args.inputs:
        print("No input files specified. Use --chat for interactive mode or provide files to review.", file=sys.stderr)
        parser.print_help()
        return 1

    input_paths = collect_input_paths(args.inputs)
    if not input_paths:
        print("No supported input files found.", file=sys.stderr)
        return 1

    print(f"Loading model: {MODEL_NAME}")
    model, tokenizer = load(MODEL_NAME)
    print(f"Backend: {current_backend()}")

    file_summaries: List[Tuple[str, str]] = []
    per_file_tables: Dict[str, List[Tuple[int, str]]] = {}
    per_file_chunks: Dict[str, List[str]] = {}
    # The manuscript's own text, kept for the citation check. The summaries
    # must not be used for that: a quotation lifted from a summary would then
    # verify against the summary and the check would confirm nothing.
    per_file_source_text: Dict[str, str] = {}
    all_manifests: List[EvidenceManifest] = []

    for path in input_paths:
        print(f"Reading: {path}")
        try:
            text, table_blocks = load_document(path)
        except Exception as e:
            file_summaries.append((path.name, f"Failed to read file: {e}"))
            per_file_tables[path.name] = []
            per_file_chunks[path.name] = []
            continue

        derived = detect_derived_input(text)
        if derived and not getattr(args, "allow_derived_input", False):
            print(
                f"\n  REFUSING: {path.name} looks like output from this pipeline, "
                f"not a manuscript.\n"
                f"  Found: {', '.join(repr(d) for d in derived[:3])}\n"
                f"  Reviewing a review produces a report that reads normally and "
                f"describes the wrong document.\n"
                f"  Point it at the original manuscript, or pass "
                f"--allow-derived-input if this really is what you meant.\n",
                file=sys.stderr,
            )
            file_summaries.append(
                (path.name, "Skipped: input appears to be output from this pipeline.")
            )
            per_file_tables[path.name] = []
            per_file_chunks[path.name] = []
            continue

        per_file_tables[path.name] = table_blocks
        per_file_source_text[path.name] = text

        # --- Phase 2: Structure evidence ---
        print(f"  Structuring evidence for {path.name}...")
        manifest = structure_evidence(path.name, text, table_blocks)
        all_manifests.append(manifest)

        mc = manifest.method_class
        print(f"  Method classification: {mc.value}")
        print(f"  Evidence: {manifest.n_tables} tables, model_spec={manifest.has_model_spec}, "
              f"equations={manifest.has_equations}, p_values={manifest.has_p_values}")

        # --- Phase 3: Get method-specific expectations ---
        method_expectations = get_method_expectations(
            mc, additional_classes=manifest.additional_method_classes,
        )
        manifest_summary = manifest.summary_text()
        file_tables_text = tables_for_prompt(table_blocks, source_text=text)
        for note in (document_extraction_quality(text),
                     missing_tables_warning(text, table_blocks)):
            if note:
                print(f"  {note.splitlines()[0]}")

        chunks = split_text(text)
        per_file_chunks[path.name] = chunks

        if not chunks:
            file_summaries.append((path.name, "No usable text extracted from this file."))
            continue

        # --- Phase 4: Method-aware chunk review ---
        chunk_outputs = []
        for i, chunk_text in enumerate(chunks, start=1):
            print(f"  Reviewing {path.name} chunk {i}/{len(chunks)}")
            chunk = DocChunk(source_name=path.name, chunk_id=i, text=chunk_text)
            try:
                reviewed = review_chunk(
                    model, tokenizer, chunk,
                    method_expectations=method_expectations,
                    manifest_summary=manifest_summary,
                )
            except Exception as e:
                reviewed = f"Chunk review failed: {e}"
            chunk_outputs.append(reviewed)

        combined = "\n\n".join(f"### Chunk {i}\n{txt}" for i, txt in enumerate(chunk_outputs, start=1))

        # --- Phase 5: File-level synthesis ---
        print(f"  Synthesising file-level review for {path.name}...")
        try:
            file_summary = synthesize_file_review(
                model, tokenizer, path.name, combined,
                tables_text=file_tables_text,
                method_expectations=method_expectations,
                manifest_summary=manifest_summary,
            )
        except Exception as e:
            file_summary = f"# File synopsis\nFile-level synthesis failed: {e}"
        file_summaries.append((path.name, file_summary))

    # --- Phase 5 continued: Final synthesis ---
    print("Synthesising final report...")
    try:
        all_table_blocks = [
            item for tables in per_file_tables.values() for item in tables
        ]
        final_report = synthesize_report(
            model, tokenizer, file_summaries,
            all_manifests=all_manifests,
            tables_text=tables_for_prompt(
                all_table_blocks,
                source_text="\n".join(per_file_source_text.values()),
            ),
        )

        # --- Phase 6: Programmatic post-checks ---
        all_corrections: List[str] = []
        for manifest in all_manifests:
            corrections = programmatic_post_checks(final_report, manifest)
            all_corrections.extend(corrections)

        if all_corrections:
            print(f"  Programmatic checks found {len(all_corrections)} correction(s):")
            for c in all_corrections:
                print(f"    - {c[:100]}...")

        # --- Phase 6 continued: LLM validation ---
        print("Validating final report against evidence...")
        final_report = validate_report_against_evidence(
            model,
            tokenizer,
            final_report,
            file_summaries,
            programmatic_corrections=all_corrections if all_corrections else None,
        )

        # --- Phase 6b: Light contradiction correction ---
        all_text = "\n\n".join(text for _, text in file_summaries)
        all_tables = [tbl for tables in per_file_tables.values() for _, tbl in tables]

        final_report = correct_review_contradictions(
            final_report,
            all_text,
            all_tables,
        )

        # --- Phase 6c: Direction-of-effect consistency guard ---
        final_report = correct_direction_of_effect_summaries(
            final_report,
            all_text,
            all_tables,
        )

        if all_manifests:
            final_report = apply_method_sensitive_critique_rules(
                final_report,
                all_manifests,
            )

        # --- Phase 6d: Conservative wording cleanup ---
        final_report = enforce_negative_constraints(final_report)

        # --- Phase 6e: Remove leaked markdown/math artifacts ---
        final_report = clean_markdown_math_artifacts(final_report)

        stale = stale_module_warning()
        if stale:
            print(f"\n  {stale}\n")
            final_report = f"> **{stale}**\n\n" + final_report

        citation_source = "\n".join(per_file_source_text.values()) + "\n" + "\n".join(
            block for tables in per_file_tables.values() for _, block in tables
        )
        final_report = annotate_concern_confidence(final_report, citation_source)
        report_problems = (verify_report_citations(final_report, citation_source)
                           + evidence_echo_problems(final_report)
                           + overclaim_problems(final_report))
        # Order matters: both checks above read the quotation marks, so the
        # marks may only be removed once they have run.
        final_report = mark_unverified_quotations(final_report, citation_source)
        final_report = report_reliability_banner(final_report) + final_report
        final_report += format_action_list(final_report)
        final_report += format_consistency_check(
            "\n".join(per_file_source_text.values()),
            [t for tables in per_file_tables.values() for t in tables],
        )
        final_report += format_citation_check(report_problems)

    except Exception as e:
        final_report = (
            "# Overall synopsis\n"
            f"Synthesis failed: {e}\n\n"
            "# File-level notes\n\n"
            + "\n\n".join(f"## {name}\n{summary}" for name, summary in file_summaries)
        )
    output_dir = Path(args.output_dir)
    report_path = write_report(output_dir, input_paths, final_report, manifests=all_manifests)
    appendix_path = write_evidence_appendix(
        output_dir, input_paths, per_file_tables, per_file_chunks, manifests=all_manifests,
    )

    print(f"Saved report: {report_path}")
    print(f"Saved evidence appendix: {appendix_path}")
    return 0

def clean_markdown_math_artifacts(report_text: str) -> str:
    """
    Remove LaTeX-style math delimiters that leak into Markdown output.
    Handles both display math ($$...$$) and inline math ($...$).
    Converts common LaTeX constructs to readable plain text.
    """
    revised = report_text

    # --- Step 1: Remove display math blocks ($$...$$) ---
    def _clean_display_math(m):
        content = m.group(1).strip()
        label_match = re.search(r"\\text\{([^}]+)\}", content)
        if label_match:
            label = label_match.group(1)
            return f"[Formula: {label} -- see original document for full expression]"
        return "[Formula -- see original document for full expression]"

    revised = re.sub(r"\$\$([\s\S]*?)\$\$", _clean_display_math, revised)

    # --- Step 2: Clean remaining inline math ($...$) ---
    def _clean_inline_math(m):
        cleaned = m.group(1)
        cleaned = re.sub(r"\\text\{([^}]*)\}", r"\1", cleaned)
        cleaned = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", cleaned)
        cleaned = re.sub(r"\\mathbf\{([^}]*)\}", r"\1", cleaned)
        cleaned = re.sub(r"\\hat\{([^}]*)\}", r"\1", cleaned)
        cleaned = re.sub(r"\\bar\{([^}]*)\}", r"\1", cleaned)
        cleaned = re.sub(r"\\frac\{([^}]*)\}\{([^}]*)\}", r"\1/\2", cleaned)
        cleaned = re.sub(r"\\sqrt\{([^}]*)\}", r"sqrt(\1)", cleaned)
        greek = {
            "\\alpha": "alpha", "\\beta": "beta", "\\gamma": "gamma",
            "\\delta": "delta", "\\epsilon": "epsilon", "\\sigma": "sigma",
            "\\mu": "mu", "\\pi": "pi", "\\theta": "theta", "\\lambda": "lambda",
            "\\eta": "eta", "\\tau": "tau", "\\omega": "omega", "\\phi": "phi",
            "\\chi": "chi", "\\rho": "rho", "\\nu": "nu", "\\kappa": "kappa",
            "\\Sigma": "Sigma", "\\Delta": "Delta", "\\Omega": "Omega",
        }
        for tex, uni in greek.items():
            cleaned = cleaned.replace(tex, uni)
        cleaned = re.sub(r"\\[,;]|\\q(?:uad)?", " ", cleaned)
        cleaned = cleaned.replace("\\cdot", "*").replace("\\times", "x")
        cleaned = cleaned.replace("\\leq", "<=").replace("\\geq", ">=").replace("\\neq", "!=")
        cleaned = cleaned.replace("\\int", "integral").replace("\\sim", "~")
        cleaned = cleaned.replace("\\approx", "approx.").replace("\\infty", "infinity")
        cleaned = re.sub(r"\^\{([^}]*)\}", r"^\1", cleaned)
        cleaned = re.sub(r"_\{([^}]*)\}", r"_\1", cleaned)
        cleaned = re.sub(r"\\(\w+)", r"\1", cleaned)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        return cleaned

    # Handle inline math -- allow up to 500 chars
    revised = re.sub(r"\$([^$\n]{1,500})\$", _clean_inline_math, revised)

    # --- Step 3: Common exponent/inequality cleanups ---
    revised = revised.replace("R^2", "R-squared")
    revised = revised.replace("r^2", "r-squared")
    revised = revised.replace("$<$", "<").replace("$>$", ">")
    revised = revised.replace("$<0$", "<0").replace("$>0$", ">0")

    return revised
    
if __name__ == "__main__":
    raise SystemExit(main())


